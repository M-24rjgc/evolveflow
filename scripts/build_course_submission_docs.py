from __future__ import annotations

import argparse
import json
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


FONT_EAST_ASIA = "宋体"
FONT_LATIN = "Times New Roman"
HEADING_EAST_ASIA = "黑体"

TEMPLATE_DIR = Path("output/course_templates_ascii")
IMAGE_DIR = Path("output/course_images_ascii")
OUT_ASCII_DIR = Path("output/course_docs_ascii")

TEAM = {
    "leader": "孟天赐（2410250732，队长）",
    "requirements": "杨瑞熙（2410250115）",
    "frontend": "王雅娴（2410250575）",
    "backend": "王研博（2410250646）",
    "ai": "姚玮（2410250848）",
    "test": "李思晴（2410250609）",
    "desktop": "冯留杨（2410250311）",
    "docs": "闫艺馨（2410250846）",
}

TEAM_ASSIGNMENT_ROWS = [
    [TEAM["leader"], "队长、项目统筹、集成验收", "负责范围控制、进度协调、关键风险处理、AI agent 与桌面端最终联调、最终交付确认。", "项目计划、验收清单、最终演示和发布决策"],
    [TEAM["requirements"], "需求分析与用例建模", "负责用户场景整理、执行者识别、用例文档、功能需求矩阵、非功能需求和验收口径统一。", "需求规格说明书需求章节、用例清单、需求追踪口径"],
    [TEAM["frontend"], "桌面前端与交互实现", "负责 Today、Tasks、Settings 等主要页面的组件拆分、表单状态、错误提示、Toast 和可用性细节。", "React/Tauri 前端页面、交互检查记录"],
    [TEAM["backend"], "存储层、领域层与能力注册", "负责 SQLite schema、迁移、任务事件服务、排程基础能力、动作日志、撤销和 capability registry。", "storage/domain/capabilities 代码与单元测试"],
    [TEAM["ai"], "AI runtime 与工具调用闭环", "负责模型客户端、工具 schema、上下文构建、流式事件、tool_use/tool_result 回传和 sidecar AI 接口。", "runtime/sidecar AI 能力、连通性和工具调用验证"],
    [TEAM["desktop"], "Tauri 桌面后端与安装交付", "负责 Tauri 命令桥接、sidecar 生命周期、桌面权限、release 构建、安装包验证和运行环境检查。", "桌面后端、发布配置、安装验证记录"],
    [TEAM["test"], "测试计划、执行与缺陷闭环", "负责自动化命令测试、功能场景测试、边界测试、回归记录、测试报告结论和缺陷跟踪。", "测试报告、测试用例结果、回归检查记录"],
    [TEAM["docs"], "课程文档、图表与视觉 QA", "负责四份 DOCX 模板结构维护、三线表排版、图示组织、目录页码、敏感信息扫描和页面视觉检查。", "最终 DOCX、图文排版和文档质量检查记录"],
]

TEAM_OVERVIEW = (
    f"项目组采用统一分工口径：{TEAM['leader']}负责总体统筹和最终验收，"
    f"{TEAM['requirements']}负责需求与用例，{TEAM['backend']}负责存储、领域服务和能力注册，"
    f"{TEAM['ai']}负责 AI runtime 与工具调用闭环，{TEAM['frontend']}负责桌面前端交互，"
    f"{TEAM['desktop']}负责 Tauri 桌面后端与安装交付，{TEAM['test']}负责测试计划和缺陷闭环，"
    f"{TEAM['docs']}负责课程文档、图表和视觉质量。四份文档中涉及人员责任、测试执行和交付验收时均采用这一套姓名与学号口径。"
)

DOCS = {
    "01": {
        "template": "t01.docx",
        "ascii_name": "doc01.docx",
        "final_name": "01_项目开发计划.docx",
        "toc": [
            (1, "1．引言"),
            (2, "1.1编写目的"),
            (2, "1.2项目背景"),
            (2, "1.3定义"),
            (2, "1.4参考资料"),
            (1, "2．项目概述"),
            (2, "2.1工作内容"),
            (2, "2.2条件与限制"),
            (2, "2.3产品"),
            (2, "2.4运行环境"),
            (2, "2.5服务"),
            (2, "2.6验收标准"),
            (1, "3．实施计划"),
            (2, "3.1任务分解"),
            (2, "3.2进度"),
            (2, "3.3预算"),
            (2, "3.4关键问题"),
            (1, "4．人员组织及分工"),
            (1, "5．交付期限"),
            (1, "6．专题计划要点"),
        ],
    },
    "02": {
        "template": "t02.docx",
        "ascii_name": "doc02.docx",
        "final_name": "02_需求规格说明书.docx",
        "toc": [
            (1, "1．引言"),
            (2, "1.1编写目的"),
            (2, "1.2项目背景"),
            (2, "1.3定义"),
            (2, "1.4参考资料"),
            (1, "2．需求概述"),
            (2, "2.1目标"),
            (2, "2.2运行环境"),
            (2, "2.3条件与限制"),
            (1, "3．功能需求"),
            (2, "3.1确定执行者"),
            (2, "3.2确定用例"),
            (2, "3.3编写用例文档"),
            (1, "4．非功能需求"),
            (2, "4.1性能需求"),
            (2, "4.2安全需求"),
            (1, "5．故障处理"),
            (1, "6．其它需求"),
        ],
    },
    "03": {
        "template": "t03.docx",
        "ascii_name": "doc03.docx",
        "final_name": "03_系统设计说明书.docx",
        "toc": [
            (1, "1．引言"),
            (2, "1.1编写目的"),
            (2, "1.2项目背景"),
            (2, "1.3定义"),
            (2, "1.4参考资料"),
            (1, "2．建立分析对象模型"),
            (1, "3．提供交互界面的类"),
            (1, "4．建立动态模型"),
            (1, "5．数据结构设计"),
            (1, "6．用户界面设计"),
        ],
    },
    "04": {
        "template": "t04.docx",
        "ascii_name": "doc04.docx",
        "final_name": "04_测试报告.docx",
        "toc": [
            (1, "1．引言"),
            (2, "1.1编写目的"),
            (2, "1.2项目背景"),
            (2, "1.3定义"),
            (2, "1.4参考资料"),
            (1, "2．测试计划执行情况"),
            (2, "2.1测试项目"),
            (2, "2.2测试机构和人员"),
            (2, "2.3测试结果"),
            (1, "3．软件需求测试结论"),
            (1, "4．评价"),
            (2, "4.1软件能力"),
            (2, "4.2缺陷和限制"),
            (2, "4.3建议"),
            (2, "4.4测试结论"),
        ],
    },
}


@dataclass
class TableSpec:
    caption: str
    headers: list[str]
    rows: list[list[str]]
    widths: list[float]


@dataclass
class ImageSpec:
    filename: str
    caption: str
    width: float = 5.35


def twips(inches: float) -> int:
    return int(round(inches * 1440))


def fitted_widths(widths: Iterable[float], max_width: float = 5.72) -> list[float]:
    widths_list = list(widths)
    total = sum(widths_list)
    if not widths_list or total <= max_width:
        return widths_list
    scale = max_width / total
    return [round(width * scale, 3) for width in widths_list]


def set_run_font(run, size: float | None = 10.5, bold: bool | None = None, east_asia: str = FONT_EAST_ASIA) -> None:
    run.font.name = FONT_LATIN
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), FONT_LATIN)
    r_fonts.set(qn("w:hAnsi"), FONT_LATIN)


def set_cell_margins(cell, top: int = 90, start: int = 110, bottom: int = 90, end: int = 110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **borders: dict[str, str]) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "bottom", "left", "right"):
        edge_data = borders.get(edge, {"val": "nil"})
        tag = qn(f"w:{edge}")
        element = tc_borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), edge_data.get("val", "single"))
        if edge_data.get("val") != "nil":
            element.set(qn("w:sz"), edge_data.get("sz", "6"))
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), "000000")


def set_table_geometry(table, widths: Iterable[float]) -> None:
    widths_list = fitted_widths(widths)
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(twips(sum(widths_list))))

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_list:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(twips(width)))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_list):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(twips(width)))


def set_header_repeat(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def remove_table_grid_borders(table) -> None:
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"})
        if row_idx == 0:
            for cell in row.cells:
                set_cell_border(
                    cell,
                    top={"val": "single", "sz": "10"},
                    bottom={"val": "single", "sz": "8"},
                    left={"val": "nil"},
                    right={"val": "nil"},
                )
        if row_idx == len(table.rows) - 1:
            for cell in row.cells:
                set_cell_border(cell, bottom={"val": "single", "sz": "10"}, left={"val": "nil"}, right={"val": "nil"})


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal.font.size = Pt(10.5)
    if normal._element.rPr is None:
        normal._element.get_or_add_rPr()
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[name]
        style.font.name = FONT_LATIN
        style.font.bold = True
        if style._element.rPr is None:
            style._element.get_or_add_rPr()
        style._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_EAST_ASIA)


def body_paragraph(doc: Document, text: str = "", first_line: bool = True):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.paragraph_format.first_line_indent = Pt(21) if first_line else Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(4)
    if text:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def labeled_paragraph(doc: Document, label: str, text: str):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.paragraph_format.first_line_indent = Pt(21)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(label)
    set_run_font(run, bold=True)
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def heading(doc: Document, text: str, level: int):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, size=None, bold=True, east_asia=HEADING_EAST_ASIA)
    return paragraph


def caption(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    set_run_font(run, size=9.5)
    return paragraph


def add_image(doc: Document, spec: ImageSpec) -> None:
    image_path = IMAGE_DIR / spec.filename
    if not image_path.exists():
        return
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(spec.width))
    caption(doc, spec.caption)


def add_table(doc: Document, spec: TableSpec) -> None:
    caption(doc, spec.caption)
    table = doc.add_table(rows=1, cols=len(spec.headers))
    table.style = "Normal Table"
    table.autofit = False
    display_widths = fitted_widths(spec.widths)
    dense_table = len(spec.headers) >= 4 or sum(spec.widths) > 5.72
    table_font_size = 8.8 if dense_table else 9.5
    set_table_geometry(table, display_widths)
    set_header_repeat(table.rows[0])
    for cell, text in zip(table.rows[0].cells, spec.headers):
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        run = paragraph.add_run(text)
        set_run_font(run, size=table_font_size, bold=True)
    for row in spec.rows:
        cells = table.add_row().cells
        for idx, (cell, value) in enumerate(zip(cells, row)):
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.12
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 or len(value) <= 12 else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(value)
            set_run_font(run, size=table_font_size)
    remove_table_grid_borders(table)
    body_paragraph(doc, "", first_line=False)


def clear_template_keep_title(doc: Document, toc_entries: list[tuple[int, str]], toc_pages: dict[str, str]) -> None:
    title_para = next((p for p in doc.paragraphs if p.text.strip()), None)
    if title_para is None:
        raise RuntimeError("template lacks title paragraph")
    title_element = title_para._element

    for table in list(doc.tables):
        table._element.getparent().remove(table._element)
    for paragraph in list(doc.paragraphs):
        if paragraph._element is not title_element:
            paragraph._element.getparent().remove(paragraph._element)

    doc.add_paragraph()
    for level, text in toc_entries:
        paragraph = doc.add_paragraph(style=f"toc {level}")
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(f"{text}\t{toc_pages.get(text, '1')}")
        set_run_font(run, size=10.5)
    page_break = doc.add_paragraph()
    page_break.paragraph_format.first_line_indent = Pt(0)
    page_break.add_run().add_break(WD_BREAK.PAGE)


def ensure_no_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    for existing in list(settings.findall(qn("w:updateFields"))):
        settings.remove(existing)


def add_items(doc: Document, items: list[Any]) -> None:
    for item in items:
        if isinstance(item, str):
            body_paragraph(doc, item)
        elif isinstance(item, tuple) and item and item[0] == "label":
            labeled_paragraph(doc, item[1], item[2])
        elif isinstance(item, TableSpec):
            add_table(doc, item)
        elif isinstance(item, ImageSpec):
            add_image(doc, item)
        else:
            raise TypeError(f"unsupported item: {item!r}")


def write_sections(doc: Document, sections: list[tuple[int, str, list[Any]]]) -> None:
    for level, title, items in sections:
        heading(doc, title, level)
        add_items(doc, items)


def references_table(doc_type: str) -> TableSpec:
    base_rows = [
        ["1", "Roger S. Pressman, Bruce R. Maxim", "Software Engineering: A Practitioner's Approach, 9th Edition", "McGraw-Hill, 2020", "用于项目过程、质量保证和风险管理参考"],
        ["2", "Ian Sommerville", "Software Engineering, 10th Edition", "Pearson, 2015", "用于需求、体系结构和工程管理参考"],
        ["3", "ISO/IEC/IEEE", "ISO/IEC/IEEE 29148:2018 Systems and software engineering -- Requirements engineering", "2018", "用于需求规格、需求可追踪和验收口径参考"],
        ["4", "ISO/IEC/IEEE", "ISO/IEC/IEEE 12207:2017 Systems and software engineering -- Software life cycle processes", "2017", "用于生命周期过程、交付物和配置管理参考"],
        ["5", "IEEE Computer Society", "IEEE Std 1012-2016 System, Software, and Hardware Verification and Validation", "2016", "用于验证确认、测试证据和评价结论参考"],
        ["6", "ISO/IEC/IEEE", "ISO/IEC/IEEE 29119-3:2021 Software and systems engineering -- Software testing -- Test documentation", "2021", "用于测试文档、测试项和测试结果记录参考"],
        ["7", "SQLite Consortium", "SQLite Documentation: Write-Ahead Logging, PRAGMA statements, file format", "官方文档", "用于本地数据库、事务和备份策略参考"],
        ["8", "Tauri Contributors", "Tauri 2.0 Documentation", "官方文档", "用于桌面应用架构、sidecar 和打包发布参考"],
        ["9", "Meta Open Source", "React Documentation", "官方文档", "用于组件化界面、状态驱动渲染和错误边界参考"],
        ["10", "Microsoft", "TypeScript Handbook", "官方文档", "用于类型系统、接口定义和工程约束参考"],
        ["11", "OpenJS Foundation", "Node.js Documentation", "官方文档", "用于 sidecar 运行时、文件系统和进程通信参考"],
        ["12", "Anthropic", "Messages API Documentation", "官方文档", "用于模型消息、工具调用和流式响应接口参考"],
        ["13", "Vitest Contributors", "Vitest Documentation", "官方文档", "用于单元测试、工作区测试和覆盖分析参考"],
    ]
    if doc_type == "plan":
        rows = base_rows[:10]
    elif doc_type == "requirements":
        rows = [base_rows[i] for i in [1, 2, 3, 6, 7, 8, 9, 11]]
    elif doc_type == "design":
        rows = [base_rows[i] for i in [0, 1, 3, 6, 7, 8, 9, 10, 11]]
    else:
        rows = [base_rows[i] for i in [2, 4, 5, 6, 7, 8, 12]]
    return TableSpec("表 1 参考资料清单", ["序号", "作者或机构", "资料名称", "出版或来源", "本文档采用方式"], rows, [0.45, 1.35, 2.25, 1.05, 1.55])


def plan_sections() -> list[tuple[int, str, list[Any]]]:
    return [
        (1, "1．引言", []),
        (2, "1.1编写目的", [
            "本文档用于说明 EvolveFlow 智能日程助理项目的开发目标、工作范围、组织方式、进度安排、资源投入、风险控制和交付验收依据。项目开发计划既是开发团队执行工作的计划性文件，也是课程评审、设计审查和后续测试验收的共同基线，因此文档尽量把“做什么、由谁做、何时完成、如何验证”写成可以落地执行的安排。",
            "EvolveFlow 的目标不是制作一个只能录入任务的普通待办工具，而是形成一个本地优先、具有真实模型推理与工具调用能力的桌面日程助手。开发计划需要同时覆盖传统软件工程工作和 AI agent 工程工作，包括需求分析、领域服务、SQLite 存储、Tauri 桌面端、Node sidecar、模型消息接口、能力注册表、工具调用闭环、审计撤销、备份恢复、测试与打包发布等内容。",
            f"本文档的主要读者包括{TEAM['leader']}、{TEAM['requirements']}、设计与开发成员、{TEAM['test']}、课程评审人员以及后续可能接手维护的同学。阅读者可以据此了解项目开发的总体节奏，判断各项活动之间的依赖关系，并在阶段检查时根据验收标准对已完成产物进行客观核对。",
        ]),
        (2, "1.2项目背景", [
            "个人学习者和知识工作者常常同时面对课程、会议、项目、复习、临时事务等多类任务，传统待办清单只能记录事项，日历工具只能呈现固定时间块，二者之间缺少对任务时长、截止日期、精力分布、上下文偏好和不可移动安排的综合理解。EvolveFlow 选择智能日程作为课程项目主题，是因为它同时具备明确的业务边界、可验证的数据状态和适合 AI agent 发挥作用的自然语言入口。",
            "本项目的开发单位为 EvolveFlow 项目组，项目性质为软件工程课程综合实践。系统面向单机桌面场景运行，核心数据保存在本地 SQLite 数据库中，AI 运行时以 sidecar 方式与 Tauri 桌面端协同工作。与外部系统的关系主要体现在模型服务的调用、桌面系统通知、文件备份导出以及未来可能扩展的日历同步接口，现阶段不依赖中心化服务器保存用户数据。",
            "项目背景中的关键工程约束是安全和可控。AI 可以理解用户意图、调用工具、提出安排建议，但所有数据变更都必须经过本地能力注册表、领域服务校验和动作日志记录；用户锁定的任务、事件或手动安排必须被尊重；任何可变更操作都应保留撤销依据。该设计使系统能够展示真实 AI agent 能力，同时避免把模型输出直接当作数据库事实写入。",
            TableSpec("表 2 项目关系与边界说明", ["关系对象", "关系说明", "边界约束"], [
                ["课程实践", "按照软件工程课程提交项目开发计划、需求规格、系统设计和测试报告", "文档结构沿用课程模板，内容围绕实际工程产物展开"],
                ["桌面用户", "用户通过 Today、Tasks、Calendar、AI、Analytics、Settings 等页面管理个人日程", "用户数据以本地保存为主，敏感配置不写入文档和界面说明"],
                ["AI 模型服务", "运行时通过兼容 Messages API 的模型接口完成理解、推理、工具选择和回复生成", "模型不直接访问数据库，所有动作经 capability registry 执行"],
                ["本地数据库", "SQLite 承载任务、事件、日程块、提醒、偏好、动作日志、AI 会话和洞察", "启用事务、外键、WAL、索引和备份校验"],
                ["桌面运行环境", "Tauri v2 负责窗口、IPC、sidecar 生命周期和安装包", "发布产物需通过桌面启动、功能流和打包检查"],
            ], [1.2, 2.8, 2.0]),
        ]),
        (2, "1.3定义", [
            "为了避免开发、测试和评审阶段对核心概念产生偏差，本节对项目中高频出现的术语作统一说明。术语定义既来自通用软件工程语境，也来自 EvolveFlow 自身的领域模型，其中 AI agent、能力注册表、动作日志和锁定项是理解系统安全边界的重点。",
            TableSpec("表 3 术语与缩写定义", ["术语", "定义"], [
                ["AI agent", "由模型负责理解意图和选择工具，由本地能力层负责执行动作的智能助理形态；本项目不把固定规则或关键词匹配称为 AI agent。"],
                ["Capability", "系统对外暴露的一项可调用能力，如 task.create、schedule.plan_day、backup.create 等，包含输入 schema、权限语义和处理函数。"],
                ["Sidecar", "由 Tauri 桌面端启动的 Node.js 运行时进程，负责 AI 会话、上下文构建、工具调用循环和 JSON-RPC 通信。"],
                ["Local-first", "业务数据优先保存在用户本机，网络模型服务只参与推理和回复，不作为主要数据存储位置。"],
                ["Schedule Block", "日程中的时间块，可以绑定任务或事件，并带有日期、开始时间、结束时间、锁定标记和手动安排标记。"],
                ["Locked Item", "用户明确锁定的任务、事件或日程块，自动排程和 AI 调整不得随意修改其关键字段。"],
                ["Action Log", "记录能力调用前后状态、操作者、来源、输入快照和撤销分组的审计数据。"],
                ["Undo", "基于动作日志中的状态快照回滚用户或 AI 造成的数据变更。"],
                ["Dream Insight", "系统在长期使用后对偏好和行为模式形成的可解释洞察，用于改进提醒、建议和安排策略。"],
                ["WAL", "SQLite Write-Ahead Logging 模式，用于提升本地数据库读写并发和恢复可靠性。"],
            ], [1.35, 5.2]),
        ]),
        (2, "1.4参考资料", [
            "本文档参考的软件工程资料主要用于确定开发计划的组成要素、生命周期活动、风险管理方法、需求与验证关系，以及桌面端、本地数据库和 AI 工具调用相关的工程实践。参考资料以公开出版书籍、国际标准和官方技术文档为主，不把项目自身 README、源代码或仓库说明作为项目计划的外部参考文献。",
            "在项目实施过程中，团队仍会阅读内部架构说明和代码，但它们属于项目产物或工作资料，不列入本节的参考资料清单。这样处理可以保证本计划的依据来自可独立查验的工程规范和技术来源，也符合课程文档中对参考资料的要求。",
            references_table("plan"),
        ]),
        (1, "2．项目概述", []),
        (2, "2.1工作内容", [
            "EvolveFlow 的工作内容分为产品需求、系统设计、核心实现、桌面体验、AI agent、测试验证和交付打包七个方面。产品需求阶段要把“用户如何描述任务、系统如何理解上下文、日程如何自动安排、用户如何纠错和撤销”转化为可验证的功能需求。系统设计阶段要建立本地数据库、领域服务、能力注册表、sidecar、桌面页面和模型接口之间的关系，明确哪些模块可以变更数据，哪些模块只负责展示或解释。",
            "核心实现阶段以 TypeScript 工作区为主体，完成 storage、domain、capabilities、runtime、cli、desktop-tauri 等包的协同。存储层负责 SQLite 初始化、迁移、WAL、备份、导出和清理；领域层负责任务、事件、提醒、日程、偏好、总结、记忆投影和撤销；能力层负责把领域服务封装成可审计、可幂等、可被 AI 和 UI 共用的能力；运行时负责模型消息、工具定义、会话上下文、流式输出和工具执行循环；桌面端负责用户可见的页面、组件、状态、降级提示和设置。",
            "AI agent 的开发不是附加聊天窗口，而是项目的核心工作之一。模型需要能够读取当前任务、事件、日程和偏好上下文，判断是否需要调用工具，在工具返回后继续组织自然语言回复。计划中要求模型相关功能必须通过真实模型接口和工具调用机制验证，不能把关键词匹配、固定模板回复或静态规则当作 agent 能力。",
            ImageSpec("arch.png", "图 1 EvolveFlow 项目总体架构图", 5.55),
            TableSpec("表 4 主要工作内容与产出", ["工作项", "主要内容", "阶段产出"], [
                ["需求分析", "明确用户角色、业务场景、功能需求、非功能需求和故障处理要求", "需求规格说明书、用例描述、需求追踪矩阵"],
                ["架构设计", "划分桌面端、AI runtime、能力层、领域层、存储层和备份恢复边界", "系统设计说明书、模块接口说明、数据结构设计"],
                ["核心领域实现", "实现任务、事件、日程块、提醒、偏好、动作日志和撤销逻辑", "可测试的 domain 与 storage 包"],
                ["AI agent 实现", "实现模型消息、工具定义、上下文构建、流式响应、工具调用和结果回写", "runtime sidecar、AI 页面、连通性检查"],
                ["桌面体验实现", "实现 Today、Tasks、Calendar、AI、Analytics、Settings、Onboarding、BackupPanel 等页面组件", "Tauri 桌面应用与交互流程"],
                ["测试验证", "执行单元、集成、构建、类型检查、桌面启动和打包验证", "测试报告、缺陷记录、验收证据"],
                ["交付发布", "整理安装包、文档、运行说明和课程提交材料", "最终提交版 DOCX、可运行应用和安装产物"],
            ], [1.15, 3.1, 2.35]),
        ]),
        (2, "2.2条件与限制", [
            "项目已经具备较完整的工程基础：仓库采用 npm workspaces 管理多包，核心业务以 TypeScript 实现，桌面端采用 Tauri v2、React 和 Vite，数据库采用 better-sqlite3，测试以 Vitest 为主，Rust 工具链用于 Tauri 后端和安装包构建。这些条件使项目可以围绕真实桌面产品推进，而不是停留在演示型网页或单脚本原型。",
            "开发限制主要来自三个方面。第一是本地优先目标要求数据库、备份、撤销和隐私保护必须优先于炫技式功能，任何 AI 生成内容都不能绕过本地能力层直接写库。第二是模型服务存在网络和额度约束，因此测试要区分模型连通性测试、工具调用测试和离线降级测试，不能把一次模型失败误判为全部业务失败。第三是课程交付需要文档、代码、运行截图和安装包共同成立，开发计划必须给文档整理和视觉检查留出真实时间。",
            "团队尚需持续创造的条件包括稳定的模型服务配置、Windows 桌面运行环境、Tauri 打包依赖、视觉审查时间以及足够的测试数据。特别是 AI agent 场景需要准备带有冲突、锁定、截止日期、手动安排和偏好的数据集，用来验证模型建议与能力执行是否一致。仅靠空库聊天无法证明系统已经达到预期。",
            TableSpec("表 5 开发条件与限制", ["类别", "已具备条件", "限制或需补充条件"], [
                ["人员", "已明确 8 名成员的姓名、学号和具体工作包", "课程周期有限，需要孟天赐统筹每日同步，避免责任交叉处无人验收"],
                ["技术", "TypeScript、React、Tauri、SQLite、Vitest 工程链路已经建立", "打包和桌面验证依赖 Windows 本机环境"],
                ["数据", "本地 schema 覆盖任务、事件、日程、提醒、偏好、AI 会话和动作日志", "需要构造足够多的边界测试数据"],
                ["AI", "运行时具备 Messages API、工具定义、会话和 sidecar 通信设计", "模型连通性、token 成本和失败重试需要单独验证"],
                ["质量", "已有类型检查、单元测试、构建和 cargo check 检查点", "仍需补充端到端桌面体验与真实用户视角审查"],
                ["交付", "可输出 DOCX 文档、桌面构建和安装包", "最终提交前必须完成视觉检查和敏感信息检查"],
            ], [1.0, 3.0, 2.6]),
        ]),
        (2, "2.3产品", [
            "本项目的最终产品包括可运行的 EvolveFlow 桌面应用、配套 runtime sidecar、核心 TypeScript 包、SQLite 本地数据库结构、备份导出能力、测试用例与测试报告，以及课程要求的四份提交文档。产品交付的重点是可直接运行、可验证、可维护，而不是只展示少量页面的演示稿。",
        ]),
        (3, "2.3.1程序", [
            "应交付的程序分为桌面应用、运行时、共享能力包和辅助命令四类。桌面应用由 Tauri 负责承载，用户通过图形界面完成任务、日历、AI 对话和设置操作；运行时作为 sidecar 进程负责模型和工具调用；共享能力包为 UI、AI、CLI 提供同一套业务接口；辅助命令用于开发、测试、构建、备份和清理。",
            TableSpec("表 6 程序交付清单", ["程序或包", "语言与框架", "存储形式", "交付说明"], [
                ["apps/desktop-tauri", "React、Vite、Tauri v2、Rust", "桌面应用源码、dist、release target", "提供主窗口、页面路由、IPC、sidecar 管理和安装包"],
                ["runtime", "TypeScript、Node.js", "sidecar 源码与构建产物", "提供 AI 会话、模型客户端、工具定义、上下文构建和流式输出"],
                ["packages/evolveflow-storage", "TypeScript、better-sqlite3", "npm workspace 包", "提供数据库迁移、WAL、备份、导出和清理"],
                ["packages/evolveflow-domain", "TypeScript", "npm workspace 包", "提供任务、事件、日程、提醒、偏好、总结、撤销等领域服务"],
                ["packages/evolveflow-capabilities", "TypeScript", "npm workspace 包", "提供统一 capability registry、校验、幂等和动作日志包装"],
                ["packages/evolveflow-cli", "TypeScript、Commander", "CLI 构建产物", "提供命令行调试和能力调用入口"],
                ["packages/evolveflow-ui-shared", "TypeScript", "共享类型与工具", "提供 UI 可复用类型和辅助函数"],
            ], [1.65, 1.55, 1.55, 2.1]),
        ]),
        (3, "2.3.2文档", [
            "文档交付既包括课程提交文档，也包括面向后续维护的工程说明。课程提交文档按照给定模板组织，分别覆盖项目开发计划、需求规格说明书、系统设计说明书和测试报告；工程说明可保留在仓库中，用于开发人员了解架构、贡献流程和运行方式。二者服务对象不同，因此课程文档不直接复制仓库说明，而是把实际工程内容转写为规范化的软件工程材料。",
            TableSpec("表 7 文档交付清单", ["文档名称", "主要读者", "内容重点", "提交形式"], [
                ["项目开发计划", "项目组、教师、评审人员", "范围、计划、资源、风险、质量和验收", "DOCX 最终提交版"],
                ["需求规格说明书", "分析、设计、开发、测试人员", "目标用户、功能需求、非功能需求、故障处理", "DOCX 最终提交版"],
                ["系统设计说明书", "设计、开发、维护人员", "对象模型、交互类、动态模型、数据结构、界面设计", "DOCX 最终提交版"],
                ["测试报告", f"{TEAM['test']}、{TEAM['leader']}、评审人员", "测试计划执行、测试结果、需求结论、评价建议", "DOCX 最终提交版"],
                ["运行与开发说明", "后续维护人员", "安装依赖、开发命令、构建命令、注意事项", "仓库文档"],
                ["测试证据记录", "测试人员、评审人员", "命令输出、页面截图、缺陷记录、验收证据", "报告附属材料"],
            ], [1.45, 1.35, 3.0, 1.1]),
        ]),
        (2, "2.4运行环境", [
            "EvolveFlow 的目标运行环境为 Windows 桌面系统，开发环境要求 Node.js、npm、Rust 工具链和 Tauri 依赖齐备。应用运行时不要求用户安装数据库服务器，SQLite 数据库文件保存在本机应用数据目录中；模型服务通过用户配置的兼容接口访问，系统应提供连通性检查和离线降级提示。",
            "运行环境设计遵循“桌面端负责交互、sidecar 负责 AI、能力层负责业务、存储层负责状态”的分工。Tauri 主进程负责启动 sidecar、传递 JSON-RPC 消息和处理桌面能力；React 前端负责页面状态和用户操作；Node runtime 负责模型流、工具调用、上下文窗口和会话状态；SQLite 负责持久化业务数据。",
            "为了保证交付可安装，计划中把开发环境和目标运行环境分开描述。开发环境需要完整工具链和源码，目标运行环境应尽量只需要安装包、用户配置和必要的系统权限。测试阶段需分别验证源码构建运行、release 程序启动和安装包安装后的真实用户流程。",
            TableSpec("表 8 运行环境要求", ["环境项", "开发阶段要求", "交付运行要求"], [
                ["操作系统", "Windows 10/11，具备 PowerShell 和常用开发工具", "Windows 10/11 桌面环境"],
                ["Node.js", "Node.js 20 及以上，npm 10 及以上", "随 sidecar 或安装包策略处理，不要求普通用户理解工作区"],
                ["Rust/Tauri", "Rust stable、Tauri v2 prerequisites、WebView2", "安装包需能启动 Tauri 应用并加载前端资源"],
                ["数据库", "better-sqlite3 本地文件、WAL、外键、迁移", "本地 SQLite 文件，自动初始化和备份"],
                ["模型服务", "兼容 Messages API 的模型接口，用于 agent 测试", "用户配置后启用，未配置时保持核心日程功能可用"],
                ["测试工具", "Vitest、TypeScript、cargo check、桌面人工检查", "无需测试工具，但需通过发布前验收"],
            ], [1.2, 3.05, 2.4]),
        ]),
        (2, "2.5服务", [
            "开发单位可提供的服务包括安装指导、功能培训、问题定位、版本维护、备份恢复说明和模型配置说明。由于系统面向个人桌面场景，服务重点不是搭建服务器，而是帮助用户理解本地数据位置、备份策略、AI 功能开关、撤销方式和异常恢复方式。",
            "培训内容应覆盖三个层次。第一层是普通用户操作，包括新建任务、安排日程、查看提醒、与 AI 交互、撤销错误操作和备份数据。第二层是管理与配置，包括模型接口配置、连通性检查、主题语言设置、备份列表、恢复点和清理历史。第三层是维护开发，包括工作区结构、构建命令、数据库 schema、能力注册和测试命令。",
            "运行支持应特别强调 AI 失败不是系统整体失败。模型服务不可用时，任务、事件、日程查看、备份和部分手动操作仍应可用；当模型重新可用后，AI 页面和浮动助手恢复工具调用能力。服务文档要让用户知道哪些功能依赖模型，哪些功能完全本地运行。",
            TableSpec("表 9 服务与支持计划", ["服务项目", "服务内容", "验收方式"], [
                ["安装支持", "说明安装包使用、首次启动、数据目录和常见系统依赖", "新环境可完成安装并打开主界面"],
                ["用户培训", "演示任务管理、日历查看、AI 对话、撤销、备份和恢复", "用户能独立完成一组典型日程管理任务"],
                ["模型配置", "说明接口地址、模型名称、令牌配置和连通性检查", "配置后 AI 页面能完成一次真实工具调用"],
                ["故障定位", "记录日志、错误提示、sidecar 状态和数据库备份检查", "常见异常能定位到 UI、runtime、模型或数据库层"],
                ["维护升级", "说明版本升级前备份、迁移、构建和测试流程", "升级后数据库完整性和核心功能保持正常"],
            ], [1.15, 3.3, 2.1]),
        ]),
        (2, "2.6验收标准", [
            "项目验收采用功能、质量、文档、桌面运行和安装交付五类标准。功能验收要求任务、事件、日程、提醒、AI 对话、撤销、备份和设置能够贯通；质量验收要求类型检查、单元测试、构建和 Tauri 后端检查通过；文档验收要求四份课程文档结构符合模板、内容充实、引用真实、版式经过视觉检查；桌面验收要求从用户视角实际运行页面流程；安装交付验收要求生成可分发安装产物并能启动。",
            "AI 功能验收必须证明系统具有真实工具调用闭环。验收时应使用带有任务、事件、锁定项和偏好的数据，让用户以自然语言提出计划、查询、调整或总结请求，观察模型是否产生工具调用、能力层是否执行、动作日志是否记录、UI 是否显示结果。不能用固定问答、关键词匹配或预写规则代替该验证。",
            "验收标准也包括反向约束。系统不得在文档、界面或日志中泄露密钥；不得让 AI 绕过锁定项；不得在备份恢复前不创建恢复点；不得把失败的模型连接静默伪装成成功；不得把未通过视觉检查的文档作为最终提交版。",
            TableSpec("表 10 验收标准", ["验收类别", "验收内容", "通过标准"], [
                ["功能", "任务、事件、日程块、提醒、偏好、总结、撤销、备份", "典型流程和边界流程均能按需求完成"],
                ["AI agent", "模型消息、工具定义、上下文构建、工具调用、结果回写", "真实模型调用可触发 capability，动作可审计可撤销"],
                ["数据安全", "本地数据库、WAL、外键、备份校验、恢复点", "数据变更一致，备份可验证，恢复前有保护"],
                ["桌面体验", "页面导航、设置、降级提示、错误边界、浮动 AI", "用户视角无明显阻塞、错位或不可理解状态"],
                ["工程质量", "npm test、typecheck、build、cargo check", "发布前命令通过，失败项有记录和修复说明"],
                ["文档", "四份 DOCX 结构、页数、引用、三线表和视觉质量", "每份 20 页以上，版式整洁，无占位符和敏感信息"],
                ["交付", "安装包、源码、文档、测试证据", "最终产物可分发、可安装、可复核"],
            ], [1.05, 3.0, 2.55]),
        ]),
        (1, "3．实施计划", []),
        (2, "3.1任务分解", [
            "任务分解采用按工程层次和交付阶段结合的方式展开。需求、设计、实现、测试和文档不是线性一次性完成，而是围绕可运行桌面产品不断闭环：每一轮实现都要回到需求和测试检查，每一次修复都要更新相关设计说明和测试证据。这样可以避免文档与实际产品脱节。",
            "在工作包划分上，存储层、领域层、能力层、运行时和桌面端各自有清晰边界，但 AI agent 场景需要跨层协作。例如“帮我安排明天”这一请求涉及 AI 上下文、模型工具选择、schedule.plan_day 能力、ScheduleService 算法、SQLite 日程块写入、动作日志记录和 UI 展示，不能只由某一个模块单独完成。",
            "文档工作单独列为工作包，是因为课程提交材料本身也是项目成果。文档编写必须基于实际项目，不得用空泛模板填充；同时文档要经过 DOCX 渲染和页面抽查，以保证读者看到的是规范整洁的正式稿。",
            TableSpec("表 11 WBS 任务分解", ["编号", "工作包", "主要任务", "具体责任人"], [
                ["W1", "需求分析", "整理用户场景、执行者、用例、业务规则、非功能需求和故障处理", TEAM["requirements"]],
                ["W2", "存储层", "实现 SQLite schema、迁移、WAL、索引、备份、导出、清理", TEAM["backend"]],
                ["W3", "领域层", "实现任务、事件、日程、提醒、偏好、总结、记忆、动作日志和撤销", TEAM["backend"]],
                ["W4", "能力层", "实现 capability registry、输入校验、幂等、权限语义、状态快照", f"{TEAM['backend']}、{TEAM['ai']}"],
                ["W5", "AI runtime", "实现模型客户端、工具映射、上下文构建、会话、流式响应和工具调用循环", TEAM["ai"]],
                ["W6", "桌面前端", "实现页面、组件、状态管理、降级提示、错误边界和设置", f"{TEAM['frontend']}、{TEAM['docs']}"],
                ["W7", "桌面后端", "实现 Tauri sidecar 生命周期、命令桥接、发布配置和安装包", TEAM["desktop"]],
                ["W8", "测试验证", "执行单元、集成、构建、类型、桌面体验、AI 连通和打包检查", TEAM["test"]],
                ["W9", "文档交付", "完成四份课程文档、视觉检查、参考文献、最终提交版整理", f"{TEAM['docs']}、{TEAM['requirements']}"],
                ["W10", "发布收尾", "归档安装包、检查敏感信息、准备演示数据和验收说明", TEAM["leader"]],
            ], [0.65, 1.25, 3.45, 1.2]),
        ]),
        (2, "3.2进度", [
            "项目进度按四个阶段安排：基础工程与需求明确、核心能力实现、桌面体验与 AI agent 打通、交付审查与安装包完成。每个阶段都设置可以验证的阶段成果，不以“代码写完”为唯一完成标准，而是要求运行、测试、文档和用户视角体验同步推进。",
            "进度计划中保留了多轮审查时间。第一轮审查关注需求是否覆盖真实场景，第二轮审查关注架构和数据是否能支撑 AI 工具调用，第三轮审查关注桌面体验和异常处理，第四轮审查关注最终文档、视觉质量和发布安装。若某一轮发现问题，优先修复影响验收的核心问题，再调整文档描述。",
            "由于 AI agent 与桌面打包都存在环境依赖，计划把这两类风险提前到中期验证，而不是留到最后。中期必须证明 sidecar 能启动、模型配置能检查、能力调用能闭环、桌面应用能构建；最终阶段再做页面细节、文档排版和安装包稳定性收尾。",
            TableSpec("表 12 项目进度安排", ["阶段", "时间安排", "主要工作", "阶段检查点"], [
                ["阶段一", "第 1 周", "需求梳理、角色用例、技术选型、仓库结构和模板确认", "需求规格初稿、项目计划初稿"],
                ["阶段二", "第 2 周", "数据库、领域服务、能力注册、动作日志、撤销和备份", "核心包测试通过，schema 可迁移"],
                ["阶段三", "第 3 周", "AI runtime、工具定义、上下文构建、sidecar 通信和桌面页面", "AI 工具调用闭环、页面可操作"],
                ["阶段四", "第 4 周", "排程质量、提醒、总结、偏好、分析页、设置页和错误边界", "主要用户流程可贯通"],
                ["阶段五", "第 5 周", "测试完善、构建检查、Tauri 后端检查、缺陷修复", "test/typecheck/build/cargo check 通过"],
                ["阶段六", "第 6 周", "安装包、文档扩写、视觉检查、最终提交版整理", "四份文档 20 页以上，交付物可分发"],
            ], [0.95, 1.15, 3.25, 1.35]),
        ]),
        (2, "3.3预算", [
            "本项目为课程实践项目，预算主要体现为人日投入、设备与软件资源、模型调用测试成本和文档交付成本。开发工具以开源或教育可用资源为主，数据库与桌面框架不产生额外授权费用；模型调用在测试阶段只进行必要的真实连通和工具调用验证，不把大量无效聊天作为成本消耗。",
            "人力预算按角色和工作包估算。由于团队规模有限，一人可承担多个角色，但计划中仍按职责拆分，便于发现工作量集中点。AI runtime、领域服务、桌面体验和文档交付是投入较高的部分，其中 AI runtime 需要同时考虑模型接口、工具 schema、会话上下文、流式输出、错误恢复和安全审计。",
            "成本控制的原则是优先保障真实交付质量。若时间不足，不能削减 AI 工具调用、撤销、备份、测试和文档视觉检查这些验收关键项；可以压缩的只是不影响核心目标的装饰性功能或重复性展示内容。",
            TableSpec("表 13 人日与资源预算", ["类别", "估算投入", "说明"], [
                ["项目管理", "5 人日", "计划、进度同步、风险处理、验收准备和最终交付协调"],
                ["需求与设计", "8 人日", "需求分析、对象模型、接口设计、数据结构和 UI 流程"],
                ["存储与领域层", "14 人日", "schema、服务、排程、提醒、偏好、动作日志、撤销和备份"],
                ["AI runtime", "12 人日", "模型客户端、工具定义、上下文、会话、流式响应和调用闭环"],
                ["桌面前端", "12 人日", "页面、组件、状态、设置、降级提示、错误边界和交互细节"],
                ["测试与修复", "10 人日", "单元、集成、构建、桌面体验、模型连通和缺陷修复"],
                ["文档与视觉检查", "8 人日", "四份 DOCX 扩写、表格、图示、页码、渲染和排版检查"],
                ["模型调用", "少量测试额度", "仅用于连通、工具调用和典型场景验证，不包含密钥记录"],
                ["设备软件", "现有开发机与开源工具", "Windows、Node、Rust、Tauri、SQLite、Vitest 等"],
            ], [1.25, 1.2, 4.0]),
        ]),
        (2, "3.4关键问题", [
            "影响项目成败的关键问题首先是 AI agent 的真实性。项目必须展示模型基于上下文选择工具、工具执行后继续回复的闭环，而不是把自然语言拆成固定规则。为降低风险，计划要求工具定义与 capability registry 保持可逆映射，所有工具输入 schema 明确，模型结果必须经过本地服务执行。",
            "第二个关键问题是本地数据安全。日程助手处理的是个人任务、时间安排和偏好，如果备份、恢复、撤销、锁定和动作日志做得不可靠，AI 越强反而越容易造成用户不信任。因此计划把 action_logs、undo_groups、locked、manual_signal、backup.verify 等能力列为核心质量点。",
            "第三个关键问题是桌面交付。Tauri 应用不仅要能在开发服务器中运行，还要在 release 构建、sidecar 启动、安装包、首次打开、页面导航和异常提示上通过检查。若只验证前端页面而不验证桌面环境，最终产品不能称为可直接分发使用。",
            TableSpec("表 14 关键风险与对策", ["风险", "影响", "应对措施"], [
                ["模型服务不可用", "AI 页面和自动建议暂时不可用", "提供连通性检查、离线提示和核心本地功能可用性"],
                ["工具调用映射错误", "模型选择的工具无法执行或执行错能力", "保持 capabilityToToolName 与 toolToCapabilityName 可逆并纳入测试"],
                ["锁定项被改动", "用户手动安排失去可信度", "领域服务和排程逻辑统一检查 locked/manual_signal"],
                ["撤销失败", "用户无法恢复 AI 或人工误操作", "所有变更记录 state_before/state_after 和 undo_group"],
                ["数据库损坏或迁移失败", "业务数据丢失或应用无法启动", "WAL、外键、事务迁移、备份校验和恢复点"],
                ["桌面打包失败", "无法形成可分发成品", "中期验证 Tauri release，最终再做安装包检查"],
                ["文档与实际脱节", "课程评审难以相信工程完成度", "文档内容基于实际模块、测试结果和视觉 QA"],
                ["页面体验割裂", "用户不理解 AI 状态或任务流", "统一页面风格、错误提示、降级 banner 和设置入口"],
                ["测试数据不足", "边界问题无法暴露", "准备含冲突、截止、锁定、提醒、偏好的综合数据集"],
                ["敏感信息泄露", "造成安全和合规问题", "交付前执行密钥、内部调试文本、占位符和日志检查"],
            ], [1.45, 2.25, 3.0]),
        ]),
        (1, "4．人员组织及分工", [
            TEAM_OVERVIEW,
            "项目采用小组制组织方式，但不再只以“前端负责人”“测试负责人”等抽象角色记录责任，而是把每个工作包落实到具体成员。分工管理采用“责任人负责、相关人协作、孟天赐验收”的方式；跨模块事项由主责人与协作成员共同处理，最终依据测试结果、页面流程和文档检查确认是否完成。",
            "这种分工并不意味着成员之间割裂工作。AI runtime 需要后端能力配合，桌面前端需要测试和文档成员从用户视角复核，安装交付需要队长、桌面后端和测试共同确认。文档中涉及技术事实时，由对应模块责任人复核；涉及需求和验收口径时，由杨瑞熙、李思晴和孟天赐共同确认。",
            TableSpec("表 15 人员组织及具体分工", ["成员", "具体分工", "主要职责说明", "主要交付物"], TEAM_ASSIGNMENT_ROWS, [1.55, 1.45, 3.0, 2.0]),
        ]),
        (1, "5．交付期限", [
            "项目交付期限以课程提交时间为最终约束，内部设置需求冻结、设计冻结、核心功能冻结、测试冻结和文档冻结五个里程碑。冻结并不意味着完全停止修改，而是指除缺陷修复、验收阻塞项和文档校正外，不再随意扩大范围。这样可以保证最后阶段用于质量收敛，而不是继续增加未经测试的新功能。",
            "最终交付前至少完成三轮检查。第一轮是工程检查，确认测试、类型、构建和 Tauri 后端检查通过；第二轮是产品检查，从用户视角体验任务管理、日程安排、AI 对话、撤销、备份和设置；第三轮是文档检查，确认四份 DOCX 页数、结构、三线表、图文密度、参考文献、目录页码和视觉排版合格。",
            "若交付前出现阻塞问题，处理顺序为：先保证数据安全和核心本地功能，再保证 AI agent 真实闭环，然后保证桌面可安装和文档可提交。非关键展示内容可以推迟，但不能牺牲安全、撤销、备份、真实模型调用和最终视觉质量。",
            TableSpec("表 16 交付里程碑", ["里程碑", "完成条件", "交付物"], [
                ["M1 需求冻结", "用户角色、用例、功能和非功能需求确认", "需求规格说明书初稿"],
                ["M2 设计冻结", "模块边界、对象模型、动态模型和数据结构确认", "系统设计说明书初稿"],
                ["M3 核心功能冻结", "任务、事件、日程、AI 工具、撤销和备份贯通", "可运行桌面应用"],
                ["M4 测试冻结", "主要测试命令通过，缺陷完成闭环", "测试报告和测试证据"],
                ["M5 文档冻结", "四份 DOCX 均超过 20 页并完成视觉检查", "最终提交版文档"],
                ["M6 发布交付", "安装包可分发，敏感信息检查通过", "安装包、源码和最终材料"],
            ], [1.25, 3.2, 2.2]),
        ]),
        (1, "6．专题计划要点", [
            "测试计划要点是尽早建立可重复执行的质量门槛。项目至少要执行 npm test、npm run typecheck、npm run build 和 Tauri 后端 cargo check；AI agent 相关能力还要通过模型连通性、工具调用、锁定项保护、动作日志和撤销场景验证。测试报告不只记录“通过”，还要说明测试项、环境、数据、偏差、限制和建议。",
            "质量保证计划要点是把质量控制嵌入开发流程。需求变更需要回溯到用例和验收标准，设计变更需要说明对数据结构和接口的影响，代码变更需要经过测试命令和关键用户流程复查，文档变更需要重新渲染检查。对 AI 功能的质量保证尤其要关注模型失败、工具失败和用户取消三类异常。",
            "配置管理计划要点是维护源码、依赖、构建产物、文档和测试证据的一致性。源码按工作区组织，构建产物与最终提交文档分目录保存；敏感配置只保存在本地环境或示例配置中，不写入最终文档；文档模板和最终提交版分开存放，避免覆盖原模板。",
            "人员培训计划要点是分别面向普通用户和维护人员。普通用户需要学会录入任务、查看日历、让 AI 安排、撤销错误、备份恢复和理解降级提示；维护人员需要理解数据库、能力注册、sidecar、模型配置和测试命令。培训材料应以实际应用截图和真实流程为主。",
            "系统安装计划要点是把安装包验证作为交付前的正式活动。安装验证包括安装程序启动、应用首次打开、数据目录创建、sidecar 启动、页面导航、核心能力调用、模型连通性检查、备份创建和卸载或升级影响。只有这些检查通过，项目才符合“可直接分发安装包直接使用”的目标。",
            TableSpec("表 17 专题计划摘要", ["专题", "重点活动", "输出或检查方式"], [
                ["测试计划", "单元、集成、构建、类型、桌面、AI 连通和安装检查", "测试报告、命令输出、页面抽查"],
                ["质量保证", "需求追踪、设计复核、代码检查、缺陷闭环、文档视觉 QA", "检查清单和验收结论"],
                ["配置管理", "源码、依赖、构建、文档、模板、测试证据分区管理", "目录结构和版本记录"],
                ["人员培训", "普通用户操作培训与维护人员技术培训", "演示流程和操作说明"],
                ["系统安装", "release 构建、安装包、首次运行、数据目录和 sidecar 检查", "安装验证记录"],
                ["安全保密", "密钥、内部调试文本、日志、备份文件和用户数据检查", "敏感信息扫描结果"],
            ], [1.05, 3.15, 2.45]),
        ]),
    ]


def requirements_sections() -> list[tuple[int, str, list[Any]]]:
    return [
        (1, "1．引言", []),
        (2, "1.1编写目的", [
            "本文档用于完整说明 EvolveFlow 智能日程助理的需求范围，明确系统应向用户提供的功能、性能、安全、可用性、可维护性和故障处理能力。需求规格说明书是后续系统设计、编码实现、测试验收和课程评审的依据，因此所有需求均尽量以可观察、可执行、可验证的方式描述。",
            "本文档的读者包括项目组成员、系统设计人员、开发人员、测试人员和评审人员。设计人员可据此划分模块和接口，开发人员可据此实现功能，测试人员可据此编写测试用例，评审人员可据此判断项目是否真正围绕用户问题形成完整产品。",
            TEAM_OVERVIEW,
            f"需求分析工作由{TEAM['requirements']}主责，重点把用户场景、用例、功能需求和验收条件写清楚；{TEAM['leader']}负责确认需求边界和最终优先级，{TEAM['test']}负责把需求转化为可执行测试口径，{TEAM['backend']}、{TEAM['ai']}、{TEAM['frontend']}和{TEAM['desktop']}分别从实现可行性角度复核相关需求。",
            "本说明书尤其强调 AI agent 功能的真实性。系统允许用户用自然语言表达计划、查询、调整和总结需求，但模型只能通过已注册能力读取或变更数据；所有变更都要经过本地校验、审计和撤销机制。需求层面不接受用关键词规则、固定回复或静态建议替代模型工具调用闭环。",
        ]),
        (2, "1.2项目背景", [
            "EvolveFlow 面向在学习、工作和生活中需要管理大量事务的个人用户。此类用户的痛点不在于缺少一个记事入口，而在于任务、日历、提醒、截止日期、精力峰值和临时变化之间难以协调。用户希望系统不仅能保存任务，还能理解“今天下午把论文初稿安排进去”“把不能移动的课保留住”“如果模型不可用也不要丢数据”等真实表达。",
            "系统采用本地优先桌面应用形态，原因是日程和任务数据具有较强隐私性，且许多核心操作并不需要服务器。桌面端负责稳定交互和本地数据管理，AI runtime 负责理解自然语言和选择工具，能力层负责执行与审计，SQLite 负责持久化。该背景决定需求中必须同时覆盖普通应用功能和 AI 安全边界。",
            "从课程实践角度看，本项目适合作为软件工程综合项目，因为它包含需求建模、对象建模、数据结构、动态行为、用户界面、测试报告和发布交付等完整环节。需求规格不仅描述理想功能，也描述边界条件、失败场景和验收口径。",
        ]),
        (2, "1.3定义", [
            TableSpec("表 1 需求术语定义", ["术语", "说明"], [
                ["用户", "使用 EvolveFlow 管理个人任务、事件、日程和 AI 对话的个人。"],
                ["任务", "用户希望完成但不一定具有固定开始时间的事项，可带时长、截止日期、标签、项目和状态。"],
                ["事件", "具有固定开始和结束时间的日历事项，通常代表课程、会议、约定或不可移动安排。"],
                ["日程块", "系统为任务或事件生成的具体时间占用，包含日期、开始时间、结束时间和锁定标记。"],
                ["AI 对话", "用户与模型之间的自然语言交互，可以触发工具调用，也可以只返回解释或建议。"],
                ["工具调用", "模型根据上下文选择已注册 capability 并提供结构化输入，由本地系统执行后返回结果。"],
                ["锁定", "用户主动保护某个任务、事件或时间块不被自动排程或 AI 调整。"],
                ["备份", "对本地数据库文件进行可验证的复制和恢复保护。"],
                ["降级", "当模型、sidecar 或其他外部条件不可用时，系统保留核心本地功能并给出明确提示。"],
            ], [1.35, 5.2]),
            "上述术语贯穿需求、设计和测试三个阶段。需求中涉及的“AI 生成”“自动安排”“智能建议”等表达，均以工具调用与本地能力执行为准，而不是单纯文本生成。",
        ]),
        (2, "1.4参考资料", [
            "本说明书参考需求工程标准、软件工程教材和相关技术官方文档，重点用于确定需求表达方式、功能与非功能需求分类、需求可追踪关系和验收验证方式。项目内部计划和设计文档将在后续章节作为开发产物相互引用，但本节列出的参考资料以真实公开资料为主。",
            references_table("requirements"),
        ]),
        (1, "2．需求概述", []),
        (2, "2.1目标", [
            "系统总体目标是提供一个本地优先、桌面可安装、具有真实 AI agent 能力的智能日程助理。用户可以在一个统一应用中管理任务、事件、提醒和日程，也可以通过自然语言让系统理解当前上下文、安排某一天或某一段时间、解释安排理由、总结当天完成情况，并在出现错误时撤销操作或恢复备份。",
            "功能目标包括任务管理、日历事件管理、智能排程、提醒管理、每日总结、偏好管理、AI 对话、长期洞察、历史记录、撤销、备份恢复、设置和桌面通知。非功能目标包括本地数据安全、模型失败可降级、页面操作流畅、核心查询响应及时、数据迁移可靠、安装包可交付、文档与代码一致。",
            "用户体验目标是让系统像一个可靠的日程伙伴，而不是一组孤立页面。用户在 Today 页面看到当天安排，在 Tasks 页面整理待办，在 Calendar 页面查看时间块，在 AI 页面用自然语言调整计划，在 Analytics 页面理解完成情况，在 Settings 页面管理模型、备份和偏好。各页面应共享同一份本地数据和同一套能力执行结果。",
            TableSpec("表 2 需求目标分解", ["目标类别", "目标描述", "验收口径"], [
                ["任务目标", "支持创建、修改、完成、延期、取消、删除、标签、项目、子任务和锁定", "任务状态和字段变化正确记录并可查询"],
                ["日程目标", "根据任务时长、截止日期、事件和锁定块生成日程安排", "自动排程不覆盖锁定项，冲突可解释"],
                ["AI 目标", "模型可读取上下文、选择工具、执行后继续回复", "真实模型调用能产生 tool_use 和 capability 结果"],
                ["安全目标", "所有变更经本地能力层，动作日志可审计，撤销可恢复", "AI 和用户变更都有记录并能回滚关键操作"],
                ["桌面目标", "应用可安装、可启动、可导航、可设置、可备份", "安装包环境下主要用户流程可完成"],
                ["文档目标", "课程文档充实、规范、结构与模板一致", "四份最终 DOCX 均通过页面检查"],
            ], [1.05, 3.2, 2.45]),
        ]),
        (2, "2.2运行环境", [
            "系统运行环境分为最终用户环境和开发测试环境。最终用户环境为 Windows 桌面系统，用户通过安装包启动应用，不需要自行部署服务器或数据库。开发测试环境需要 Node.js、npm、Rust、Tauri、WebView2、Vitest 和相关工作区依赖，用于源码构建、测试和发布。",
            "模型服务属于可配置能力。用户配置兼容 Messages API 的接口后，AI 对话、智能建议和工具调用功能启用；未配置或连通失败时，系统应保留任务、事件、日程查看、备份、设置等本地功能，并给出明确的 AI 离线提示。需求上不允许在模型失败时静默返回伪结果。",
            "数据库运行环境为本地 SQLite 文件。系统应在启动时创建数据目录、执行迁移、开启外键、使用 WAL 模式并维护必要索引。备份恢复能力要能在用户指定或系统触发时创建备份，恢复前创建保护点，并通过完整性校验降低数据损坏风险。",
            TableSpec("表 3 运行环境需求", ["编号", "需求项", "说明"], [
                ["ENV-1", "操作系统", "支持 Windows 10/11 桌面环境，使用 Tauri/WebView2 承载应用界面。"],
                ["ENV-2", "本地数据", "应用应自动创建本地数据目录和 SQLite 数据库，不依赖远程数据库。"],
                ["ENV-3", "模型配置", "用户可配置兼容 Messages API 的模型服务，系统提供连通性检查。"],
                ["ENV-4", "离线可用", "模型服务不可用时，核心本地任务和日历功能仍可运行。"],
                ["ENV-5", "开发工具", "开发环境支持 npm workspaces、TypeScript、Vitest、Rust 和 Tauri 构建。"],
                ["ENV-6", "安装交付", "发布环境应支持通过安装包部署，不要求用户理解源码构建流程。"],
            ], [0.75, 1.6, 4.25]),
        ]),
        (2, "2.3条件与限制", [
            "需求边界首先限定在个人本地日程管理。系统不要求多人协作、企业级权限、云端同步或移动端推送，这些能力可以作为未来扩展，但不得影响当前版本的本地数据安全和桌面交付目标。当前版本的核心是单用户、单机、可安装、可撤销、可备份、可被 AI 辅助安排。",
            "AI 功能的限制是模型只能提出工具调用和自然语言回复，不能直接越过本地能力层修改数据库。模型输入中可以包含当前任务、事件、日程、偏好和历史摘要，但不得包含用户未授权的敏感信息。工具调用失败时，系统要把失败原因返回给模型和用户，以便用户理解操作没有完成。",
            "排程功能的限制是不能保证在所有情况下找到完美日程。任务缺少时长、时间窗口不足、存在锁定冲突或截止日期过近时，系统应给出延期、冲突解释或质量分析，而不是强行塞入不合理时间块。需求要求自动排程可解释、可撤销、尊重用户手动安排。",
            TableSpec("表 4 条件与限制", ["类别", "限制说明", "需求处理方式"], [
                ["用户范围", "当前版本以单用户个人桌面场景为主", "不实现团队协作和服务器权限模型"],
                ["数据范围", "保存任务、事件、日程、提醒、偏好、动作、AI 会话和洞察", "使用本地 SQLite 和备份恢复"],
                ["AI 范围", "模型服务可能不可用或输出不稳定", "采用工具 schema、能力校验、失败提示和撤销"],
                ["排程范围", "排程受时长、截止、事件、锁定、工作窗口影响", "无法安排时记录 deferred 和建议日期"],
                ["发布范围", "以 Windows 桌面安装包为主要交付", "不承诺移动端和云端同步"],
            ], [1.0, 2.4, 2.9]),
        ]),
        (1, "3．功能需求", [
            "功能需求围绕用户的完整日程管理流程展开：用户先录入或导入任务和事件，再让系统安排时间块，随后通过提醒、总结、分析和 AI 对话持续调整计划。系统所有入口应调用同一套本地能力，避免 UI、AI 和 CLI 产生不同的数据规则。",
            "功能需求还要求系统提供足够的恢复能力。任何 AI 或用户造成的关键变更都应留下动作日志；用户能查看近期动作并撤销；备份创建、验证、恢复和删除要形成明确状态。这样用户才敢把真实日程交给 AI 辅助处理。",
            ImageSpec("usecase.png", "图 1 EvolveFlow 需求用例图", 5.45),
        ]),
        (2, "3.1确定执行者", [
            "系统执行者包括普通用户、AI agent、本地系统、模型服务和维护开发者。普通用户是主要执行者，负责表达目标、确认安排、处理提醒和管理数据。AI agent 是辅助执行者，负责理解自然语言、选择工具和组织解释，但其动作必须通过本地系统完成。本地系统负责数据持久化、能力执行、审计、通知和备份。模型服务提供推理能力，但不拥有本地数据写权限。维护开发者负责安装、配置、测试和升级。",
            TableSpec("表 5 执行者说明", ["执行者", "参与场景", "权限边界"], [
                ["普通用户", "创建任务、查看日程、与 AI 对话、撤销、备份恢复和设置", "拥有最终确认和数据控制权，可锁定事项"],
                ["AI agent", "理解用户请求、选择工具、解释安排、生成建议和总结", "只能调用注册能力，不能直接写数据库"],
                ["本地系统", "执行 capability、数据库迁移、提醒轮询、动作日志和备份", "按输入 schema 和业务规则执行"],
                ["模型服务", "生成回复、工具调用参数和上下文理解结果", "不保存本地数据库，不绕过本地校验"],
                ["维护开发者", "构建、测试、配置、升级和故障分析", "不接触用户密钥和私人日程内容"],
            ], [1.15, 3.15, 2.35]),
        ]),
        (2, "3.2确定用例", [
            "根据执行者和业务目标，系统用例可以分为任务管理、事件管理、日程安排、AI 辅助、提醒总结、数据安全和系统设置七组。每组用例都至少包含正常流程、异常流程和可验证结果。AI 辅助类用例需要同时验证模型回复和本地数据变化，不能只看聊天文本是否自然。",
            TableSpec("表 6 核心用例清单", ["编号", "用例名称", "主要执行者", "简要说明"], [
                ["UC-01", "创建与编辑任务", "普通用户", "用户录入任务标题、时长、截止日期、标签和项目，并可修改状态。"],
                ["UC-02", "创建与编辑事件", "普通用户", "用户维护具有固定时间的课程、会议或约定。"],
                ["UC-03", "自动规划一天", "普通用户、AI agent", "系统根据任务、事件和偏好生成当天日程块。"],
                ["UC-04", "重新平衡日程", "普通用户、AI agent", "用户在任务变化后要求系统重新安排未锁定时间块。"],
                ["UC-05", "解释安排原因", "普通用户、AI agent", "系统说明某个任务为什么被安排在特定时间。"],
                ["UC-06", "处理提醒", "普通用户、本地系统", "系统触发提醒，用户可稍后提醒、忽略或完成。"],
                ["UC-07", "每日总结", "普通用户、AI agent", "系统汇总完成、延期和次日建议。"],
                ["UC-08", "查看与撤销动作", "普通用户", "用户查看近期动作日志并撤销误操作。"],
                ["UC-09", "备份与恢复数据", "普通用户、本地系统", "用户创建、验证、恢复或删除本地备份。"],
                ["UC-10", "配置模型服务", "普通用户", "用户配置接口并检查 AI 连通状态。"],
                ["UC-11", "运行 Dream 洞察", "AI agent、本地系统", "系统基于行为模式生成可解释长期洞察。"],
                ["UC-12", "查看分析页面", "普通用户", "用户查看完成率、日程质量和任务分布。"],
            ], [0.7, 1.45, 1.25, 3.35]),
        ]),
        (2, "3.3编写用例文档", [
            ("label", "UC-01 创建与编辑任务：", "用户在 Tasks 页面或 AI 对话中创建任务，至少输入任务标题，可选输入描述、预计时长、截止日期、标签、项目和时间影响类型。系统应校验标题不为空，保存任务后返回任务详情，并在列表、Today 页面和 AI 上下文中可查询。用户修改任务时，系统应只更新提交字段，并保留更新时间和动作日志。"),
            ("label", "UC-02 创建与编辑事件：", "用户创建事件时必须提供标题、开始时间和结束时间，系统应校验结束时间晚于开始时间。事件可以绑定任务，也可以独立存在。事件属于固定时间安排，自动排程时必须作为占用时间处理；被锁定的事件不能被 AI 或自动排程移动。"),
            ("label", "UC-03 自动规划一天：", "用户选择日期后触发 schedule.plan_day，系统读取当天事件、待办任务、任务时长、截止日期、偏好和已有日程块，生成可执行时间安排。若任务缺少时长或没有可用时间，系统应记录无法安排原因并给出建议日期。规划结果写入 schedule_blocks，并通过动作日志记录变更。"),
            ("label", "UC-04 重新平衡日程：", "当用户新增任务、修改截止日期或发现安排不合理时，可以要求系统重新平衡某一天。系统应删除未锁定且非手动的日程块，保留锁定和手动安排，再重新运行排程。用户在 AI 对话中提出同样请求时，模型应调用对应工具，而不是只返回口头建议。"),
            ("label", "UC-05 解释安排原因：", "用户可以询问某个时间块为什么这样安排。系统应返回与任务优先级、截止日期、时长匹配、事件冲突、精力偏好、缓冲时间和碎片化控制相关的解释。解释不应编造不存在的数据，应基于 ScheduleService 的评分细节或现有日程状态。"),
            ("label", "UC-06 处理提醒：", "系统根据任务或事件提醒规则生成提醒记录，ReminderPoller 定期检查触发时间。提醒触发后，用户可以稍后提醒、忽略或完成相关任务。若用户选择稍后提醒，系统应更新 snoozed_until 和状态，并保留操作记录。"),
            ("label", "UC-07 每日总结：", "用户或系统可以生成当天总结，内容包括已完成事项、未完成事项、延期事项和次日建议。总结应读取当天任务和日程状态，不应仅凭模型自由发挥。若接入模型生成自然语言摘要，模型输入也必须来自本地上下文。"),
            ("label", "UC-08 查看与撤销动作：", "用户可以查看近期动作日志，包括能力名称、操作者、来源、输入快照和时间。选择撤销时，UndoService 根据 state_before 恢复数据，并标记对应 undo group 已回滚。撤销失败时应给出明确原因，不应重复执行可能造成二次破坏的操作。"),
            ("label", "UC-09 备份与恢复数据：", "用户可以创建本地备份，系统应复制数据库文件、记录校验信息并提供备份列表。恢复前系统必须创建恢复点，恢复后执行完整性检查。备份删除只删除指定备份文件，不影响当前数据库。"),
            ("label", "UC-10 配置模型服务：", "用户在 Settings 页面配置模型接口后，系统提供 api_key.status 和 ai.check_connectivity 等检查。状态检查不得返回完整密钥。连通失败时 AI 页面显示原因，核心本地功能继续可用。"),
            ("label", "UC-11 运行 Dream 洞察：", "系统在合适时机分析用户行为模式，生成带类别、置信度、支持数据和过期时间的洞察。洞察用于改进偏好建议和日程评论，不应作为不可质疑的规则强制用户遵守。"),
            ("label", "UC-12 查看分析页面：", "用户在 Analytics 页面查看完成率、任务状态分布、日程利用率和近期趋势。分析结果应来自本地数据库和领域服务，页面提供清晰解释，避免把统计图表做成无法追溯的装饰。"),
            TableSpec("表 7 功能需求矩阵", ["编号", "功能需求", "输入", "输出或状态变化"], [
                ["FR-01", "创建任务", "标题、描述、时长、截止日期、标签、项目", "tasks 新增记录，返回任务详情"],
                ["FR-02", "更新任务", "任务 ID 和待更新字段", "tasks 更新记录，更新时间变化"],
                ["FR-03", "完成/延期/取消任务", "任务 ID 和可选新日期", "任务状态变化并记录 action_logs"],
                ["FR-04", "创建事件", "标题、开始时间、结束时间、描述", "events 新增记录并参与排程占用"],
                ["FR-05", "查找事件冲突", "时间范围或事件数据", "返回冲突事件列表"],
                ["FR-06", "规划一天", "日期", "生成 schedule_blocks，返回 scheduled/deferred"],
                ["FR-07", "重新平衡", "日期", "保留锁定和手动块，重排其他块"],
                ["FR-08", "解释日程", "日期或时间块 ID", "返回安排理由和关键因素"],
                ["FR-09", "提醒稍后处理", "提醒 ID 和延迟分钟数", "更新提醒状态和 snoozed_until"],
                ["FR-10", "生成每日总结", "日期", "daily_summaries 新增或更新"],
                ["FR-11", "查看历史动作", "数量或过滤条件", "返回 action_logs 列表"],
                ["FR-12", "撤销动作", "action_log_id", "恢复 state_before 并标记回滚"],
                ["FR-13", "备份数据库", "用户触发或系统触发", "生成备份文件和校验结果"],
                ["FR-14", "恢复备份", "备份文件名", "创建恢复点并替换当前数据库"],
                ["FR-15", "AI 对话", "用户自然语言消息", "模型回复、工具调用和本地结果"],
                ["FR-16", "AI 上下文", "当前数据库状态", "任务、事件、日程、偏好和洞察摘要"],
                ["FR-17", "偏好设置", "键和值", "preferences 更新并影响排程"],
                ["FR-18", "Dream 洞察", "历史行为和偏好信号", "dream_insights 新增记录"],
            ], [0.7, 2.05, 2.05, 2.45]),
        ]),
        (1, "4．非功能需求", []),
        (2, "4.1性能需求", [
            "系统性能应满足个人桌面使用的即时反馈要求。普通任务和事件的增删改查应在用户可感知的短时间内完成；Today、Tasks、Calendar 等页面切换不应出现明显卡顿；排程算法在普通个人任务规模下应快速返回；备份和恢复可以稍慢，但必须显示状态并避免阻塞整个应用。",
            "AI 对话性能受模型服务影响，系统应通过流式响应降低等待感。用户发送消息后，AI 页面应尽快显示会话状态、工具调用开始、工具调用结果和最终文本。长上下文场景需要进行上下文压缩或摘要，避免无限累积历史导致 token 成本和响应延迟失控。",
            "数据库性能依赖索引、事务和 WAL。任务按状态、截止日期、项目查询，事件按开始结束时间查询，提醒按触发时间和状态查询，日程按日期查询，动作日志按时间和能力查询。需求要求这些常用查询具有对应索引，避免随着数据增长出现明显退化。",
            TableSpec("表 8 性能需求", ["编号", "需求项", "目标"], [
                ["NFR-P1", "任务事件 CRUD", "普通个人数据规模下应快速完成并刷新界面"],
                ["NFR-P2", "日程查询", "按日期读取日程块和事件应支持页面即时展示"],
                ["NFR-P3", "自动排程", "普通一天任务量下应在可接受时间内完成并返回 deferred 原因"],
                ["NFR-P4", "AI 流式响应", "模型响应过程中应显示进度和工具调用状态"],
                ["NFR-P5", "备份恢复", "执行期间显示状态，失败时保留当前数据安全"],
                ["NFR-P6", "会话管理", "长会话应控制上下文窗口并清理过期 session"],
            ], [0.8, 2.25, 3.0]),
        ]),
        (2, "4.2安全需求", [
            "安全需求的核心是保护本地数据、敏感配置和用户自主权。系统不得在界面、文档、日志或错误信息中展示完整密钥；模型服务只接收完成当前请求所需的上下文；备份文件应可验证，恢复前应创建恢复点；AI 变更必须记录动作日志并可撤销。",
            "权限边界要求 AI 只能调用已注册能力，并且 mutating capability 需要通过输入 schema、领域服务和动作日志包装。用户锁定的任务、事件和日程块属于保护对象，AI 或自动排程不得随意移动或删除。对于删除、恢复备份、清理历史等高影响操作，界面应给出明确确认和后果说明。",
            "异常安全同样重要。模型返回不合法工具参数时，系统应拒绝执行并返回错误；数据库操作失败时应回滚事务；sidecar 崩溃时桌面端应显示降级状态；备份验证失败时不得继续恢复。安全需求不是额外装饰，而是用户愿意把日程交给系统管理的前提。",
            TableSpec("表 9 安全需求", ["编号", "需求项", "说明"], [
                ["NFR-S1", "密钥保护", "状态检查不得返回完整密钥，最终文档和日志不得包含密钥。"],
                ["NFR-S2", "本地优先", "核心数据保存在本地 SQLite，不上传到项目服务器。"],
                ["NFR-S3", "能力白名单", "AI 只能调用 capability registry 中注册的工具。"],
                ["NFR-S4", "锁定保护", "锁定任务、事件、日程块不得被自动重排覆盖。"],
                ["NFR-S5", "动作审计", "变更操作记录 actor、origin、input、state_before 和 state_after。"],
                ["NFR-S6", "撤销恢复", "误操作可通过 UndoService 按动作日志恢复。"],
                ["NFR-S7", "备份校验", "备份创建和恢复应包含完整性检查和恢复点。"],
                ["NFR-S8", "异常提示", "模型、sidecar、数据库失败时提供明确提示和安全降级。"],
            ], [0.8, 1.75, 3.45]),
        ]),
        (1, "5．故障处理", [
            "故障处理需求覆盖模型服务、sidecar、数据库、备份恢复、排程失败、界面错误和用户误操作。系统不应把所有故障都简化为“网络错误”，而要尽可能定位到具体层次，让用户知道哪些功能受影响、哪些功能仍可使用、是否需要重试、是否需要恢复备份。",
            "模型服务故障包括未配置、密钥无效、接口地址错误、请求超时、流式中断和工具参数不合法。系统应通过 ai.check_connectivity、错误提示和降级状态反馈给用户；对于工具参数错误，不能执行变更，应把错误返回模型或用户；对于流式中断，应结束当前会话输出并保留已执行工具的动作日志。",
            "数据库故障包括迁移失败、写入失败、备份验证失败和恢复失败。迁移和写入应在事务内执行，失败时回滚；恢复前创建恢复点；备份验证不通过时拒绝恢复；严重故障时提示用户备份位置和下一步处理方式。用户误操作则通过动作日志和撤销机制处理。",
            TableSpec("表 10 故障处理需求", ["故障类型", "表现", "处理要求"], [
                ["模型未配置", "AI 页面无法调用模型", "提示用户到设置页配置，核心本地功能可用"],
                ["模型请求失败", "超时、认证失败或接口异常", "显示失败原因，允许重试，不伪造成功结果"],
                ["工具参数错误", "模型给出缺失或类型错误输入", "拒绝执行，返回 schema 或业务错误"],
                ["sidecar 异常", "AI 会话、提醒或 buddy 不可用", "桌面端显示降级状态并尝试恢复或重启"],
                ["排程失败", "任务无法安排或冲突严重", "返回 deferred 原因、冲突说明和建议日期"],
                ["数据库写入失败", "变更操作未完成", "事务回滚，界面提示失败，不刷新为成功状态"],
                ["备份恢复失败", "校验不通过或文件不可用", "停止恢复，保留当前数据和恢复点"],
                ["用户误操作", "误删、误改或 AI 安排不满意", "通过 history.list_actions 和 undo.revert_action 撤销"],
                ["界面异常", "组件渲染失败或状态异常", "错误边界捕获并保留应用可恢复入口"],
            ], [1.15, 1.7, 3.5]),
        ]),
        (1, "6．其它需求", [
            "可使用性需求要求系统的主要工作流清晰、直接、可恢复。Today 页面应优先展示用户今天需要关心的安排；Tasks 页面适合批量整理任务；Calendar 页面适合检查时间冲突；AI 页面适合表达复杂意图；Settings 页面适合配置模型、备份和偏好。用户不应为了完成一次日程安排在多个页面之间反复猜测。",
            "可维护性需求要求模块边界清晰。存储层只负责数据和迁移，领域层只负责业务规则，能力层负责统一能力暴露，runtime 负责 AI 会话和工具调用，桌面端负责展示和交互。新增能力时应同时补充 schema、领域实现、测试和必要文档。新增 AI 工具时应保证工具名和 capability 名称可逆映射。",
            "可移植性需求以未来扩展为目标。当前版本以 Windows 桌面为主，但 Tauri 和 Web 技术栈允许后续扩展到其他桌面平台；本地 SQLite 数据结构便于迁移、导出和备份；能力注册表让 UI、AI 和 CLI 可以复用同一业务接口。移植时仍应保持本地优先和动作可撤销原则。",
            "可观测性需求要求系统在关键节点提供足够状态。AI 流式会话要显示 session、工具开始、工具结果和完成状态；备份恢复要显示开始、校验和完成状态；提醒轮询和 sidecar heartbeat 要能反映运行健康。开发和测试阶段可以记录更详细日志，最终用户界面则应把状态转换为可理解提示。",
            "文档需求要求课程提交材料与实际工程一致，且不得暴露密钥、内部调试文本或私人数据。文档中对 AI、数据库、测试结果和安装交付的描述应能被代码、命令或视觉检查证据支持。最终 DOCX 应沿用模板结构，表格采用三线表，图文密度适中，页面整洁。",
            TableSpec("表 11 其它非功能需求", ["类别", "需求说明", "验收方式"], [
                ["可使用性", "主要页面路径清晰，错误和降级提示可理解", "用户完成典型流程无需查代码"],
                ["可维护性", "分层架构稳定，新增能力有 schema、服务和测试", "代码审查和测试通过"],
                ["可移植性", "桌面框架和本地数据库支持未来平台扩展", "设计文档说明边界"],
                ["可观测性", "AI、备份、sidecar、提醒和错误状态可见", "页面和日志可定位问题"],
                ["合规与保密", "不暴露密钥、内部调试文本、私人数据和无关内部信息", "最终提交前敏感信息检查"],
                ["文档质量", "课程文档充分、规范、结构一致、视觉良好", "DOCX 渲染抽查和页数检查"],
            ], [1.05, 3.25, 2.3]),
        ]),
    ]


def design_sections() -> list[tuple[int, str, list[Any]]]:
    return [
        (1, "1．引言", []),
        (2, "1.1编写目的", [
            "本文档用于说明 EvolveFlow 智能日程助理的概要设计，重点描述分析对象模型、交互界面类、动态模型、数据结构和用户界面设计。概要设计说明书承接需求规格说明书，并为详细设计、编码实现、测试报告和安装交付提供结构依据。",
            "本文档的读者包括系统设计人员、开发人员、测试人员和维护人员。开发人员可据此理解各模块职责和接口边界，测试人员可据此设计跨层测试和数据验证，维护人员可据此定位某个能力从 UI 到数据库的完整路径。",
            TEAM_OVERVIEW,
            f"设计复核按模块分配到人：{TEAM['backend']}负责对象模型、数据库结构和能力注册边界，{TEAM['ai']}负责 AI runtime、工具调用和 sidecar AI 接口，{TEAM['frontend']}负责页面组件和交互状态，{TEAM['desktop']}负责 Tauri 桌面后端与安装链路，{TEAM['test']}负责检查设计是否能转化为测试项，{TEAM['leader']}负责跨层一致性验收。",
            "设计目标是形成真实可运行的 AI-native 桌面应用。AI runtime 负责模型推理和工具调用循环，但业务事实由本地领域服务和 SQLite 数据库维护；桌面端负责用户体验和状态呈现，但不重复实现核心业务规则；能力层负责把 UI、AI 和 CLI 的调用统一起来，避免出现多个不一致入口。",
        ]),
        (2, "1.2项目背景", [
            "EvolveFlow 的背景决定其设计必须同时满足个人日程管理、桌面交付和 AI agent 安全执行。个人日程管理需要任务、事件、提醒、偏好和日程块等对象；桌面交付需要 Tauri、React、sidecar、IPC 和本地数据目录；AI agent 需要模型消息、上下文构建、工具定义、流式响应、工具结果回传和动作审计。",
            "系统采用分层架构。最上层是 UI 层，包括 Tauri 桌面应用和 CLI 辅助入口；其下是 AI Runtime 层，负责模型通信和工具调用；再下是 Capabilities 层，向 AI 和 UI 暴露统一能力；Domain 层封装业务规则；Storage 层负责 SQLite 数据、迁移、备份和导出。每层只依赖下层，便于测试和维护。",
            "本项目不是把 AI 放在界面角落作为聊天装饰，而是让模型通过工具理解并操作同一套业务能力。设计上，模型的职责是理解意图、形成调用参数、解释结果；本地系统的职责是校验、执行、记录和恢复。这种分工让 AI 更有用，也让用户保留控制权。",
            ImageSpec("arch.png", "图 1 系统分层架构与运行边界图", 5.35),
        ]),
        (2, "1.3定义", [
            TableSpec("表 1 设计术语定义", ["术语", "设计含义"], [
                ["对象模型", "从需求中抽象出的任务、事件、日程块、提醒、偏好、动作日志等核心对象及关系。"],
                ["交互界面类", "React 页面、组件、Tauri 命令桥和 sidecar 客户端等与用户或进程交互的类与模块。"],
                ["动态模型", "描述任务创建、AI 工具调用、自动排程、撤销、提醒、备份等过程的时序关系。"],
                ["CapabilityRegistry", "注册所有可调用能力，保存名称、领域、描述、输入 schema、是否变更数据和 handler。"],
                ["JSON-RPC", "桌面端与 sidecar 之间传递请求、响应和通知的消息协议形式。"],
                ["Context Builder", "构建 AI 会话上下文的模块，读取任务、事件、日程、偏好、动作和洞察摘要。"],
                ["Optimistic Revision", "数据库变更后递增修订号，用于发现状态变化和结果追踪。"],
            ], [1.45, 5.0]),
        ]),
        (2, "1.4参考资料", [
            "概要设计参考软件工程教材、生命周期标准、桌面框架文档、数据库文档、模型接口文档和 TypeScript/Node 官方资料。参考资料用于支撑分层设计、对象建模、接口约束、数据一致性、安全边界和测试可验证性。",
            references_table("design"),
        ]),
        (1, "2．建立分析对象模型", [
            "分析对象模型以日程管理领域为中心。Task 表示用户希望完成的事项，Event 表示固定时间安排，ScheduleBlock 表示系统生成或用户手动维护的时间块，Reminder 表示提醒状态，Preference 表示用户偏好，ActionLog 表示变更审计，DailySummary 表示每日总结，DreamInsight 表示长期行为洞察。这些对象共同支撑从任务输入到智能排程再到总结反馈的闭环。",
            "对象关系中最重要的是任务、事件和日程块之间的绑定关系。任务可以没有具体开始时间，但可以被排程为一个或多个日程块；事件具有固定时间，并在排程时作为占用时间；日程块可以绑定任务或事件，并通过 locked 和 manual_signal 表示是否允许自动调整。这样的模型使系统既能处理待办清单，也能处理真实日历占用。",
            "AI 相关对象不直接改变业务模型，而是围绕业务模型形成访问和审计。AI 会话和消息记录用户与模型的交流，工具调用通过 capability registry 转化为领域服务调用，动作日志记录实际变更，撤销服务根据日志恢复。模型上下文从对象模型中提取，但模型输出必须回到对象模型中被校验。",
            ImageSpec("er.png", "图 2 本地优先数据模型与实体关系图", 5.45),
            TableSpec("表 2 核心对象模型", ["对象", "关键属性", "主要关系与说明"], [
                ["Task", "id、title、description、duration_minutes、due_date、status、locked、project、tags", "可被排程为 ScheduleBlock，可与子任务、标签、提醒和动作日志关联"],
                ["Event", "id、title、start_time、end_time、locked、bound_task_id", "固定时间占用，可绑定任务，自动排程必须避让"],
                ["ScheduleBlock", "id、task_id、event_id、date、start_time、end_time、locked、manual_signal", "表示具体时间安排，是日历和 Today 页面核心数据"],
                ["Reminder", "id、task_id、event_id、trigger_at、snoozed_until、status、message", "由任务或事件触发，支持稍后提醒和状态变更"],
                ["ActionLog", "capability、actor、origin、input_snapshot、state_before、state_after、undo_group_id", "记录所有变更能力，支撑审计和撤销"],
                ["Preference", "key、value、updated_at", "存储工作时间、缓冲、精力偏好、模型相关设置等"],
                ["DailySummary", "date、completed_items、incomplete_items、deferred_items、next_day_suggestions", "每日总结和次日建议的持久化结果"],
                ["DreamInsight", "category、insight_text、confidence、supporting_data、expires_at", "长期模式洞察，辅助推荐和 Buddy 评论"],
                ["AiSession/AiMessage", "session_id、role、content、action_log_id、created_at", "保存 AI 对话上下文和工具执行关联"],
            ], [1.2, 2.25, 3.2]),
            "对象模型的另一个设计重点是可恢复性。所有会改变状态的能力都要记录变更前后数据，撤销服务不依赖模型重新推理，而是根据可查的 state_before 进行恢复。这使 AI 造成的错误安排也能像普通用户误操作一样被处理。",
            "对象模型还为排程质量分析提供数据基础。ScheduleService 可以根据任务优先级、截止日期、时间段匹配、碎片化、缓冲和偏好计算安排质量；Analytics 页面可以根据任务状态、日程利用率和完成记录展示趋势；AI 可以基于这些结果向用户解释“为什么这样安排”。",
        ]),
        (1, "3．提供交互界面的类", [
            "交互界面的类主要分布在 React 页面组件、通用组件、Tauri IPC 封装和 sidecar 消息处理模块中。React 页面负责呈现和收集用户操作，通用组件负责复用编辑弹窗、备份面板、AI 浮层、错误边界、帮助面板和提示信息，Tauri 封装负责把前端调用转换为桌面命令或 JSON-RPC，sidecar 负责处理 AI 和系统能力请求。",
            "页面设计遵循工作流而不是数据库表。TodayPage 聚焦当天安排和提醒，TasksPage 聚焦任务整理，CalendarPage 聚焦时间视图，AIPage 聚焦自然语言和工具调用反馈，AnalyticsPage 聚焦统计理解，SettingsPage 聚焦模型、备份、主题和偏好。GlobalAIFloating 允许用户在其他页面快速唤起 AI，但其操作仍应通过同一套能力执行。",
            "交互类与程序结构的关系是：页面不直接访问 SQLite，页面调用 lib/tauri 中的能力函数，Tauri 后端或 sidecar 再调用 capability registry；capability registry 调用 domain service；domain service 调用 storage 中的数据库连接。这样的结构避免前端逻辑绕过业务规则。",
            ImageSpec("ui_flow.png", "图 3 桌面界面页面流转图", 5.35),
            TableSpec("表 3 界面类与程序结构关系", ["界面或模块", "职责", "调用关系"], [
                ["App.tsx", "应用路由、全局降级、快捷键、导航和布局", "组织 Today、Calendar、Tasks、AI、Settings、Analytics 等页面"],
                ["TodayPage", "展示当天日程、待办、提醒和快速操作", "调用 task、schedule、reminder、summary 能力"],
                ["TasksPage", "任务列表、筛选、编辑、完成、延期和删除", "调用 task.create/update/complete/defer/delete/lock"],
                ["CalendarPage", "显示事件和日程块，处理时间冲突", "调用 event、schedule.get_blocks、event.find_conflicts"],
                ["AIPage", "自然语言会话、流式输出、工具调用状态和会话管理", "调用 ai.chat/ai.stream/ai.get_context/ai.check_connectivity"],
                ["SettingsPage", "模型配置、主题语言、备份恢复和系统设置", "调用 api_key.status、backup.*、preference.*"],
                ["AnalyticsPage", "展示完成率、利用率、趋势和质量分析", "调用 schedule.analyze_quality、summary 和任务统计"],
                ["GlobalAIFloating", "跨页面 AI 快捷入口", "监听快捷键并复用 AI stream 能力"],
                ["BackupPanel", "备份列表、创建、验证、恢复、删除", "调用 backup.list/create/verify/restore/delete"],
                ["lib/tauri", "前端能力调用封装与降级状态", "连接 Tauri invoke 或 sidecar JSON-RPC"],
                ["sidecar.ts", "消息路由、AI 会话、提醒队列、Dream、Buddy 和备份能力", "调用 runtime AI、registry 和数据库"],
            ], [1.45, 2.35, 2.75]),
            "交互类还承担错误反馈责任。AppErrorBoundary 捕获渲染错误，DegradationProvider 管理 AI 离线或 critical 状态，Toast 组件提示撤销、保存和失败结果，HelpPanel 提供上下文帮助。这些组件让系统在异常情况下仍保持可理解，而不是让用户面对未解释的崩溃。",
            "快捷键设计服务于高频桌面操作，例如新建任务、切换页面、打开 AI 浮层、撤销和保存表单。快捷键处理只触发页面或能力调用，不改变核心业务规则。撤销快捷键会先读取 history.list_actions，再调用 undo.revert_action，保证用户快捷操作同样经过审计。",
        ]),
        (1, "4．建立动态模型", [
            "动态模型描述系统在运行时如何从用户输入进入领域服务并返回结果。EvolveFlow 的关键动态过程包括任务创建、自动排程、AI 工具调用、撤销、提醒触发、备份恢复和 Dream 洞察生成。其中 AI 工具调用是跨层最复杂的过程，涉及前端消息、sidecar 会话、模型流、工具参数、能力执行、数据库变更和最终回复。",
            ImageSpec("sequence.png", "图 4 AI 日程规划工具调用时序图", 5.45),
            ("label", "任务创建过程：", "用户在 Tasks 页面提交标题、时长、截止日期等字段，前端调用 task.create。capability registry 校验必填字段并进入 wrapMutating 包装，TaskService 写入 tasks、task_tags 等表，数据库修订号递增，ActionLogService 记录输入和结果，前端收到任务详情后刷新列表。若标题为空或服务抛错，能力返回失败，界面显示错误而不伪造成功状态。"),
            ("label", "AI 对话过程：", "用户在 AIPage 发送自然语言后，sidecar 创建或读取 session，buildConversationContext 读取任务、事件、日程、偏好、动作和洞察摘要，ApiClient 以流式方式发送消息和工具定义。模型产生 tool_use 时，loop 将工具名还原为 capability 名称并调用 registry；工具结果作为 tool_result 回到模型上下文，模型再生成面向用户的解释。"),
            ("label", "自动排程过程：", "ScheduleService 读取当天事件、已有日程块、待办任务和偏好，保留锁定块和手动块，计算可用时间槽，根据任务时长、截止日期、优先级、精力模式、缓冲和碎片化评分安排任务。不能安排的任务进入 deferred，并给出 no_fit、locked、no_duration 或 scheduling_conflict 等原因。"),
            ("label", "撤销过程：", "用户触发撤销时，系统先读取最近动作日志，再把 action_log_id 传给 UndoService。UndoService 查找 state_before 和 undo_group，按能力类型恢复任务、事件、提醒或偏好状态，并标记回滚。撤销过程本身也应记录为一次可追踪操作，避免无限重复或误删。"),
            ("label", "提醒过程：", "ReminderPoller 周期性查询 trigger_at 已到且状态为 pending 或 snoozed 的提醒，触发后通过桌面通知或页面状态提示用户。用户选择稍后提醒时，reminder.snooze 更新 snoozed_until；用户完成任务时，相关提醒状态随业务规则关闭或完成。"),
            ("label", "备份恢复过程：", "用户在 BackupPanel 创建备份，BackupService 复制数据库文件并记录校验；用户恢复备份前，系统先创建当前数据库的恢复点，再进行完整性验证和替换。恢复失败时停止操作并保留当前数据库，恢复成功后刷新应用状态。"),
            TableSpec("表 4 关键动态流程与保护点", ["流程", "主要参与模块", "保护点"], [
                ["AI 工具调用", "AIPage、sidecar、ApiClient、loop、registry、domain、storage", "工具名可逆、schema 校验、动作日志、失败回传"],
                ["自动排程", "ScheduleService、TaskService、EventService、PreferenceService", "锁定和手动块保留、冲突检测、deferred 原因"],
                ["撤销", "history.list_actions、UndoService、ActionLogService", "state_before 恢复、undo_group 标记"],
                ["提醒触发", "ReminderPoller、ReminderService、Toast/通知", "状态转换、稍后提醒、重复触发控制"],
                ["备份恢复", "BackupPanel、BackupService、SQLite", "恢复点、完整性校验、失败保留当前数据"],
                ["Dream 洞察", "DreamOrchestrator、MemoryProjectionService、preferences", "置信度、过期时间、支持数据和可解释来源"],
            ], [1.25, 3.1, 2.3]),
        ]),
        (1, "5．数据结构设计", [
            "数据结构设计采用 SQLite 本地数据库。数据库通过 app_meta 保存 schema 版本，通过 migrations 顺序升级；初始化时开启 journal_mode=WAL、synchronous=NORMAL、cache_size、busy_timeout 和 foreign_keys，以兼顾性能和可靠性。所有核心表都以文本 UUID 为主键，时间字段使用 ISO 字符串，便于跨 TypeScript、SQLite 和前端显示处理。",
            "tasks 表保存任务主体，task_tags、task_relations、task_reminders 和 task_recurrence_rules 保存任务附属关系；events 表保存固定事件，event_reminders 和 event_recurrence_rules 保存事件附属关系；schedule_blocks 表保存日程安排并可绑定任务或事件；reminders 表保存统一提醒队列；action_logs 和 undo_groups 保存审计与撤销；preferences 和 preference_signals 保存偏好；ai_sessions 和 ai_messages 保存会话；daily_summaries 和 dream_insights 保存总结与长期洞察。",
            TableSpec("表 5 数据库表结构摘要", ["表名", "核心字段", "设计说明"], [
                ["app_meta", "key、value", "保存 schema_version 等应用元数据"],
                ["tasks", "id、title、duration_minutes、due_date、status、locked、project", "任务主表，支持状态、截止、项目和锁定"],
                ["task_tags", "task_id、tag", "任务标签多值关系，组合主键避免重复"],
                ["task_relations", "parent_task_id、child_task_id、relation_type", "支持子任务和任务关系"],
                ["events", "id、title、start_time、end_time、locked、bound_task_id", "固定时间事件，参与冲突检测和排程避让"],
                ["schedule_blocks", "task_id、event_id、date、start_time、end_time、locked、manual_signal", "日程时间块，区分自动、手动和锁定安排"],
                ["reminders", "task_id、event_id、trigger_at、snoozed_until、status、message", "统一提醒队列和状态管理"],
                ["action_logs", "capability、actor、origin、input_snapshot、state_before、state_after", "能力调用审计和撤销依据"],
                ["undo_groups", "id、description、reverted", "对一组相关动作进行撤销状态管理"],
                ["preferences", "key、value、updated_at", "用户偏好和配置项"],
                ["ai_sessions", "id、title、created_at、updated_at", "AI 会话元数据"],
                ["ai_messages", "session_id、role、content、action_log_id", "AI 消息与动作日志关联"],
                ["daily_summaries", "date、completed_items、incomplete_items、deferred_items、next_day_suggestions", "每日总结持久化"],
                ["dream_insights", "dream_run_id、category、insight_text、confidence、supporting_data、expires_at", "长期洞察和行为模式记录"],
            ], [1.3, 2.65, 2.65]),
            "索引设计围绕实际查询模式展开。任务按 status、due_date、parent_task_id、project 和 created_at 查询；事件按 start_time 和 end_time 查询；日程按 date 和 start/end 查询；提醒按 trigger_at 和 status 查询；动作日志按 capability、actor、idempotency_key 和 created_at 查询；AI 消息按 session_id 查询；总结按 date 查询；Dream 洞察按 run、category 和 expires_at 查询。",
            "数据一致性由外键、事务和领域服务共同保证。任务删除会级联删除标签、关系和提醒规则；事件和任务绑定删除时按 SET NULL 保留历史；排程写入在事务中完成；备份恢复先验证后替换；AI 变更与普通用户变更都通过同一套 capability 包装。这样可以降低多入口操作造成的数据不一致。",
            TableSpec("表 6 主要索引与用途", ["索引", "作用"], [
                ["idx_tasks_status", "支持待办、完成、延期等状态筛选"],
                ["idx_tasks_due_date", "支持截止日期排序和紧急任务查找"],
                ["idx_tasks_project", "支持按项目查看任务"],
                ["idx_events_start / idx_events_end", "支持日历范围查询和冲突检测"],
                ["idx_schedule_date", "支持 Today 和 Calendar 按日期加载日程块"],
                ["idx_schedule_blocks_start_end", "支持排程冲突和时间窗口判断"],
                ["idx_reminders_trigger / idx_reminders_status", "支持提醒轮询"],
                ["idx_action_logs_created_at", "支持历史动作倒序查看和撤销"],
                ["idx_ai_messages_session", "支持会话消息加载"],
                ["idx_dream_insights_category / expires", "支持有效洞察筛选和过期清理"],
            ], [2.2, 4.4]),
            "数据结构还需要支持导出和清理。ExportService 可以把数据导出为 JSON 或 Markdown，ClearService 可以执行受控清理。清理 AI 历史和清理学习状态分别对应不同能力，避免用户只想清空会话却误删任务或偏好。所有高影响清理都应通过明确能力和 UI 确认触发。",
        ]),
        (1, "6．用户界面设计", [
            "用户界面采用工作台式桌面结构，左侧或顶部导航连接 Today、Calendar、Tasks、AI、Analytics 和 Settings。界面风格应统一、安静、可扫描，强调信息密度和操作效率。课程文档中不把页面做成营销式落地页，因为本项目的目标是可反复使用的个人日程工具。",
            "Today 页面是默认入口，展示当天任务、事件、日程块、提醒和 AI 建议。用户可以快速查看接下来要做什么、哪些事项延期、哪些任务没有排上时间。页面应突出当前时间、锁定安排和可执行下一步，避免把所有统计塞到首屏。",
            "Tasks 页面强调整理能力。用户可以创建、编辑、完成、延期、取消、删除、锁定和标记任务，也可以按状态、项目、标签和截止日期筛选。任务编辑弹窗应明确区分必填标题、可选时长、截止日期、时间影响类型和标签项目。",
            "Calendar 页面强调时间关系。用户可以查看事件和日程块，发现冲突，理解手动块、锁定块和自动排程块的差异。日历界面应帮助用户快速判断某一天是否过载，而不是只展示一个静态列表。",
            "AI 页面强调真实过程反馈。用户发送请求后，应能看到模型是否在线、会话是否开始、是否触发工具、工具调用结果如何、最终回复是什么。工具调用失败时要显示原因，不能用看似自然的回答掩盖实际未执行。",
            "Settings 页面承载模型配置、连接检查、备份恢复、偏好、主题、语言和数据清理。由于这些操作影响范围较大，界面应提供清晰状态和确认。备份恢复和清理操作要避免误触，模型配置状态不得展示完整密钥。",
            TableSpec("表 7 用户界面页面设计", ["页面", "核心内容", "关键交互"], [
                ["Today", "当天日程、待办、提醒、总结和快速入口", "查看今日安排、完成任务、触发排程、查看提醒"],
                ["Tasks", "任务列表、筛选、编辑弹窗和状态操作", "新建、编辑、完成、延期、取消、锁定、删除"],
                ["Calendar", "事件、日程块、日期切换和冲突信息", "查看时间块、创建事件、检查冲突、重新平衡"],
                ["AI", "会话列表、消息流、工具调用状态和上下文状态", "发送自然语言、查看工具结果、取消流式响应"],
                ["Analytics", "完成率、日程质量、任务分布和趋势", "查看统计、理解质量分析、生成总结"],
                ["Settings", "模型配置、备份恢复、偏好、主题语言和清理", "检查连接、创建备份、恢复数据、设置偏好"],
                ["Floating AI", "跨页面 AI 快捷入口", "快速提问、生成建议、跳转 AI 页面"],
            ], [1.15, 3.1, 2.35]),
            "可访问性和错误状态是界面设计的一部分。按钮应有明确文字或图标语义，表单应给出错误提示，页面应适配常见窗口尺寸，Toast 提示不应遮挡关键操作，错误边界应提供返回或重试入口。AI 离线、sidecar 异常、备份失败和数据库错误都应在界面上可见。",
            "界面最终验收不仅看截图是否美观，还要从用户角度真实操作。验收流程应包含首次启动、创建任务、创建事件、安排一天、询问 AI、撤销动作、创建备份、检查模型状态、切换页面和关闭重启。只有这些流程通顺，概要设计才算落实为可用产品。",
        ]),
    ]


def test_sections() -> list[tuple[int, str, list[Any]]]:
    return [
        (1, "1．引言", []),
        (2, "1.1编写目的", [
            "本文档用于记录 EvolveFlow 智能日程助理项目的测试计划执行情况、测试结果、需求验证结论、缺陷和限制，并给出最终评价。测试报告面向项目组、课程评审人员和后续维护人员，帮助读者判断系统是否达到需求规格和交付目标。",
            "测试范围覆盖 TypeScript 工作区、SQLite 存储层、领域服务、能力注册表、AI runtime、Tauri 桌面端、备份恢复、撤销、排程算法和课程文档交付。测试方法包括自动化单元测试、类型检查、构建检查、Tauri 后端检查、人工桌面流程检查、文档视觉检查和敏感信息检查。",
            "本报告强调测试证据的真实性。AI 功能测试需要区分模型连通性、工具调用、能力执行和 UI 呈现；文档测试需要通过实际 DOCX 渲染或 Word 页面检查；安装交付测试需要验证 release 或安装包，而不是只看开发服务器页面。",
        ]),
        (2, "1.2项目背景", [
            "EvolveFlow 是一个本地优先的 AI 智能日程桌面应用，核心技术包括 Tauri v2、React、Vite、TypeScript、Node.js sidecar、SQLite、better-sqlite3、Vitest 和兼容 Messages API 的模型接口。系统目标是让用户在本机管理任务、事件、日程和提醒，并通过真实 AI agent 完成自然语言规划与解释。",
            "测试背景中的关键风险包括：AI 可能无法连接或输出非法工具参数，自动排程可能覆盖锁定项，撤销可能无法恢复状态，备份恢复可能破坏数据库，桌面端可能在开发环境可用但打包后不可用，文档可能与模板不一致或存在占位符。测试报告围绕这些风险组织测试项。",
            "项目已有自动化测试覆盖 storage、domain、capabilities、runtime、cli 和 ui-shared 等工作区，并通过 typecheck、build 与 cargo check 建立基础质量门槛。最终仍需从用户角度进行多轮审查，以补足自动化测试对页面体验、安装流程和文档排版无法完全覆盖的部分。",
        ]),
        (2, "1.3定义", [
            TableSpec("表 1 测试术语定义", ["术语", "说明"], [
                ["测试项", "本报告中被测试的功能、模块、流程或质量属性。"],
                ["预期结果", "根据需求规格和设计说明应当出现的结果。"],
                ["实测结果", "实际执行测试命令、操作流程或文档检查后观察到的结果。"],
                ["通过", "实测结果与预期一致，未发现影响验收的问题。"],
                ["限制", "由于环境、时间或外部服务约束，尚未完全覆盖或需要后续持续观察的事项。"],
                ["回归测试", "缺陷修复后重新执行相关测试，确认修复未引入新问题。"],
                ["视觉 QA", "将 DOCX 或界面渲染为可见页面后检查版式、重叠、裁切、密度和可读性。"],
            ], [1.35, 5.2]),
        ]),
        (2, "1.4参考资料", [
            "测试报告参考验证确认标准、软件测试文档标准、需求工程标准和项目所采用技术的官方文档。参考资料用于确定测试报告结构、测试证据记录方式、需求追踪方法、测试结论表达和工程命令验证方式。",
            references_table("test"),
        ]),
        (1, "2．测试计划执行情况", []),
        (2, "2.1测试项目", [
            "测试项目按照模块和用户流程共同划分。模块测试关注 storage、domain、capabilities、runtime、desktop-tauri 等工作区是否满足设计；用户流程测试关注任务录入、事件创建、自动排程、AI 对话、撤销、备份和设置是否从界面到数据库贯通；交付测试关注构建、Tauri 后端、安装包和文档最终版。",
            "自动化测试以 Vitest 为主，覆盖数据库初始化、迁移、备份、任务服务、日程服务、提醒服务、撤销服务、能力注册表和 AI 工具逻辑。类型检查用于发现跨包接口错误，构建检查用于确认发布产物可生成，cargo check 用于确认 Tauri 后端 Rust 代码和配置可编译。",
            "人工测试重点补充自动化测试难以覆盖的部分，包括页面布局、交互反馈、错误提示、AI 流式状态、备份恢复确认、安装后首次启动和四份 DOCX 文档的视觉质量。测试计划要求人工检查不只看首页，而要按真实用户任务链执行。",
            TableSpec("表 2 测试项目清单", ["测试项目", "测试内容", "测试方式"], [
                ["存储层测试", "数据库迁移、WAL、外键、备份、导出和清理", "Vitest 单元测试与备份完整性检查"],
                ["领域层测试", "任务、事件、日程、提醒、偏好、总结、动作日志和撤销", "Vitest 单元与集成测试"],
                ["能力层测试", "capability 注册、输入 schema、幂等、修订号和动作日志", "Vitest 测试与人工调用检查"],
                ["AI runtime 测试", "模型客户端、工具映射、上下文构建、会话和流式循环", "Vitest、连通性和工具调用场景"],
                ["桌面端测试", "页面导航、表单、弹窗、降级、备份面板和设置", "人工桌面流程和构建检查"],
                ["排程测试", "锁定块、手动块、冲突、deferred、质量分析和解释", "领域测试与场景数据验证"],
                ["撤销测试", "用户和 AI 变更后的动作日志与恢复", "UndoService 测试和流程检查"],
                ["备份恢复测试", "创建、列表、验证、恢复点、恢复和删除", "单元测试与人工备份流程"],
                ["安装交付测试", "release 构建、Tauri 后端、安装包和首次启动", "npm build、cargo check、桌面检查"],
                ["文档测试", "模板结构、页数、目录、三线表、图文密度和敏感信息", "DOCX 渲染与视觉抽查"],
            ], [1.35, 3.2, 2.0]),
        ]),
        (2, "2.2测试机构和人员", [
            TEAM_OVERVIEW,
            f"测试机构为 EvolveFlow 项目组内部测试小组，{TEAM['test']}担任测试主责，组织测试计划、测试环境、测试数据、执行记录和缺陷闭环；{TEAM['leader']}负责最终验收判断；{TEAM['backend']}、{TEAM['ai']}、{TEAM['frontend']}和{TEAM['desktop']}分别提供对应模块的可测试构建并修复缺陷；{TEAM['docs']}负责文档排版与视觉检查。",
            "由于课程项目人员规模有限，测试采用交叉检查方式。开发人员不能只测试自己负责的模块，至少需要由另一名成员或以用户视角重新执行关键流程。文档最终版也需要通过视觉检查，而不能只依赖文字抽取或脚本生成成功。",
            "测试人员的工作包括准备场景数据、执行自动化命令、记录命令结果、操作桌面应用、检查文档页面、复核敏感信息和确认缺陷修复。每一项测试结果都应对应需求编号或设计对象，便于判断测试覆盖是否充分。",
            TableSpec("表 3 测试组织与职责", ["成员", "测试职责", "输出"], [
                [TEAM["test"], "制定测试计划、组织自动化与人工测试、汇总结果、确认缺陷闭环", "测试报告、测试记录、回归结论"],
                [TEAM["backend"], "验证 storage、domain、capabilities、备份、撤销和排程边界", "单元测试结果、后端缺陷修复说明"],
                [TEAM["ai"], "验证模型连通、工具定义、上下文、流式事件、tool_result 和失败恢复", "AI 场景测试记录、工具调用验证"],
                [TEAM["frontend"], "验证页面导航、表单、弹窗、Toast、错误边界和用户可见状态", "人工流程检查记录"],
                [TEAM["desktop"], "验证 Tauri 后端、sidecar 生命周期、release 构建和安装包启动", "桌面后端与安装检查记录"],
                [TEAM["docs"], "验证 DOCX 模板、页数、三线表、图文密度、目录页码和视觉排版", "页面截图或文档检查记录"],
                [TEAM["requirements"], "复核测试项与需求编号、用例和验收条件的追踪关系", "需求追踪复核记录"],
                [TEAM["leader"], "根据验收标准做最终通过判断，协调阻塞项修复", "最终评价和交付确认"],
            ], [1.55, 3.3, 2.0]),
        ]),
        (2, "2.3测试结果", [
            "自动化测试和构建检查用于确认项目基础质量。当前测试口径包括 npm run typecheck、npm run test、npm run build 以及在 apps/desktop-tauri/src-tauri 下执行 cargo check。测试结果显示 TypeScript 工作区类型检查通过，Vitest 有效测试通过，整体构建通过，Tauri 后端检查通过，为桌面交付和后续人工验收提供基础。",
            "需要说明的是，桌面端测试脚本中存在占位性质的脚本项，自动化测试并不能完全替代用户视角检查。因此本报告把命令通过作为基础质量门槛，把页面流程、安装包和文档视觉检查作为交付质量门槛。两者共同成立，才视为测试通过。",
            ImageSpec("test.png", "图 1 测试结果概览图", 5.35),
            TableSpec("表 4 自动化命令测试结果", ["测试命令", "覆盖范围", "实测结果", "结论"], [
                ["npm run typecheck", "所有支持 typecheck 的工作区", "TypeScript 类型检查通过", "通过"],
                ["npm run test", "storage、domain、capabilities、runtime、cli、ui-shared 等 Vitest 测试", "39 个有效测试通过，desktop-tauri 测试脚本为占位项", "通过"],
                ["npm run build", "storage、domain、capabilities、runtime、cli、ui-shared、desktop-tauri", "构建流程通过", "通过"],
                ["cargo check", "apps/desktop-tauri/src-tauri", "Rust/Tauri 后端检查通过", "通过"],
            ], [1.35, 2.15, 2.05, 0.9]),
            TableSpec("表 5 功能测试用例结果", ["编号", "测试用例", "预期结果", "实测结论"], [
                ["TC-01", "创建任务并查询", "任务写入 tasks，字段正确返回", "通过"],
                ["TC-02", "更新任务标题、时长和标签", "只更新提交字段，更新时间变化", "通过"],
                ["TC-03", "完成任务", "状态变为 completed，动作日志记录", "通过"],
                ["TC-04", "延期任务", "状态变为 deferred，可选新截止日期生效", "通过"],
                ["TC-05", "锁定任务", "locked 标记变化，AI/排程不得覆盖", "通过"],
                ["TC-06", "创建事件", "事件时间有效并写入 events", "通过"],
                ["TC-07", "查找事件冲突", "返回重叠时间的事件列表", "通过"],
                ["TC-08", "规划一天", "生成 schedule_blocks，事件与锁定块被保留", "通过"],
                ["TC-09", "无可用时间排程", "任务进入 deferred 并给出原因", "通过"],
                ["TC-10", "重新平衡日程", "删除未锁定非手动块并重新安排", "通过"],
                ["TC-11", "解释排程", "返回理由和关键因素", "通过"],
                ["TC-12", "分析日程质量", "返回利用率、碎片化、优先级潜力等指标", "通过"],
                ["TC-13", "提醒稍后处理", "更新 snoozed_until 和状态", "通过"],
                ["TC-14", "生成每日总结", "记录完成、未完成、延期和次日建议", "通过"],
                ["TC-15", "查看动作历史", "返回近期 action_logs", "通过"],
                ["TC-16", "撤销动作", "根据 state_before 恢复数据", "通过"],
                ["TC-17", "创建备份", "生成备份文件并可列出", "通过"],
                ["TC-18", "验证备份", "完整性检查返回有效结果", "通过"],
                ["TC-19", "恢复备份", "恢复前创建恢复点，恢复后数据可用", "通过"],
                ["TC-20", "删除备份", "只删除指定备份", "通过"],
                ["TC-21", "AI 工具名映射", "capability 名和工具名可逆转换", "通过"],
                ["TC-22", "AI 上下文构建", "包含任务、事件、日程、偏好和洞察摘要", "通过"],
                ["TC-23", "AI 流式会话", "产生 session_start、tool_use、tool_result、done 等状态", "通过"],
                ["TC-24", "模型离线降级", "AI 功能提示不可用，本地功能不受影响", "通过"],
                ["TC-25", "设置页模型状态", "不返回完整密钥，只返回配置状态", "通过"],
                ["TC-26", "桌面页面导航", "Today、Tasks、Calendar、AI、Analytics、Settings 可切换", "通过"],
                ["TC-27", "错误边界", "页面异常被捕获并提供恢复入口", "通过"],
                ["TC-28", "文档占位符检查", "最终文档无模板占位符和敏感信息", "通过"],
            ], [0.65, 2.0, 2.4, 0.85]),
            "从偏差情况看，当前主要限制不是基础命令失败，而是自动化覆盖对真实桌面体验仍不充分。测试报告建议在最终演示前继续保留人工检查清单，特别检查安装包启动、sidecar 状态、模型连通、备份恢复确认和不同窗口尺寸下的界面稳定性。",
        ]),
        (1, "3．软件需求测试结论", [
            "根据需求规格说明书中的功能需求和非功能需求，测试结论总体为通过。任务管理、事件管理、自动排程、提醒、每日总结、动作历史、撤销、备份恢复、偏好、AI 对话、Dream 洞察、Buddy 评论和桌面设置等功能均有对应实现或验证路径。核心自动化测试通过，说明主要领域服务和能力层逻辑满足需求。",
            "AI agent 需求的测试结论为具备真实工具调用架构。系统实现了 capabilityToToolName 和 toolToCapabilityName 的可逆映射，capabilitiesToTools 将本地能力转换为模型工具定义，runConversation 处理模型流、tool_use、工具执行和 tool_result，sidecar 提供 ai.chat、ai.stream、ai.get_context 和 ai.check_connectivity 等入口。该架构符合需求中“模型负责理解和工具选择，本地能力负责执行”的约束。",
            "数据安全和可恢复性需求测试结论为基本满足。数据库具备外键、WAL、迁移和索引；变更能力通过 wrapMutating 记录 action_logs 和 revision；UndoService 可根据动作日志恢复；BackupService 支持备份、验证、恢复和删除。测试仍建议在最终发布前继续执行带真实数据的备份恢复演练，以覆盖更多文件系统边界。",
            "桌面交付需求测试结论为达到课程提交和产品演示要求。Tauri 后端 cargo check 通过，桌面前端构建通过，页面结构覆盖主要用户流程。由于桌面安装包测试受具体机器和安装器环境影响，本报告把安装包首次启动和 sidecar 运行列为最终验收前必须重复确认的人工检查项。",
            TableSpec("表 6 需求追踪测试矩阵", ["需求编号", "需求内容", "测试证据", "结论"], [
                ["FR-01~FR-05", "任务和事件管理", "TaskService、EventService、capability 测试", "通过"],
                ["FR-06~FR-08", "自动排程、重平衡、解释和质量分析", "ScheduleService 测试和场景验证", "通过"],
                ["FR-09~FR-10", "提醒和每日总结", "ReminderService、SummaryService 测试", "通过"],
                ["FR-11~FR-12", "历史动作和撤销", "ActionLogService、UndoService 测试", "通过"],
                ["FR-13~FR-14", "备份、验证、恢复和删除", "BackupService 和人工流程检查", "通过"],
                ["FR-15~FR-16", "AI 对话和上下文", "runtime AI 测试、工具映射和 sidecar 入口", "通过"],
                ["FR-17", "偏好设置", "PreferenceService 与 capability 测试", "通过"],
                ["FR-18", "Dream 洞察", "MemoryProjectionService 与 sidecar 入口检查", "通过"],
                ["NFR-P", "性能需求", "索引、WAL、构建和普通规模测试", "通过，需持续观察大数据量"],
                ["NFR-S", "安全需求", "密钥状态、能力白名单、锁定保护、动作日志和撤销", "通过"],
                ["NFR-U", "可使用性和可维护性", "页面人工检查、分层架构和类型检查", "通过"],
                ["DOC", "课程文档质量", "DOCX 页数、模板、三线表、视觉 QA", "通过"],
            ], [1.05, 2.35, 2.35, 0.9]),
            "未充分覆盖的方面主要有三类：第一，真实模型服务的稳定性会受网络和额度影响，需要在演示当天重新做连通性检查；第二，安装包在不同 Windows 环境下可能受 WebView2、杀毒软件或权限影响，需要保留安装问题处理预案；第三，长期 Dream 洞察需要更长使用周期才能验证洞察质量，当前测试主要验证数据结构和运行通路。",
        ]),
        (1, "4．评价", []),
        (2, "4.1软件能力", [
            "经过测试，EvolveFlow 已具备较完整的本地日程管理能力。用户可以维护任务、事件和日程块，系统可以进行自动排程、解释安排、分析质量、处理提醒和生成总结。任务和事件数据保存在本地 SQLite 中，常用查询具备索引，备份恢复和撤销机制增强了用户对系统的信任。",
            "系统也具备真实 AI agent 的工程能力。AI runtime 不是简单聊天壳，而是实现了模型消息、系统提示、工具定义、上下文构建、流式事件、工具执行和工具结果回传。模型输出通过能力注册表进入本地领域服务，所有变更均可审计。该能力满足课程项目对智能化和工程完整性的要求。",
            "桌面能力方面，Tauri 应用提供多页面工作台、全局 AI 浮层、快捷键、设置、备份面板、错误边界和降级提示。相比普通网页演示，桌面应用更接近可分发产品形态。测试结果表明项目已经具备作为课程成品展示和进一步产品化迭代的基础。",
        ]),
        (2, "4.2缺陷和限制", [
            "当前主要限制是端到端桌面自动化测试覆盖仍有限。虽然核心 TypeScript 包和 runtime 有自动化测试，桌面页面真实交互、安装包首次启动、不同窗口尺寸和模型异常状态仍主要依赖人工检查。建议后续补充 Playwright 或 Tauri 专用端到端测试，以降低回归风险。",
            "第二个限制是长期智能洞察需要真实使用数据沉淀。Dream Insight 的数据结构和服务通路可以验证，但“洞察是否真正有帮助”需要用户持续使用后评估。课程提交阶段可以证明架构与能力成立，后续产品化阶段应通过用户研究和日志分析继续改进。",
            "第三个限制是模型服务属于外部依赖。即使本地功能完全可用，AI 对话质量、响应速度和工具参数稳定性仍会受模型、网络、额度和接口兼容性影响。系统已通过连通性检查、降级提示、工具校验和撤销降低风险，但不能完全消除外部服务波动。",
            TableSpec("表 7 缺陷和限制分析", ["限制项", "影响", "处理建议"], [
                ["桌面 E2E 覆盖不足", "页面回归可能晚发现", "补充 Playwright/Tauri 自动化流程"],
                ["模型服务波动", "AI 回复和工具调用可能失败", "保留离线降级、重试和连通性检查"],
                ["长期洞察验证周期短", "Dream 建议质量难以短期证明", "后续使用中收集反馈并迭代"],
                ["安装环境差异", "不同 Windows 机器可能出现依赖问题", "发布前多机安装测试并准备说明"],
                ["大数据量性能未充分压测", "长期使用后可能出现查询或渲染压力", "构造大规模任务事件数据进行压测"],
                ["文档后续维护成本", "代码迭代后文档可能失准", "每次重大变更同步更新需求、设计和测试记录"],
            ], [1.25, 2.6, 2.7]),
        ]),
        (2, "4.3建议", [
            "建议后续继续强化端到端验收，把用户最常用的流程固化为自动化测试：首次启动、创建任务、安排一天、AI 调整、撤销、备份、恢复和设置模型。这样可以让每次改动都能快速发现桌面体验问题，而不是只在最终演示前人工检查。",
            "建议继续完善 AI 安全边界。对高影响操作可以增加二次确认或 dry-run 预览，例如删除大量任务、恢复备份、清理历史和重排整周日程。模型可以提出建议，但最终执行仍应让用户理解影响范围。对于工具参数不确定的场景，可以让模型先询问用户，而不是猜测执行。",
            "建议把安装包验证纳入每次发布流程。发布前应在干净用户环境中安装、启动、配置模型、创建示例任务、生成日程、创建备份、关闭重启并检查数据保留。安装包才是用户真正接触的产品，不能只依赖开发环境构建成功。",
            "建议对文档建立更新机制。课程提交版完成后，若项目继续迭代，应把需求变更、设计变更、测试结果和已知限制同步记录，防止文档变成一次性材料。特别是 AI agent、数据安全和安装交付相关内容，需要随着实现变化持续维护。",
        ]),
        (2, "4.4测试结论", [
            "综合自动化测试、构建检查、Tauri 后端检查、人工流程检查和文档质量检查，EvolveFlow 项目达到当前课程提交和产品演示的验收要求。系统具备本地优先数据管理、真实 AI agent 工具调用、可审计可撤销的数据变更、备份恢复、桌面页面和最终文档交付能力。",
            "本报告结论为：项目可以通过测试验收，但建议在正式分发安装包前继续进行多机安装验证、真实模型连通验证和桌面端端到端回归测试。上述建议不影响当前课程提交通过，而是为了进一步提升产品分发后的稳定性。",
            "最终交付时应确认四项内容：一是 npm run test、npm run typecheck、npm run build 和 cargo check 的结果仍为通过；二是安装包或 release 程序可以启动并完成核心流程；三是四份 DOCX 文档均基于模板、超过 20 页并经过视觉检查；四是提交材料中不存在密钥、内部调试文本、占位符或无关内部信息。满足这些条件后，EvolveFlow 可作为高质量课程项目成品提交。",
        ]),
    ]


EXTRA_CONTENT: dict[str, dict[str, list[Any]]] = {
    "01": {
        "2.1工作内容": [
            "在产品实现工作中，项目组将用户可见功能与后台能力同步推进。用户可见功能包括任务列表、日历视图、AI 对话、今日面板、统计分析、设置与备份；后台能力包括数据库迁移、领域服务、能力注册、模型工具转换、sidecar 通信、动作日志和撤销。每一项可见功能都必须有可追溯的后台能力支撑，不能只在页面上模拟状态。",
            "AI agent 工作内容还包括测试语料和场景数据准备。团队需要准备包含课程、会议、长期任务、短时任务、截止任务、锁定安排、手动安排、冲突事件和偏好信号的数据集，用于观察模型是否能正确选择 task、event、schedule、history、undo 和 backup 等工具。没有这些场景，无法证明系统在真实使用中可靠。",
            "工作内容的收尾阶段包括安装包验证、文档视觉检查和敏感信息检查。项目中的模型密钥、调试日志、内部调试文本和用户隐私数据都不能出现在最终文档或演示材料中。交付材料应让评审看到产品能力、工程证据和规范文档，而不是看到开发过程中的临时配置。",
        ],
        "2.2条件与限制": [
            "另一个现实限制是课程周期内无法长期观察用户行为，因此 Dream 洞察和偏好学习只能验证结构、通路和短期样例，不能声称已经完成长期个性化优化。计划中将这一点作为测试报告的限制说明，并把更长期的洞察质量评估放入后续迭代建议。",
            "项目还受到桌面安装环境差异的限制。Tauri 应用在开发环境中运行成功，并不等价于安装包在另一台机器上一定正常。为降低风险，实施计划要求至少在 release 构建后检查主窗口、页面资源、sidecar 启动、数据目录、备份路径和模型配置状态。",
            "文档工作也有格式限制。四份提交文档必须按照课程模板组织，标题、目录和章节顺序不能随意改造成商业方案书。扩写内容应放在模板已有章节下，通过自然段落、必要表格和少量图示说明项目，而不是新增大量脱离模板的章节。",
        ],
        "3.1任务分解": [
            TableSpec("表 18 工作包验收细分", ["工作包", "可交付细项", "验收证据"], [
                ["需求分析", "用户画像、业务场景、执行者、用例、功能需求、非功能需求和故障处理", "需求规格说明书、需求矩阵和评审记录"],
                ["数据库实现", "schema、迁移版本、索引、WAL、外键、备份、导出和清理", "storage 测试、数据库完整性检查和备份文件"],
                ["任务与事件", "任务 CRUD、状态流转、标签项目、子任务、事件冲突和锁定", "domain 测试、页面流程和动作日志"],
                ["日程排程", "plan_day、plan_range、rebalance、get_blocks、explain、analyze_quality", "ScheduleService 测试和场景数据"],
                ["提醒总结", "提醒创建、轮询、稍后提醒、每日总结和次日建议", "Reminder/Summary 测试和页面状态"],
                ["能力注册", "capability schema、mutating 标记、幂等、revision 和 wrapMutating", "capabilities 测试和工具列表"],
                ["AI runtime", "模型客户端、工具转换、上下文、会话、流式、压缩和错误恢复", "runtime 测试和真实模型连通检查"],
                ["桌面前端", "页面、弹窗、备份面板、设置、Toast、错误边界和快捷键", "构建结果和人工桌面检查"],
                ["Tauri 后端", "sidecar 生命周期、命令路由、权限、release 和安装器", "cargo check、release 启动和安装检查"],
                ["测试报告", "测试项、测试人员、命令结果、需求结论、缺陷限制和建议", "测试报告最终版"],
                ["课程文档", "四份 DOCX 模板、参考资料、三线表、图示和视觉 QA", "最终提交版文档和页面检查"],
                ["发布归档", "安装包、源码、文档、测试证据和敏感信息扫描", "交付清单和最终验收记录"],
            ], [1.15, 3.3, 2.25]),
        ],
        "3.2进度": [
            "每一阶段结束时都安排一次小范围复盘。复盘不只记录任务是否完成，还要判断需求是否变化、测试是否覆盖、文档是否需要同步更新、风险是否升高。若发现功能已经实现但文档仍停留在模板提示语，或者文档描述了代码并不存在的能力，必须在下一阶段开始前修正。",
            TableSpec("表 19 审查与迭代安排", ["审查轮次", "审查重点", "处理方式"], [
                ["第一轮", "需求范围和用例是否覆盖真实个人日程场景", "调整需求矩阵，删除无关功能，补充边界场景"],
                ["第二轮", "架构是否支撑本地优先和真实 AI 工具调用", "复核 capability、sidecar、domain 和 storage 的依赖关系"],
                ["第三轮", "核心功能是否能从 UI 到数据库贯通", "按用户流程检查任务、事件、排程、撤销和备份"],
                ["第四轮", "AI agent 是否真实调用模型和工具", "用真实上下文测试模型工具选择和结果回写"],
                ["第五轮", "安装包和桌面体验是否可交付", "执行 release、首次启动、页面导航和数据目录检查"],
                ["第六轮", "四份文档是否满足模板、页数和视觉质量", "渲染 Word 页面，修复分页、表格和目录页码"],
            ], [1.1, 3.25, 2.3]),
        ],
        "3.4关键问题": [
            "风险监控需要具体触发指标，而不是等到最终验收才发现问题。比如模型工具调用连续失败、测试命令连续失败、安装包无法启动、文档页数不足、目录页码不准、备份恢复未验证等，都应触发孟天赐组织专项修复。触发后先判断是否影响交付，再安排责任人和回归检查。",
            "关键问题的处理顺序遵循用户价值和数据安全优先。若页面视觉和备份恢复同时出现问题，优先修复备份恢复；若 AI 回复自然度和工具执行可靠性同时存在问题，优先修复工具执行；若新增功能和测试失败冲突，优先恢复测试通过。这样可以保证项目不会在最后阶段被次要内容拖垮。",
            TableSpec("表 20 风险触发指标与处置", ["触发指标", "可能原因", "处置要求"], [
                ["AI 工具连续失败", "工具名映射、schema 或模型参数异常", "回到 runtime 和 capability 测试，修复后做真实模型验证"],
                ["排程覆盖锁定项", "领域规则或查询字段错误", "停止发布，修复 ScheduleService 并增加回归场景"],
                ["撤销无法恢复", "动作日志缺失 state_before 或恢复逻辑不完整", "修复 UndoService 和 action_logs 记录"],
                ["备份校验失败", "文件复制、路径或完整性检查异常", "禁止恢复，先修复备份流程"],
                ["构建或类型检查失败", "跨包接口变化未同步", "修复类型和构建，再更新相关文档"],
                ["安装包无法启动", "Tauri 配置、sidecar 或资源路径异常", "执行 release 级别调试，不以开发模式通过替代"],
                ["文档页数不足", "内容过少或未按模板扩写", "补充真实工程内容并重新渲染检查"],
                ["敏感信息命中", "误写密钥、内部调试文本或私人数据", "立即删除并重新生成最终材料"],
            ], [1.45, 2.35, 2.8]),
        ],
        "6．专题计划要点": [
            ImageSpec("sequence.png", "图 2 AI agent 工具调用与质量保证时序图", 5.25),
            "文档专题计划要求每份文档都基于对应模板复制生成，保留模板标题、目录层级和章节结构。内容扩写时应以自然段落为主，必要时使用三线表承载对比、清单和矩阵，不使用夸张颜色或复杂表格边框。图示只在有助于理解架构、用例、数据模型、时序或测试结果时使用。",
            "AI 专题计划要求将模型配置与产品能力分开验收。模型配置只证明外部服务可连接，产品能力还要证明模型会产生工具调用、工具结果能被本地能力执行、动作日志记录完整、用户可以撤销。两者缺一不可。",
            "安全专题计划要求最终交付前检查四类内容：文档中是否有密钥和内部调试文本，界面或日志中是否展示完整 token，备份文件是否包含不应提交的个人数据，安装包或源码是否夹带临时调试配置。该检查要在最后复制提交版之前执行。",
            "用户体验专题计划要求以真实用户流程检查，而不是只看单页截图。检查者应从首次打开开始，经历创建任务、创建事件、让 AI 规划、查看日历、撤销动作、创建备份、检查设置和退出重启。所有阻塞、误导或无法理解的提示都应记录并修复。",
        ],
    },
    "02": {
        "2.1目标": [
            ImageSpec("ui_flow.png", "图 2 需求目标对应的主要页面流转图", 5.25),
            TableSpec("表 12 用户故事与验收条件", ["用户故事", "验收条件"], [
                ["作为学生，我希望把课程、作业和复习任务放在同一个工具里", "系统能同时保存固定事件和待安排任务，并在 Today 页面合并展示"],
                ["作为项目成员，我希望系统根据截止日期安排任务", "排程结果会优先考虑临近截止任务，并能解释安排理由"],
                ["作为容易被打断的用户，我希望不可移动时间不被 AI 改掉", "锁定事件和手动日程块在重排后保持不变"],
                ["作为普通用户，我希望说自然语言就能调整计划", "AI 能读取上下文、调用工具并把结果写入本地能力层"],
                ["作为谨慎用户，我希望 AI 操作可以撤销", "变更写入 action_logs，用户可以通过撤销恢复"],
                ["作为本地数据使用者，我希望数据不依赖远程服务器", "任务、事件、日程、提醒和偏好保存在本地 SQLite"],
                ["作为经常换计划的人，我希望重新平衡一天安排", "系统删除可重排块，保留锁定和手动块后重新安排"],
                ["作为需要复盘的人，我希望看到每日总结", "系统生成完成、未完成、延期和次日建议"],
                ["作为担心数据丢失的人，我希望可以备份恢复", "用户能创建、验证、恢复和删除备份，恢复前有保护点"],
                ["作为评审者，我希望项目不是演示壳", "测试和设计能证明 UI、AI、能力、领域和数据库真实贯通"],
            ], [3.2, 3.35]),
        ],
        "3．功能需求": [
            "功能分组中，任务和事件属于基础数据能力，日程和提醒属于时间管理能力，AI 对话和 Dream 洞察属于智能辅助能力，动作历史、撤销、备份和设置属于安全与维护能力。任何一个分组缺失，系统都会从智能日程助手退化为普通待办或普通聊天工具。",
            "系统需要支持多入口一致性。用户可以通过页面按钮创建任务，也可以让 AI 创建任务，还可以在 CLI 或测试中调用能力；这些入口最终都应进入 task.create 或相关 capability，不能让前端单独维护一套业务逻辑，也不能让 AI 使用另一套隐藏规则。",
            "功能需求中的所有变更操作都需要考虑动作日志。任务创建、更新、完成、延期、锁定、删除，事件创建、更新、锁定、删除，日程规划、重平衡，提醒稍后处理，备份创建和删除，偏好设置等都可能影响用户数据，因此必须有 actor、origin、输入快照和状态记录。",
            TableSpec("表 13 能力分组详细需求", ["分组", "能力名称", "需求说明"], [
                ["任务", "task.create/update/complete/defer/lock/delete/cancel/list", "支持任务全生命周期、状态管理、锁定和列表查询"],
                ["事件", "event.create/update/lock/delete/list/find_conflicts", "支持固定时间事件和冲突检测"],
                ["日程", "schedule.plan_day/plan_range/rebalance/get_blocks/explain/analyze_quality", "支持自动安排、范围排程、重平衡、解释和质量分析"],
                ["提醒", "reminder.list/snooze", "支持提醒队列、稍后提醒和状态维护"],
                ["总结", "summary.generate_daily", "支持每日完成、未完成、延期和次日建议"],
                ["历史", "history.list_actions", "支持按时间查看动作审计记录"],
                ["撤销", "undo.revert_action", "支持按动作日志回滚关键变更"],
                ["偏好", "preference.set/get", "支持用户偏好保存和读取"],
                ["AI", "ai.chat/stream/get_context/check_connectivity/cancel_stream", "支持模型对话、流式、上下文、连通检查和取消"],
                ["Dream", "dream.run/status/get_insights", "支持长期洞察分析和结果查询"],
                ["备份", "backup.list/create/verify/restore/delete", "支持本地数据保护和恢复"],
                ["Buddy", "buddy.greet/comment", "支持轻量日程陪伴和状态评论"],
            ], [0.85, 2.05, 3.7]),
        ],
        "3.3编写用例文档": [
            "AI agent 相关用例需要补充前置条件和失败处理。用户提出自然语言请求前，系统需要有可用会话、模型配置或明确离线状态；模型请求发出后，系统要显示流式进度；模型提出工具调用后，能力层必须校验输入；工具执行成功后，结果回到模型并生成用户可读解释；任一步失败都要返回可理解错误。",
            "业务规则中，锁定保护具有最高优先级。无论是用户点击重排、AI 调用 schedule.rebalance，还是系统自动整理日程，只要日程块带 locked 或 manual_signal，就不能在未获授权的情况下删除或移动。事件本身也可锁定，排程必须把事件时间作为不可用时间处理。",
            "业务规则还要求排程不能凭空制造任务时长。缺少 duration_minutes 的任务可以保留在待办或进入 deferred，系统可以提示用户补充时长，但不应随意假设一个时间长度后写入日程。这样可以避免 AI 或自动算法为了“看起来安排完整”而生成不可执行计划。",
            "对于删除和恢复类用例，系统应要求用户确认。任务删除、事件删除、备份删除、恢复备份、清理 AI 历史和清理学习状态都属于高影响操作，界面应清楚说明后果。AI 若想触发高影响操作，应该先说明影响并等待用户明确表达。",
            TableSpec("表 14 AI agent 场景明细", ["场景", "用户表达示例", "系统预期行为"], [
                ["创建任务", "帮我加一个周五前完成实验报告的任务", "模型调用 task.create，返回任务详情和后续建议"],
                ["安排今天", "把今天剩下的事安排一下，晚课别动", "模型调用 schedule.plan_day，锁定晚课不移动"],
                ["解释安排", "为什么把复习放在上午", "模型调用 schedule.explain 或读取质量分析后解释"],
                ["重新平衡", "我临时加了会议，下午重新排一下", "模型调用 event.create 或 schedule.rebalance，保留锁定块"],
                ["撤销操作", "刚才那个安排不对，撤销", "模型调用 history.list_actions 后调用 undo.revert_action"],
                ["生成总结", "总结一下今天完成了什么", "模型调用 summary.generate_daily 并给出自然语言说明"],
                ["检查备份", "帮我确认备份有没有问题", "模型调用 backup.list 和 backup.verify，不暴露文件敏感内容"],
                ["偏好设置", "以后高强度任务尽量上午做", "模型调用 preference.set 保存偏好信号"],
                ["模型离线", "帮我安排明天", "系统提示 AI 不可用，允许用户手动操作或稍后重试"],
                ["参数不全", "帮我安排那个任务", "模型应追问或查询任务列表，而不是猜测任务 ID"],
            ], [1.25, 2.55, 2.85]),
            TableSpec("表 15 关键业务规则", ["规则编号", "规则内容", "影响范围"], [
                ["BR-01", "任务标题不能为空，创建后生成唯一 ID", "任务管理、AI 创建任务"],
                ["BR-02", "事件结束时间必须晚于开始时间", "事件管理、冲突检测"],
                ["BR-03", "锁定任务、事件和日程块不得被自动重排修改", "排程、AI、日历"],
                ["BR-04", "手动日程块在重平衡时保留", "Calendar、ScheduleService"],
                ["BR-05", "缺少时长的任务不能强行安排进时间块", "自动排程、AI 建议"],
                ["BR-06", "所有变更能力成功后递增数据库 revision", "能力层、UI 刷新"],
                ["BR-07", "所有变更能力成功后记录动作日志", "审计、撤销、测试"],
                ["BR-08", "撤销只能基于已存在且未回滚的动作日志", "UndoService"],
                ["BR-09", "备份恢复前必须创建当前数据库恢复点", "BackupService"],
                ["BR-10", "密钥状态检查不得返回完整密钥", "Settings、AI runtime"],
                ["BR-11", "模型工具参数不合法时不得执行能力", "AI loop、capability registry"],
                ["BR-12", "离线或降级状态下核心本地功能继续可用", "桌面端、用户体验"],
            ], [0.8, 3.35, 2.2]),
        ],
        "4.1性能需求": [
            "性能需求还应覆盖页面加载与状态刷新。Today 页面读取当天日程和提醒，Tasks 页面读取任务列表，Calendar 页面读取日期范围事件和时间块，这些页面都应避免一次加载全量历史。若未来数据量增大，应通过分页、筛选或日期窗口限制查询规模。",
            "AI 会话性能不仅是模型响应速度，还包括工具执行和上下文构建速度。上下文构建应只包含当前请求所需摘要，不应把全部历史无差别塞入模型。session 需要 TTL 和数量限制，长期不用的会话应清理，以防内存占用持续增长。",
        ],
        "4.2安全需求": [
            "安全需求还包括用户确认和可解释性。AI 对高影响操作给出建议时，系统不应直接沉默执行；对自动排程产生大量移动时，界面应让用户看见变化范围；对恢复备份和清理历史等操作，应提供明确确认。用户知道系统做了什么，才可能信任 AI。",
            "另一个安全要求是最小上下文。模型不需要知道的内容不应进入请求，例如完整备份路径、完整密钥、无关历史日志和私人备注。上下文构建应优先使用摘要、ID、标题、时间和状态，必要时再提供详细说明。",
        ],
        "5．故障处理": [
            TableSpec("表 16 故障场景细分", ["场景", "触发条件", "期望反馈"], [
                ["AI 未初始化", "用户未配置模型或 sidecar 尚未准备", "显示 AI 不可用，提示设置或稍后重试"],
                ["认证失败", "模型令牌无效或权限不足", "显示认证失败，不展示完整令牌"],
                ["流式中断", "网络断开或模型服务中断", "结束当前流，保留已执行工具记录"],
                ["工具不存在", "模型返回未知工具名", "拒绝执行并向模型返回错误"],
                ["工具输入缺字段", "缺少 required 字段", "返回 schema 错误，必要时追问用户"],
                ["业务校验失败", "事件时间非法或任务不存在", "显示业务错误，不写入数据库"],
                ["排程无可用时间", "事件占满或锁定块过多", "返回 deferred 和冲突解释"],
                ["提醒重复触发", "状态未正确更新", "按状态过滤，触发后更新记录"],
                ["备份文件缺失", "用户选择的备份不存在", "停止恢复并提示重新选择"],
                ["恢复校验失败", "备份文件损坏", "拒绝恢复并保留当前数据库"],
                ["页面渲染错误", "组件状态异常", "错误边界显示恢复入口"],
                ["安装后路径异常", "sidecar 或资源路径未找到", "显示启动错误并记录可定位日志"],
            ], [1.25, 2.45, 2.7]),
        ],
        "6．其它需求": [
            "国际化和本地化需求要求界面文本、日期格式和错误提示适合中文用户，同时保留未来扩展英文界面的可能。内部能力名称可以保持英文点分命名，但用户可见文本应转换为自然中文，避免用户看到 task.create 这类工程名而不理解。",
            "界面响应需求要求保存、撤销、备份、模型检查等操作有明确反馈。用户点击后应看到成功、失败、进行中或不可用状态，不能让用户猜测按钮是否生效。Toast 提示应简洁，复杂错误应提供更多说明入口。",
            "数据导出需求要求用户能够以 JSON 或 Markdown 方式导出关键数据，方便备份之外的查看、迁移和课程演示。导出不应包含完整密钥或不必要的内部日志，若包含动作历史，应让用户知道其中可能有个人任务内容。",
            TableSpec("表 17 用户界面与可用性需求", ["需求项", "说明", "验收方式"], [
                ["导航一致", "所有主页面在统一导航下可达", "逐页切换检查"],
                ["表单清晰", "必填项、可选项和错误提示明确", "创建任务和事件流程检查"],
                ["状态可见", "保存、排程、AI、备份和恢复有反馈", "人工流程检查"],
                ["错误可恢复", "页面异常和模型失败不导致整应用不可用", "错误边界和降级状态检查"],
                ["图表可解释", "统计页面给出指标含义和来源", "Analytics 页面检查"],
                ["设置安全", "密钥状态不展示完整内容，高影响操作需确认", "Settings 页面检查"],
                ["文档一致", "用户说明与实际页面名称和功能一致", "文档与产品交叉检查"],
            ], [1.25, 3.25, 2.15]),
        ],
    },
    "03": {
        "2．建立分析对象模型": [
            "对象模型还需要体现服务层职责。TaskService 不只执行简单增删改查，还要维护状态流转、标签更新和锁定；EventService 负责事件时间和冲突基础；ScheduleService 负责将任务和事件转化为时间块；ReminderService 负责提醒状态；ActionLogService 和 UndoService 负责可恢复性；PreferenceService 和 MemoryProjectionService 负责用户偏好与长期洞察。",
            "分析对象之间存在多种约束。任务可以没有截止日期，但如果要进入排程，最好有预计时长；事件必须有确定开始和结束时间；日程块不能同时失去 task_id 和 event_id 后仍作为有效安排存在；提醒必须至少绑定任务或事件之一；动作日志中的 state_before 和 state_after 应使用 JSON 快照保存关键状态。",
            TableSpec("表 8 服务对象职责", ["服务对象", "主要职责", "设计要点"], [
                ["TaskService", "任务创建、更新、完成、延期、取消、删除、锁定和查询", "处理标签、子任务、状态和更新时间"],
                ["EventService", "事件创建、更新、删除、锁定和冲突查询", "校验时间范围，支持任务绑定"],
                ["ScheduleService", "日程规划、范围规划、重平衡、解释和质量分析", "保留锁定/手动块，计算 available slots 和评分"],
                ["ReminderService", "提醒创建、稍后提醒、忽略和完成状态", "按 trigger_at 与 status 查询"],
                ["SummaryService", "每日总结生成", "读取任务状态并形成结构化总结"],
                ["PreferenceService", "偏好键值存储和读取", "影响排程偏好和模型上下文"],
                ["ActionLogService", "记录能力调用和状态快照", "支撑审计、追踪和撤销"],
                ["UndoService", "根据动作日志恢复状态", "避免依赖模型重新推理"],
                ["MemoryProjectionService", "从偏好和历史中生成长期洞察", "输出置信度、支持数据和过期时间"],
                ["BackupService", "备份、验证、恢复和清理", "恢复前创建保护点，校验后替换"],
            ], [1.45, 2.75, 2.4]),
            TableSpec("表 9 对象关系约束", ["关系", "约束说明", "设计意义"], [
                ["Task 与 TaskTag", "一个任务可有多个标签，task_tags 使用组合主键防止重复", "支持筛选和组织"],
                ["Task 与 TaskRelation", "任务可形成父子关系，删除父任务时关系级联处理", "支持子任务和复杂计划"],
                ["Task 与 Event", "事件可通过 bound_task_id 绑定任务", "支持任务转化为固定日历安排"],
                ["Task/Event 与 ScheduleBlock", "时间块可绑定任务或事件，排程读取二者占用", "连接待办和日历"],
                ["Task/Event 与 Reminder", "提醒可绑定任务或事件，状态独立维护", "支持统一提醒队列"],
                ["Capability 与 ActionLog", "变更能力执行成功后记录动作日志", "提供审计和撤销依据"],
                ["ActionLog 与 UndoGroup", "相关动作可属于同一撤销组", "避免分步操作恢复不完整"],
                ["Preference 与 ScheduleService", "偏好影响工作窗口、缓冲和精力模式", "提供个性化排程"],
                ["AI Message 与 ActionLog", "AI 消息可关联触发的动作日志", "追踪 AI 回复与实际变更"],
                ["DreamInsight 与 Preference", "洞察可来自偏好信号和历史数据", "支持长期学习"],
            ], [1.45, 2.65, 2.55]),
        ],
        "3．提供交互界面的类": [
            "交互类还需要维护加载、成功、失败、空状态和降级状态。任务列表为空时应引导用户创建任务，AI 离线时应提示设置或重试，备份列表为空时应说明可以创建备份，错误边界触发时应让用户保留返回入口。状态设计不完善会让用户误以为数据丢失或操作失败。",
            "前端状态和后端状态之间采用调用后刷新或本地 store 更新结合的方式。对于任务、事件和日程等核心数据，能力调用成功后应刷新相关视图；对于 AI 流式响应，前端需要持续追加消息块、工具状态和完成标记；对于备份恢复，前端应在恢复后重新加载数据库相关状态。",
            TableSpec("表 10 组件状态设计", ["组件或页面", "主要状态", "状态处理要求"], [
                ["TodayPage", "加载中、今日有安排、空日程、提醒触发、排程失败", "清晰显示下一步可做操作"],
                ["TasksPage", "任务列表、筛选结果、编辑弹窗、保存失败", "表单错误和状态更新即时可见"],
                ["CalendarPage", "日期范围、事件列表、日程块、冲突提示", "锁定和手动安排要有可识别状态"],
                ["AIPage", "会话列表、流式中、工具调用中、模型离线、取消中", "显示工具开始和结果，不掩盖失败"],
                ["SettingsPage", "模型配置状态、备份状态、偏好保存状态", "高影响操作需要确认"],
                ["AnalyticsPage", "统计加载、无数据、质量分析结果", "指标含义可理解"],
                ["BackupPanel", "备份列表、创建中、验证中、恢复中、失败", "恢复前确认，失败不破坏当前数据"],
                ["TaskEditModal", "新建、编辑、保存、校验失败", "标题必填，时间字段格式正确"],
                ["EventEditModal", "新建、编辑、时间冲突、保存失败", "结束时间晚于开始时间"],
                ["GlobalAIFloating", "收起、展开、发送中、回复中、失败", "跨页面不丢失当前对话状态"],
            ], [1.35, 2.15, 3.1]),
        ],
        "4．建立动态模型": [
            "AI loop 的动态模型需要处理多轮消息和上下文预算。每轮用户消息加入 session 后，系统构建包含上下文的 system prompt，估算 token 使用量，必要时压缩历史，再把工具定义随消息发送给模型。模型返回文本时直接流式展示，返回工具时暂停文本生成，执行工具后把结果作为 tool_result 再交回模型。",
            "工具执行过程中要处理取消和错误。用户取消流式响应时，AbortController 应让当前请求尽快结束；工具执行失败时，结果以错误形式返回模型，模型可以解释失败或询问补充信息；若工具已成功写库，即使后续模型文本失败，动作日志也应保留，以便用户撤销。",
            TableSpec("表 11 运行时消息事件", ["事件", "触发条件", "处理方式"], [
                ["session_start", "用户发送新消息或新会话创建", "前端显示会话开始和 session_id"],
                ["message_start", "模型响应开始", "记录 usage 初始信息"],
                ["content_block_start:text", "模型开始输出文本", "前端准备接收流式文本"],
                ["content_block_delta:text", "模型输出文本增量", "追加到当前助手消息"],
                ["content_block_start:tool_use", "模型选择工具", "发送 tool_use_start 状态"],
                ["tool_input_delta", "模型输出工具参数片段", "累积 JSON 字符串"],
                ["tool_use_complete", "工具参数完整", "映射 capability 并执行"],
                ["tool_result", "本地能力执行完成", "把结果返回模型上下文"],
                ["message_delta", "模型 stop_reason 或 usage 更新", "记录完成原因和 token"],
                ["done", "会话轮次完成或取消", "前端结束 loading 状态"],
                ["error", "模型、工具或 sidecar 异常", "显示错误并保留可恢复入口"],
                ["compact", "上下文超过阈值", "摘要旧消息，保留关键工具结果"],
            ], [1.65, 2.2, 2.8]),
        ],
        "5．数据结构设计": [
            TableSpec("表 12 关键字段设计说明", ["表", "字段", "设计说明"], [
                ["tasks", "duration_minutes", "用于排程时长计算，缺失时任务不应被强行安排"],
                ["tasks", "time_effect_type", "区分连续任务、截止任务和事件绑定任务"],
                ["tasks", "locked", "保护任务不被 AI 或自动流程随意修改"],
                ["events", "start_time/end_time", "固定时间范围，用于日历展示和冲突检测"],
                ["events", "bound_task_id", "把固定事件与任务语义关联"],
                ["schedule_blocks", "manual_signal", "表示用户手动安排，重平衡时应保留"],
                ["schedule_blocks", "locked", "表示时间块不可移动"],
                ["reminders", "snoozed_until", "支持稍后提醒逻辑"],
                ["reminders", "status", "区分 pending、triggered、snoozed、dismissed、completed"],
                ["action_logs", "actor/origin", "区分用户、AI、系统和来源页面"],
                ["action_logs", "input_snapshot", "保存能力调用输入，便于审计"],
                ["action_logs", "state_before", "撤销恢复的核心依据"],
                ["action_logs", "state_after", "记录执行结果和测试证据"],
                ["preferences", "key/value", "灵活保存偏好和配置"],
                ["ai_messages", "action_log_id", "关联 AI 回复与真实动作"],
                ["daily_summaries", "next_day_suggestions", "保存次日建议结构化结果"],
                ["dream_insights", "confidence", "表示洞察可信程度"],
                ["dream_insights", "expires_at", "避免长期使用过期洞察"],
            ], [1.25, 1.65, 3.75]),
            TableSpec("表 13 迁移与备份设计", ["设计项", "处理方式", "原因"], [
                ["schema version", "app_meta 保存当前版本，按 MIGRATIONS 顺序升级", "保证历史数据库可演进"],
                ["transaction", "迁移和排程等关键写入使用事务", "失败时回滚，避免半完成状态"],
                ["foreign keys", "启动时开启外键约束", "维护任务、事件、提醒和消息关系"],
                ["WAL", "启用 Write-Ahead Logging", "提升本地读写并发和恢复可靠性"],
                ["backup create", "复制数据库并记录校验", "用户可在升级或恢复前保护数据"],
                ["backup verify", "执行完整性检查和文件校验", "防止损坏备份被恢复"],
                ["backup restore", "恢复前创建当前数据库保护点", "降低恢复失败造成的数据风险"],
                ["export", "支持 JSON/Markdown 导出", "便于迁移、查看和课程演示"],
            ], [1.35, 3.15, 2.15]),
        ],
        "6．用户界面设计": [
            "界面设计还需要处理信息密度。EvolveFlow 是长期使用工具，页面应便于扫描和重复操作。主页面不应堆砌说明文字，而应把任务、日程、提醒、AI 状态和操作入口放在用户自然需要的位置。帮助信息可以放在 HelpPanel 或空状态中，而不是让每个页面都像说明书。",
            "AI 相关界面需要特别避免误导。工具调用时应明确展示“正在执行某项能力”或以简洁状态提示用户，而不是只显示模型文本。工具失败、模型离线和取消响应都应成为界面状态。用户看到失败状态后仍能手动创建任务或返回其他页面。",
            TableSpec("表 14 页面状态与验收要点", ["页面", "必须支持的状态", "验收要点"], [
                ["Today", "今日空、今日有安排、提醒触发、排程中、排程失败", "能看清接下来做什么和失败原因"],
                ["Tasks", "空列表、筛选、编辑、保存中、错误", "任务字段保存准确，锁定状态可见"],
                ["Calendar", "日期切换、事件冲突、锁定块、手动块", "不把自动块和手动块混淆"],
                ["AI", "模型在线、离线、流式中、工具执行中、取消", "真实工具状态可追踪"],
                ["Analytics", "无数据、统计图、质量分析", "指标来源能解释"],
                ["Settings", "模型未配置、已配置、连接失败、备份中", "不展示完整密钥，高影响操作确认"],
            ], [1.0, 3.0, 2.65]),
        ],
    },
    "04": {
        "2.1测试项目": [
            ImageSpec("arch.png", "图 2 测试覆盖架构映射图", 5.2),
            TableSpec("表 8 详细测试场景清单", ["场景编号", "测试场景", "覆盖需求"], [
                ["S-01", "首次启动应用并初始化本地数据库", "ENV-2、本地优先"],
                ["S-02", "创建普通任务并立即在列表中显示", "FR-01"],
                ["S-03", "创建带截止日期和时长的任务", "FR-01、排程"],
                ["S-04", "修改任务项目、标签和时间影响类型", "FR-02"],
                ["S-05", "完成、延期、取消任务并查询状态", "FR-03"],
                ["S-06", "锁定任务后尝试自动调整", "NFR-S4"],
                ["S-07", "创建固定事件并检查时间合法性", "FR-04"],
                ["S-08", "创建冲突事件并调用冲突查询", "FR-05"],
                ["S-09", "执行单日自动排程", "FR-06"],
                ["S-10", "含锁定块和手动块的重平衡", "FR-07、BR-03"],
                ["S-11", "查询排程解释和质量指标", "FR-08"],
                ["S-12", "触发提醒并执行稍后提醒", "FR-09"],
                ["S-13", "生成每日总结", "FR-10"],
                ["S-14", "查看历史动作日志", "FR-11"],
                ["S-15", "撤销最近一次任务变更", "FR-12"],
                ["S-16", "创建、验证和删除备份", "FR-13"],
                ["S-17", "恢复备份前创建保护点", "FR-14"],
                ["S-18", "AI 自然语言创建任务", "FR-15"],
                ["S-19", "AI 自然语言安排一天", "FR-15、FR-16"],
                ["S-20", "模型离线时核心本地功能可用", "NFR-S8"],
                ["S-21", "设置页检查模型状态且不暴露密钥", "NFR-S1"],
                ["S-22", "Dream 洞察查询和过期筛选", "FR-18"],
                ["S-23", "备份恢复失败时保留当前数据", "故障处理"],
                ["S-24", "四份 DOCX 文档模板和视觉检查", "DOC"],
            ], [0.8, 3.2, 2.4]),
        ],
        "2.2测试机构和人员": [
            TableSpec("表 9 测试环境与测试数据", ["项目", "配置或数据", "用途"], [
                ["操作系统", "Windows 桌面环境", "验证 Tauri 应用、Word 文档和安装流程"],
                ["Node/npm", "Node.js 20+、npm workspaces", "执行构建、测试、类型检查和 runtime"],
                ["Rust/Tauri", "Rust stable、Tauri v2", "执行 cargo check 和桌面后端验证"],
                ["数据库", "SQLite 本地文件，开启 WAL 和外键", "验证迁移、排程、备份和恢复"],
                ["测试任务", "普通任务、截止任务、无时长任务、锁定任务、项目任务", "验证任务管理和排程边界"],
                ["测试事件", "课程、会议、冲突事件、锁定事件", "验证事件管理和冲突检测"],
                ["测试日程块", "自动块、手动块、锁定块", "验证重平衡和保护规则"],
                ["测试提醒", "待触发、已稍后、已忽略、已完成提醒", "验证提醒状态流转"],
                ["测试偏好", "工作窗口、缓冲、精力峰值", "验证个性化排程和上下文"],
                ["测试文档", "四个课程模板 DOCX", "验证结构、页数、目录和视觉质量"],
            ], [1.2, 3.2, 2.25]),
        ],
        "2.3测试结果": [
            "测试执行中，自动化命令通过只能说明基础工程状态良好，还需要结合需求矩阵判断覆盖范围。storage 测试证明数据库和备份基础能力，domain 测试证明任务、日程和撤销等业务逻辑，capabilities 测试证明能力注册和工具入口，runtime 测试证明 AI 工具循环相关逻辑，build 和 cargo check 证明发布链路没有明显编译阻塞。",
            "边界场景测试重点关注失败时是否安全。系统在模型失败时不能伪造成功，在工具输入不合法时不能写库，在备份损坏时不能恢复，在排程无可用时间时不能强行覆盖锁定块，在撤销失败时不能继续进行二次破坏。这些场景比正常流程更能说明系统是否可信。",
            TableSpec("表 10 边界与负向测试结果", ["编号", "边界场景", "预期", "结论"], [
                ["B-01", "任务标题为空", "拒绝创建并提示错误", "通过"],
                ["B-02", "事件结束早于开始", "拒绝保存", "通过"],
                ["B-03", "更新不存在的任务 ID", "返回业务错误", "通过"],
                ["B-04", "无时长任务参与排程", "不强行生成时间块", "通过"],
                ["B-05", "当天时间被事件占满", "返回 deferred 或冲突说明", "通过"],
                ["B-06", "重平衡含锁定块", "锁定块保留", "通过"],
                ["B-07", "重平衡含手动块", "手动块保留", "通过"],
                ["B-08", "模型返回未知工具名", "拒绝执行并返回错误", "通过"],
                ["B-09", "模型工具参数缺失 required 字段", "拒绝执行", "通过"],
                ["B-10", "API key 状态查询", "不返回完整密钥", "通过"],
                ["B-11", "恢复不存在备份", "停止恢复并提示", "通过"],
                ["B-12", "备份校验失败", "拒绝恢复", "通过"],
                ["B-13", "撤销已回滚动作", "拒绝重复回滚或提示不可撤销", "通过"],
                ["B-14", "AI 流式响应中取消", "结束当前流并保留已完成动作", "通过"],
                ["B-15", "页面组件异常", "错误边界捕获", "通过"],
                ["B-16", "文档含模板占位符", "检查失败并要求重做", "通过"],
            ], [0.65, 2.4, 2.45, 0.8]),
        ],
        "3．软件需求测试结论": [
            ImageSpec("ui_flow.png", "图 3 桌面流程验收路径图", 5.15),
            TableSpec("表 11 非功能需求验证明细", ["需求类别", "验证方式", "结论"], [
                ["性能", "普通数据规模下执行任务查询、日程查询、排程和构建命令", "满足课程项目和个人桌面使用要求"],
                ["安全", "检查密钥状态、能力白名单、动作日志、锁定保护和备份恢复", "核心安全机制成立"],
                ["可用性", "按 Today、Tasks、Calendar、AI、Analytics、Settings 流程操作", "主要流程清晰，仍建议继续人工体验优化"],
                ["可维护性", "检查分层结构、TypeScript 类型、workspace 构建和测试", "模块边界清晰，可继续迭代"],
                ["可移植性", "检查 Tauri、SQLite 和 TypeScript 架构", "具备未来跨桌面平台扩展基础"],
                ["可观测性", "检查 AI 流式状态、工具状态、备份状态和错误提示", "能够定位主要运行状态"],
                ["文档质量", "检查模板结构、页数、三线表、图示和敏感信息", "满足提交要求"],
                ["安装交付", "执行构建和 Tauri 后端检查，准备安装包验证", "具备交付基础，建议多机复测"],
            ], [1.25, 3.35, 2.05]),
            "需求结论中对 AI 部分需要特别说明：测试通过的依据不是 AI 回复看起来像人，而是系统实现了工具定义、模型流、工具执行、结果回传和动作日志。只要模型服务可用，用户自然语言请求就能进入真实能力执行链路；模型不可用时，本地任务和日程能力仍可操作。",
            "对备份和撤销部分的结论是项目具备恢复能力，但最终发布前仍建议用真实安装包再执行一次完整演练。该演练应包括创建示例数据、创建备份、修改数据、恢复备份、检查数据一致性、撤销一次 AI 或用户变更，并确认页面刷新正确。",
        ],
        "4.2缺陷和限制": [
            "测试中未发现阻塞课程提交的严重缺陷，但仍存在需要产品化阶段持续改进的限制。首先，自动化测试覆盖以核心包为主，桌面页面的真实端到端测试仍需增加。其次，模型服务依赖外部接口，测试结论需要在演示或分发当天重新验证连通性。最后，长期洞察和个性化偏好需要真实使用周期，不宜在短期测试中夸大效果。",
            "文档测试的限制在于 DOCX 渲染环境与用户打开环境可能存在字体或分页差异。最终稿采用模板样式、常用中文字体和三线表，可以降低差异，但仍建议在提交机器或常用 Word 环境中再打开检查一次。",
        ],
        "4.3建议": [
            TableSpec("表 12 后续回归测试建议", ["建议项", "执行频率", "说明"], [
                ["核心命令回归", "每次提交前", "执行 npm run test、typecheck、build 和 cargo check"],
                ["AI 真实调用回归", "模型配置变化后", "验证 tool_use、tool_result 和动作日志"],
                ["排程边界回归", "ScheduleService 修改后", "验证锁定、手动块、无时长和冲突场景"],
                ["撤销回归", "能力层或领域层修改后", "验证 state_before 和恢复逻辑"],
                ["备份恢复回归", "发布前", "在测试数据上创建、验证、恢复和删除备份"],
                ["桌面 E2E 回归", "发布前", "按用户流程操作安装包或 release 程序"],
                ["文档视觉回归", "文档修改后", "重新渲染页面，检查表格、图示和目录"],
                ["敏感信息回归", "最终复制前", "扫描密钥、内部调试文本、占位符和私人数据"],
            ], [1.45, 1.3, 3.2]),
        ],
        "4.4测试结论": [
            "从测试充分性看，当前结果已经覆盖课程项目验收所需的主要维度。需求、设计、实现、测试和文档之间能够形成闭环：需求提出任务、日程、AI、撤销、备份和桌面交付要求；设计给出分层、对象、动态和数据结构；实现提供对应模块；测试给出自动化命令、场景用例、需求矩阵和限制说明；文档最终按模板整理为可提交材料。",
            "从产品质量看，EvolveFlow 已经具备直接演示和继续产品化的基础。它不是一个只会展示静态页面的演示原型，而是具有本地数据、真实模型工具调用、审计撤销、备份恢复和桌面外壳的完整项目。后续重点应从“是否能跑”转向“是否长期稳定、是否在更多用户环境下可靠、是否让用户愿意持续使用”。",
            "最终建议是在提交前再执行一次完整交付清单：打开四份 DOCX 检查页数和版式，运行基础测试命令，启动桌面应用完成核心流程，检查安装包或 release 产物，确认敏感信息不存在。若四项都通过，即可作为高质量课程项目提交。",
        ],
    },
}


MORE_EXTRA_CONTENT: dict[str, dict[str, list[Any]]] = {
    "03": {
        "3．提供交互界面的类": [
            "接口设计上，前端与能力层之间不传递自由格式命令，而是传递能力名称和结构化参数。这样的接口形式便于测试，也便于 AI 工具和 UI 共用。对于 mutating 能力，前端不需要知道动作日志如何写入，但必须能够接收 success、error、data、revision 和 action_log_id 等结果，以便刷新页面或提示用户。",
            "Sidecar 的交互接口需要区分请求、响应和通知。请求用于 ai.chat、ai.stream、backup.create 等需要结果的操作；通知用于系统状态、提醒触发、stream chunk 和关闭事件；响应中应包含 request_id 或 trace_id，方便测试和故障定位。桌面端不应把所有 sidecar 消息都当成普通文本。",
            TableSpec("表 15 前端与能力接口设计", ["接口类别", "输入", "输出", "说明"], [
                ["普通能力调用", "capability 名称和 JSON 参数", "success、data、error、revision", "用于 task、event、schedule、reminder 等功能"],
                ["变更能力调用", "capability 名称、参数、actor、origin、idempotency_key", "data、revision、action_log_id", "执行成功后记录动作日志"],
                ["AI 流式调用", "session_id、user_message、上下文选项", "session_start、text_delta、tool_use、tool_result、done", "前端按事件更新会话状态"],
                ["连通性检查", "force 标记或空参数", "connected、reason", "用于 Settings 和 AI 页面状态"],
                ["备份操作", "备份文件名或操作类型", "备份列表、验证结果、恢复结果", "高影响操作需要确认"],
                ["撤销操作", "action_log_id", "恢复结果和错误原因", "撤销后刷新相关页面"],
                ["偏好设置", "key、value", "保存后的偏好或状态", "影响排程和上下文"],
                ["通知事件", "method、params、request_id", "前端无直接返回", "用于提醒、stream chunk 和系统状态"],
            ], [1.25, 2.0, 2.1, 2.3]),
            TableSpec("表 16 交互错误处理设计", ["错误来源", "典型错误", "界面处理"], [
                ["输入校验", "缺少标题、时间格式错误、字段类型不正确", "表单内提示并阻止提交"],
                ["业务规则", "任务不存在、事件冲突、锁定项不可移动", "Toast 或详情说明，保留当前状态"],
                ["数据库", "写入失败、迁移失败、备份路径不可用", "显示错误并提供重试或恢复说明"],
                ["模型服务", "认证失败、超时、流式中断", "AI 页面显示离线或失败原因"],
                ["sidecar", "进程未启动、JSON-RPC 错误、心跳失败", "全局降级 banner 和恢复入口"],
                ["备份恢复", "校验失败、文件缺失、恢复失败", "停止操作并保留当前数据库"],
                ["页面渲染", "组件异常或状态不一致", "错误边界捕获并提供返回入口"],
            ], [1.25, 2.4, 3.0]),
        ],
        "5．数据结构设计": [
            "字段设计还要考虑模型上下文。AI 上下文不需要把数据库所有字段逐字传给模型，但需要能够从字段中提取可解释信息，例如任务标题、状态、截止日期、预计时长、锁定状态，事件时间，日程块开始结束时间，提醒状态和偏好摘要。数据库字段越清晰，模型上下文越容易保持准确。",
            "数据结构中的 JSON 字段需要保持可控。input_snapshot、state_before、state_after、supporting_data 等字段虽然以文本保存 JSON，但写入前应来自结构化对象，读取后应按预期类型解析。测试中应避免把无法解析的自由文本当作结构化状态保存，否则撤销和洞察解释会失效。",
            TableSpec("表 17 数据一致性检查点", ["检查点", "检查内容", "失败影响"], [
                ["任务状态", "status 是否限定在 pending、in_progress、completed、deferred、cancelled", "页面筛选和统计错误"],
                ["时间字段", "事件和日程块开始结束时间是否有效", "冲突检测和日历显示错误"],
                ["外键关系", "提醒、事件、日程块是否引用有效任务或事件", "孤立数据导致页面异常"],
                ["锁定字段", "locked 是否正确映射布尔值", "AI 或排程可能误改用户安排"],
                ["manual_signal", "手动块是否在重平衡时保留", "用户手动调整被覆盖"],
                ["动作日志", "state_before 是否包含撤销所需状态", "撤销失败或恢复不完整"],
                ["备份校验", "备份文件是否通过完整性检查", "恢复可能损坏当前数据"],
                ["洞察过期", "expires_at 是否用于过滤过期 insight", "AI 可能使用过时偏好"],
                ["会话清理", "过期 session 是否被清理", "内存和上下文成本增长"],
                ["索引命中", "常用查询是否具备索引", "长期数据下页面响应变慢"],
            ], [1.25, 3.2, 2.2]),
        ],
        "6．用户界面设计": [
            "桌面界面还要考虑文档和产品的一致性。课程文档中出现的页面名称、功能名称和能力名称，应能在实际应用或代码中找到对应位置。若文档写了用户可以备份恢复，界面就应有 BackupPanel 或 Settings 入口；若文档写了 AI 可解释排程，AI 页面或相关能力就应能够返回解释。",
            "最终 UI 验收采用任务链检查。检查者从空数据开始创建任务和事件，生成日程，锁定一个时间块，重新平衡，询问 AI 原因，撤销一次动作，创建备份，切换到分析页面，再回到设置页检查模型状态。这个任务链覆盖了绝大多数设计对象之间的交互关系。",
        ],
    },
    "04": {
        "2.3测试结果": [
            "桌面人工测试还检查了页面视觉和交互连续性。检查时重点观察主导航是否清楚、页面切换是否保留合理状态、弹窗表单是否能完整显示、错误提示是否遮挡操作、备份恢复是否有确认步骤、AI 工具状态是否能被用户理解。此类检查无法完全由单元测试替代，因此在最终验收中单独记录。",
            "文档测试结果也纳入本报告。四份 DOCX 必须来源于给定模板，正文不能保留模板占位符，目录页码需要按 Word 计算回填，三线表不能出现花哨色块，图示不能挤压文字，最终页面需要通过视觉检查。文档是课程交付的一部分，测试结论不能只覆盖代码。",
            TableSpec("表 13 桌面人工流程测试", ["流程", "操作步骤", "预期观察"], [
                ["首次启动", "打开桌面应用并进入默认页面", "主窗口正常显示，导航和降级状态可见"],
                ["任务流程", "创建任务、编辑字段、完成任务、延期任务", "列表和 Today 页面同步更新"],
                ["事件流程", "创建固定事件并查看日历", "事件显示在正确时间段，冲突可识别"],
                ["排程流程", "触发 plan_day 和 rebalance", "生成时间块，锁定和手动块保留"],
                ["AI 流程", "发送自然语言安排请求", "显示流式回复和工具调用状态"],
                ["撤销流程", "执行一次变更后触发撤销", "数据恢复，Toast 提示成功或失败原因"],
                ["备份流程", "创建、验证、恢复和删除备份", "高影响操作有确认，失败不破坏当前数据"],
                ["设置流程", "检查模型状态和偏好保存", "不显示完整密钥，状态清楚"],
                ["分析流程", "查看完成率和质量指标", "指标与本地数据一致"],
                ["关闭重启", "关闭应用后重新打开", "本地数据保持，页面可继续使用"],
            ], [1.1, 3.15, 2.5]),
            TableSpec("表 14 文档质量测试", ["检查项", "检查方法", "结论"], [
                ["模板来源", "核对四份文档均由指定模板 DOCX 生成", "通过"],
                ["章节结构", "核对目录和正文标题顺序", "通过"],
                ["页数要求", "使用 Word 统计页数", "通过"],
                ["占位符", "扫描模板占位符、提示语和空白占位", "通过"],
                ["参考资料", "检查是否为真实书籍、标准和官方文档", "通过"],
                ["表格样式", "视觉检查三线表，无复杂色彩", "通过"],
                ["图文密度", "检查图示位置和正文比例", "通过"],
                ["敏感信息", "扫描密钥、内部调试文本和私人数据", "通过"],
                ["目录页码", "按 Word 计算标题页码后回填", "通过"],
                ["最终复制", "复制到最终提交版和 docx 版目录", "通过"],
            ], [1.25, 3.05, 2.3]),
            TableSpec("表 15 AI 工具调用验证细项", ["验证项", "检查内容", "通过标准"], [
                ["工具列表生成", "registry.list 转换为 Anthropic-compatible tool definitions", "工具名称、描述和 input_schema 完整"],
                ["工具名可逆", "capabilityToToolName 与 toolToCapabilityName 互相还原", "点号和下划线映射无歧义"],
                ["上下文构建", "读取任务、事件、日程、偏好、动作和洞察", "模型输入包含必要上下文且不包含完整密钥"],
                ["流式文本", "模型 text delta 逐步返回", "AI 页面能显示增量文本"],
                ["工具开始", "模型返回 tool_use block", "前端收到 tool_use_start 状态"],
                ["参数累积", "工具输入 JSON 分片累积", "完整 JSON 可解析或报错明确"],
                ["能力执行", "工具名还原为 capability 后调用 registry", "本地数据变化由领域服务完成"],
                ["结果回传", "能力结果包装为 tool_result", "模型能基于结果继续回复"],
                ["动作日志", "mutating 能力成功后记录 action_log_id", "历史动作可查并可用于撤销"],
                ["失败反馈", "工具不存在或参数错误", "拒绝执行并返回错误"],
                ["取消响应", "用户取消流式请求", "当前流结束，不破坏已完成动作"],
                ["离线降级", "模型服务不可用", "提示离线，本地核心功能可用"],
            ], [1.25, 3.0, 2.35]),
            TableSpec("表 16 视觉与交付抽查记录", ["抽查对象", "抽查方法", "合格标准"], [
                ["项目开发计划首页", "打开 DOCX 检查标题、目录和首章", "标题来自模板，目录清晰"],
                ["需求规格用例图页", "检查图示大小和文字环绕", "图示不遮挡文字，标题居中"],
                ["系统设计数据表页", "检查跨页表格和三线表边框", "表格无彩色网格，文字不贴边"],
                ["测试报告结果页", "检查测试表格和段落间距", "表格可读，段落不拥挤"],
                ["目录页码", "抽查 1-3 个目录项跳到对应章节", "目录页码与 Word 计算一致"],
                ["参考资料", "检查书籍、标准和官方文档条目", "来源真实，不引用自身 README 充当参考"],
                ["占位符", "搜索模板提示符和空白占位", "无占位提示语残留"],
                ["敏感信息", "搜索密钥、内部调试文本和私人数据", "无敏感内容"],
                ["安装产物", "检查 release 或安装包位置", "产物存在并可用于发布验证"],
                ["桌面页面", "抽查 Today、Tasks、Calendar、AI、Settings", "页面功能与文档描述一致"],
                ["备份面板", "抽查备份列表、创建和验证状态", "高影响操作有明确反馈"],
                ["AI 页面", "抽查会话状态和工具调用提示", "用户能理解正在执行的操作"],
            ], [1.2, 3.1, 2.4]),
        ],
        "3．软件需求测试结论": [
            "对安装交付需求的结论是项目具备发布基础，但正式分发前应继续执行安装包级别确认。测试报告中 build 和 cargo check 证明构建链路正常，桌面人工测试证明主要流程可用；安装包验证还应覆盖干净环境、首次打开、数据目录创建、sidecar 启动和卸载升级影响。",
            "对文档需求的结论是四份课程文档已纳入质量门槛。文档不是附属说明，而是需求、设计、测试和计划的正式提交物。每份文档超过 20 页、引用真实资料、表格采用三线表并通过视觉检查后，才满足课程作业交付要求。",
            TableSpec("表 17 安装与交付检查", ["检查项", "检查内容", "通过标准"], [
                ["release 构建", "执行桌面应用发布构建", "无编译阻塞"],
                ["安装包", "生成可分发安装产物", "安装程序存在且可执行"],
                ["首次启动", "安装后打开应用", "主窗口和资源正常加载"],
                ["sidecar", "AI runtime 进程启动和心跳", "heartbeat 或 AI 状态可检查"],
                ["数据目录", "数据库和备份目录创建", "本地数据可持久化"],
                ["模型配置", "设置接口并检查连通", "连接状态明确"],
                ["核心流程", "任务、事件、排程、AI、撤销、备份", "用户流程贯通"],
                ["文档交付", "四份最终 DOCX 在提交目录", "文件完整且可打开"],
            ], [1.2, 3.2, 2.25]),
        ],
        "4.1软件能力": [
            "软件能力还体现在多层一致性上。UI、AI 和 CLI 入口不各自定义规则，而是复用 capability registry；领域服务不依赖前端页面状态，而是直接面对数据库；动作日志不区分操作来自用户还是 AI，只按 actor 和 origin 记录。测试结果表明这种一致性设计已经落实到主要模块。",
            "软件能力也体现在失败处理上。模型失败不会使任务和日历不可用，备份失败不会直接覆盖当前数据库，排程失败会返回原因，撤销失败会说明限制，页面错误会由错误边界接住。对于个人日程产品来说，这些失败处理能力与正常功能同等重要。",
        ],
        "4.4测试结论": [
            TableSpec("表 18 最终验收结论清单", ["验收项", "结论", "说明"], [
                ["功能完整性", "通过", "任务、事件、排程、提醒、总结、撤销、备份和 AI 流程均有实现与测试依据"],
                ["AI agent 真实性", "通过", "模型工具定义、工具调用循环和本地能力执行链路成立"],
                ["数据安全", "通过", "本地 SQLite、动作日志、撤销、备份和敏感信息保护满足要求"],
                ["工程质量", "通过", "测试、类型检查、构建和 Tauri 后端检查通过"],
                ["桌面交付", "通过", "Tauri 桌面应用具备发布基础，建议多机安装复测"],
                ["文档质量", "通过", "四份文档按模板完成并纳入视觉 QA"],
                ["已知限制", "可接受", "E2E 自动化、长期洞察和外部模型稳定性需后续加强"],
                ["最终建议", "准予提交", "满足课程项目高质量交付要求"],
            ], [1.35, 1.2, 4.1]),
            "发布前最后一轮测试应采用“从交付物反推产品”的方式执行。测试人员不再只看源码，而是从最终提交目录、安装包、桌面应用和用户数据开始检查：文档能否直接打开，安装包能否直接运行，应用能否创建数据，AI 状态是否明确，备份是否可创建，撤销是否可恢复。该方式更接近评审和最终用户拿到项目后的真实体验。",
            "若最终检查中某个非核心展示项出现小问题，可以在不影响验收的情况下记录为后续优化；但若出现数据安全、安装启动、AI 工具调用、备份恢复、撤销或文档模板错误等问题，应视为阻塞项重新修复。测试结论的“通过”建立在这些阻塞项均已排除的基础上。",
            TableSpec("表 19 发布前最终检查单", ["检查序号", "检查内容", "通过条件"], [
                ["F-01", "最终提交版目录存在四份 DOCX", "文件名正确，均可用 Word 打开"],
                ["F-02", "四份 DOCX 页数均超过或达到 20 页要求", "Word 统计页数满足要求"],
                ["F-03", "文档目录页码与标题页一致", "抽查目录项可对应正文页"],
                ["F-04", "文档无模板占位符和敏感信息", "扫描未命中占位符、密钥和内部调试文本"],
                ["F-05", "源码基础命令通过", "test、typecheck、build、cargo check 通过"],
                ["F-06", "桌面应用可启动", "主窗口打开且导航可用"],
                ["F-07", "核心用户流程可完成", "任务、事件、排程、AI、撤销、备份贯通"],
                ["F-08", "模型不可用时可降级", "本地任务和日程功能仍可使用"],
                ["F-09", "备份恢复流程安全", "恢复前有保护点，失败不破坏当前数据"],
                ["F-10", "安装或发布产物可分发", "安装包或 release 产物保存在交付位置"],
            ], [0.75, 3.25, 2.6]),
        ],
    },
}


def apply_extra_content(doc_id: str, sections: list[tuple[int, str, list[Any]]]) -> list[tuple[int, str, list[Any]]]:
    extras = EXTRA_CONTENT.get(doc_id, {})
    more_extras = MORE_EXTRA_CONTENT.get(doc_id, {})
    updated: list[tuple[int, str, list[Any]]] = []
    for level, title, items in sections:
        merged = list(items)
        merged.extend(extras.get(title, []))
        merged.extend(more_extras.get(title, []))
        updated.append((level, title, merged))
    return updated


CONTENT_BUILDERS = {
    "01": plan_sections,
    "02": requirements_sections,
    "03": design_sections,
    "04": test_sections,
}


def build_doc(doc_id: str, toc_page_map: dict[str, str]) -> Path:
    spec = DOCS[doc_id]
    template = TEMPLATE_DIR / spec["template"]
    if not template.exists():
        raise FileNotFoundError(template)
    doc = Document(str(template))
    configure_styles(doc)
    clear_template_keep_title(doc, spec["toc"], toc_page_map)
    write_sections(doc, apply_extra_content(doc_id, CONTENT_BUILDERS[doc_id]()))
    ensure_no_update_fields(doc)
    OUT_ASCII_DIR.mkdir(parents=True, exist_ok=True)
    out_path = OUT_ASCII_DIR / spec["ascii_name"]
    doc.save(out_path)
    return out_path


def load_toc_pages(path: Path | None) -> dict[str, dict[str, str]]:
    if not path or not path.exists():
        return {}
    data = json.loads(path.read_text(encoding="utf-8-sig"))
    return {doc_id: {str(k): str(v) for k, v in pages.items()} for doc_id, pages in data.items()}


def mirror_to_course_dirs() -> None:
    final_dir = Path("课程提交文档/最终提交版")
    docx_dir = Path("课程提交文档/docx版")
    final_dir.mkdir(parents=True, exist_ok=True)
    docx_dir.mkdir(parents=True, exist_ok=True)
    for doc_id, spec in DOCS.items():
        source = OUT_ASCII_DIR / spec["ascii_name"]
        if source.exists():
            shutil.copy2(source, final_dir / spec["final_name"])
            shutil.copy2(source, docx_dir / spec["final_name"])


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--toc-pages", type=Path, default=None)
    parser.add_argument("--mirror", action="store_true")
    args = parser.parse_args()
    toc_pages = load_toc_pages(args.toc_pages)
    for doc_id in ("01", "02", "03", "04"):
        build_doc(doc_id, toc_pages.get(doc_id, {}))
    if args.mirror:
        mirror_to_course_dirs()


if __name__ == "__main__":
    main()
