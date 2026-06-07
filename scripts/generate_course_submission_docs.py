from __future__ import annotations

import shutil
import sys
from datetime import datetime
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


PROJECT = "EvolveFlow 智能日程助手"
COURSE = "软件工程"
DATE_TEXT = "2026年6月1日"
VERSION = "V1.0"
LEADER = "2410250732 孟天赐（组长）"
MEMBERS = [
    ("2410250732", "孟天赐", "组长；总体架构、AI 运行时、桌面集成、发布验收"),
    ("2410250115", "杨瑞熙", "需求调研、任务与事件管理需求、测试用例整理"),
    ("2410250575", "王雅娴", "界面原型、交互流程、课程文档排版"),
    ("2410250646", "王研博", "数据库设计、存储层、备份恢复方案"),
    ("2410250848", "姚玮", "日程规划、提醒机制、性能与稳定性测试"),
    ("2410250609", "李思晴", "AI 助手提示词、会话上下文、用户材料"),
    ("2410250311", "冯留杨", "能力注册层、接口校验、撤销与历史记录"),
    ("2410250846", "闫艺馨", "测试执行、缺陷跟踪、交付物整理"),
]

TEST_COUNTS = [
    ("capabilities", 7),
    ("cli", 1),
    ("domain", 18),
    ("storage", 6),
    ("ui-shared", 1),
    ("runtime", 2),
]


def font_path(preferred: Sequence[str]) -> str:
    for name in preferred:
        path = Path("C:/Windows/Fonts") / name
        if path.exists():
            return str(path)
    return ""


FONT_REGULAR = font_path(["msyh.ttc", "simsun.ttc", "simhei.ttf", "arial.ttf"])
FONT_BOLD = font_path(["msyhbd.ttc", "simhei.ttf", "msyh.ttc", "arialbd.ttf"])


def pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    path = FONT_BOLD if bold and FONT_BOLD else FONT_REGULAR
    if path:
        return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    lines: list[str] = []
    for paragraph in text.split("\n"):
        line = ""
        for ch in paragraph:
            candidate = line + ch
            if draw.textlength(candidate, font=font) <= max_width:
                line = candidate
            else:
                if line:
                    lines.append(line)
                line = ch
        if line:
            lines.append(line)
    return lines or [""]


def draw_centered_text(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.FreeTypeFont,
    fill: str = "#1f2937",
    line_gap: int = 8,
) -> None:
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, font, x2 - x1 - 16)
    heights = []
    widths = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        widths.append(bbox[2] - bbox[0])
        heights.append(bbox[3] - bbox[1])
    total_h = sum(heights) + line_gap * (len(lines) - 1)
    y = y1 + ((y2 - y1) - total_h) / 2
    for line, width, height in zip(lines, widths, heights):
        draw.text((x1 + ((x2 - x1) - width) / 2, y), line, font=font, fill=fill)
        y += height + line_gap


def draw_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    body: str = "",
    fill: str = "#ffffff",
    outline: str = "#2563eb",
    title_fill: str = "#111827",
) -> None:
    draw.rounded_rectangle(box, radius=22, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = box
    width = x2 - x1
    title_font = pil_font(28 if width >= 210 else 24, True)
    body_font = pil_font(20 if width >= 260 else 18)
    draw_centered_text(draw, (x1 + 8, y1 + 12, x2 - 8, y1 + 58), title, title_font, title_fill)
    if body:
        draw_centered_text(draw, (x1 + 18, y1 + 70, x2 - 18, y2 - 14), body, body_font, "#374151", 4)


def draw_entity_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    fields: Sequence[str],
    fill: str,
    outline: str = "#475569",
) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=3)
    title_font = pil_font(27, True)
    field_font = pil_font(20)
    draw_centered_text(draw, (x1 + 8, y1 + 12, x2 - 8, y1 + 55), title, title_font, "#111827", 2)
    y = y1 + 78
    for field in fields:
        draw.text((x1 + 24, y), field, font=field_font, fill="#334155")
        y += 30


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#4b5563") -> None:
    draw.line([start, end], fill=color, width=4)
    sx, sy = start
    ex, ey = end
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex >= sx else -1
        pts = [(ex, ey), (ex - direction * 18, ey - 10), (ex - direction * 18, ey + 10)]
    else:
        direction = 1 if ey >= sy else -1
        pts = [(ex, ey), (ex - 10, ey - direction * 18), (ex + 10, ey - direction * 18)]
    draw.polygon(pts, fill=color)


def save_architecture(path: Path) -> None:
    img = Image.new("RGB", (1700, 1050), "#f8fafc")
    draw = ImageDraw.Draw(img)
    title = pil_font(42, True)
    draw.text((60, 40), "EvolveFlow 总体架构图", font=title, fill="#0f172a")
    layers = [
        ("用户界面层", "Tauri v2 桌面端 / React + Vite / CLI\n今日、任务、日历、AI、统计、设置页面", "#dbeafe", "#2563eb"),
        ("Sidecar 与 AI Agent 层", "JSON-RPC 通道 / 会话上下文 / 工具调用循环\nAnthropic 兼容接口、Dream 记忆、Buddy 反馈", "#dcfce7", "#16a34a"),
        ("能力注册层", "统一 capability registry\n权限、输入校验、幂等、修订号、动作日志", "#fef3c7", "#d97706"),
        ("领域服务层", "Task / Event / Schedule / Reminder / Undo\nSummary / Preference / Memory Projection", "#fce7f3", "#db2777"),
        ("本地存储层", "SQLite + better-sqlite3\n迁移、索引、WAL、备份、恢复、导出", "#ede9fe", "#7c3aed"),
    ]
    y = 125
    for title_text, body, fill, outline in layers:
        draw_box(draw, (240, y, 1390, y + 135), title_text, body, fill, outline)
        if y < 700:
            draw_arrow(draw, (815, y + 135), (815, y + 170))
        y += 170
    draw_box(draw, (60, 320, 200, 505), "用户", "自然语言\n键盘鼠标", "#ffffff", "#64748b")
    draw_arrow(draw, (200, 410), (240, 190), "#64748b")
    draw_box(draw, (1435, 290, 1645, 505), "AI API", "DeepSeek / Anthropic\n兼容 Messages API", "#ffffff", "#64748b")
    draw_arrow(draw, (1390, 360), (1435, 360), "#64748b")
    draw_box(draw, (1435, 720, 1645, 905), "系统", "通知\n文件系统", "#ffffff", "#64748b")
    draw_arrow(draw, (1390, 795), (1435, 795), "#64748b")
    img.save(path, quality=95)


def save_use_case(path: Path) -> None:
    img = Image.new("RGB", (1600, 1000), "#ffffff")
    draw = ImageDraw.Draw(img)
    draw.text((70, 45), "需求用例图", font=pil_font(42, True), fill="#0f172a")
    actor_font = pil_font(28, True)
    draw_box(draw, (80, 230, 250, 390), "用户", "学生 / 个人\n日程管理者", "#f8fafc", "#475569")
    draw_box(draw, (1320, 230, 1510, 390), "AI Agent", "理解意图\n调用工具", "#f8fafc", "#475569")
    draw_box(draw, (1320, 600, 1510, 760), "系统提醒", "到点提醒\n备份任务", "#f8fafc", "#475569")
    cases = [
        ((420, 145, 725, 265), "管理任务", "创建、编辑、完成\n延期、锁定"),
        ((850, 145, 1155, 265), "管理事件", "日历事件\n冲突检查、锁定"),
        ((420, 320, 725, 440), "自动规划", "按日期生成时间块\n并解释原因"),
        ((850, 320, 1155, 440), "AI 对话", "自然语言拆解需求\n并调用能力"),
        ((420, 495, 725, 615), "提醒处理", "创建提醒、稍后提醒\n解除提醒"),
        ((850, 495, 1155, 615), "撤销与历史", "查看动作日志\n并回滚操作"),
        ((420, 670, 725, 790), "备份恢复", "创建、校验、恢复\n删除备份"),
        ((850, 670, 1155, 790), "统计与总结", "日总结、质量分析\nDream 洞察"),
    ]
    for box, title, body in cases:
        draw_box(draw, box, title, body, "#eef6ff", "#2563eb")
        draw_arrow(draw, (250, 310), (box[0], (box[1] + box[3]) // 2), "#64748b")
    for box, _, _ in cases[2:]:
        if box[0] > 800:
            draw_arrow(draw, (1320, 310), (box[2], (box[1] + box[3]) // 2), "#16a34a")
    draw_arrow(draw, (1320, 685), (1130, 555), "#16a34a")
    draw.text((70, 900), "说明：AI Agent 不是规则替代层，而是通过真实模型理解用户意图，再由受控 capability 调用修改本地数据。", font=pil_font(24), fill="#334155")
    img.save(path, quality=95)


def save_er(path: Path) -> None:
    img = Image.new("RGB", (1800, 1250), "#f8fafc")
    draw = ImageDraw.Draw(img)
    draw.text((70, 45), "数据库 ER / 表结构关系图", font=pil_font(42, True), fill="#0f172a")
    boxes = {
        "tasks": ((80, 165, 420, 355), "tasks", ["id, title, status", "duration, due_date", "parent_task_id", "project, locked"]),
        "events": ((520, 165, 860, 355), "events", ["id, title", "start_time, end_time", "location", "bound_task_id, locked"]),
        "ai": ((960, 165, 1300, 355), "ai_sessions / messages", ["session_id, role", "content, tool_calls", "created_at"]),
        "prefs": ((1400, 165, 1740, 355), "preferences / signals", ["key, value, scope", "source", "learned signals"]),
        "schedule": ((300, 495, 680, 705), "schedule_blocks", ["date, start_time", "end_time, source", "task_id, event_id", "locked"]),
        "reminders": ((780, 495, 1160, 705), "reminders", ["trigger_at, status", "snooze_until", "task_id, event_id"]),
        "actions": ((1260, 495, 1640, 705), "action_logs / undo", ["capability, actor", "input, result", "state_before", "group_id"]),
        "extensions": ((80, 845, 420, 1035), "task extensions", ["recurrence_rules", "relations", "tags", "task_reminders"]),
        "event_ext": ((520, 845, 860, 1035), "event extensions", ["event_recurrence_rules", "event_reminders"]),
        "dream": ((960, 845, 1300, 1035), "daily / dream", ["daily_summaries", "dream_insights", "confidence, evidence"]),
        "meta": ((1400, 845, 1740, 1035), "app_meta / backup", ["schema_version", "migrations", "backup files"]),
    }
    colors = ["#dbeafe", "#dcfce7", "#e0f2fe", "#fce7f3", "#fef3c7", "#fee2e2", "#ede9fe", "#f1f5f9", "#f1f5f9", "#ecfccb", "#ffffff"]
    for (name, (box, title_text, fields)), fill in zip(boxes.items(), colors):
        draw_entity_box(draw, box, title_text, fields, fill)
    relations = [
        ("tasks", "schedule"),
        ("events", "schedule"),
        ("tasks", "reminders"),
        ("events", "reminders"),
        ("ai", "actions"),
        ("actions", "tasks"),
        ("prefs", "dream"),
        ("tasks", "extensions"),
        ("events", "event_ext"),
    ]
    centers = {k: ((v[0][0] + v[0][2]) // 2, (v[0][1] + v[0][3]) // 2) for k, v in boxes.items()}
    for a, b in relations:
        ax, ay = centers[a]
        bx, by = centers[b]
        draw_arrow(draw, (ax, ay), (bx, by), "#94a3b8")
    relation_text = [
        "tasks/events → schedule_blocks：排程结果引用任务或事件",
        "tasks/events → reminders：统一提醒触发与稍后提醒",
        "ai_sessions/messages → action_logs：AI 工具调用可审计",
        "preferences/signals → dream_insights：偏好沉淀为长期洞察",
    ]
    y = 1095
    for text in relation_text:
        draw.text((90, y), f"• {text}", font=pil_font(22), fill="#334155")
        y += 34
    draw.text((900, 1130), "核心约束：外部 AI 不能直接改库，必须经过 capability、领域服务、动作日志和备份保护。", font=pil_font(24), fill="#0f172a")
    img.save(path, quality=95)


def save_sequence(path: Path) -> None:
    img = Image.new("RGB", (1800, 1050), "#ffffff")
    draw = ImageDraw.Draw(img)
    draw.text((70, 45), "AI 日程规划动态模型（时序图）", font=pil_font(42, True), fill="#0f172a")
    participants = ["用户", "桌面 UI", "Sidecar", "AI 模型", "能力注册层", "领域服务", "SQLite"]
    xs = [130, 385, 640, 895, 1150, 1405, 1660]
    top = 150
    for x, name in zip(xs, participants):
        draw_box(draw, (x - 85, top, x + 85, top + 70), name, "", "#f8fafc", "#2563eb")
        draw.line((x, top + 70, x, 930), fill="#cbd5e1", width=3)
    steps = [
        (0, 1, "输入：帮我安排今天"),
        (1, 2, "JSON-RPC：ai.chat / schedule.plan_day"),
        (2, 3, "携带上下文调用真实模型"),
        (3, 2, "返回 tool_use 与说明"),
        (2, 4, "校验 capability 与输入"),
        (4, 5, "调用 ScheduleService"),
        (5, 6, "读取任务、事件、偏好并写入计划块"),
        (6, 5, "返回计划块"),
        (5, 4, "结果与修订号"),
        (4, 2, "动作日志与响应"),
        (2, 1, "自然语言解释 + 结构化数据"),
        (1, 0, "刷新今日视图"),
    ]
    y = 265
    for frm, to, label in steps:
        draw_arrow(draw, (xs[frm], y), (xs[to], y), "#475569")
        draw.text((min(xs[frm], xs[to]) + 12, y - 34), label, font=pil_font(20), fill="#334155")
        y += 58
    draw.rounded_rectangle((70, 950, 1730, 1010), radius=16, fill="#ecfeff", outline="#0891b2", width=2)
    draw_centered_text(draw, (85, 950, 1715, 1010), "关键约束：AI 负责理解与决策建议，真正的数据修改必须经过 capability 白名单、输入校验、领域服务和动作日志。", pil_font(24), "#155e75", 4)
    img.save(path, quality=95)


def save_test_summary(path: Path) -> None:
    img = Image.new("RGB", (1600, 950), "#f8fafc")
    draw = ImageDraw.Draw(img)
    draw.text((70, 45), "测试结果概览（2026-06-01）", font=pil_font(42, True), fill="#0f172a")
    total = sum(count for _, count in TEST_COUNTS)
    draw_box(draw, (80, 130, 480, 290), "自动化测试", f"{total} 个测试用例全部通过\nnpm test", "#dcfce7", "#16a34a")
    draw_box(draw, (600, 130, 1000, 290), "类型检查", "7 个工作空间通过\nnpm run typecheck", "#dbeafe", "#2563eb")
    draw_box(draw, (1120, 130, 1520, 290), "生产构建", "前端与各包构建通过\nnpm run build", "#fef3c7", "#d97706")
    max_count = max(count for _, count in TEST_COUNTS)
    chart_x, chart_y = 160, 390
    bar_w, gap = 160, 55
    for i, (name, count) in enumerate(TEST_COUNTS):
        x = chart_x + i * (bar_w + gap)
        h = int(360 * count / max_count)
        draw.rounded_rectangle((x, chart_y + 360 - h, x + bar_w, chart_y + 360), radius=12, fill="#2563eb")
        draw.text((x + 20, chart_y + 375), name, font=pil_font(22, True), fill="#1f2937")
        draw.text((x + 55, chart_y + 315 - h), str(count), font=pil_font(30, True), fill="#0f172a")
    draw.line((120, chart_y + 360, 1500, chart_y + 360), fill="#94a3b8", width=3)
    draw.text((80, 830), "交付物：MSI 安装包约 41.10 MB，NSIS 安装程序约 27.44 MB；AI 能力通过 sidecar + Anthropic 兼容接口接入。", font=pil_font(24), fill="#334155")
    img.save(path, quality=95)


def save_ui_flow(path: Path) -> None:
    img = Image.new("RGB", (1600, 1000), "#ffffff")
    draw = ImageDraw.Draw(img)
    draw.text((70, 45), "桌面端页面与交互流转图", font=pil_font(42, True), fill="#0f172a")
    draw_box(draw, (640, 120, 960, 235), "应用外壳", "侧边导航、全局 AI 浮窗、错误边界", "#eef2ff", "#4f46e5")
    pages = [
        ((120, 350, 395, 470), "今日", "快速添加、建议、时间块"),
        ((500, 350, 775, 470), "任务", "筛选、详情、编辑、完成"),
        ((880, 350, 1155, 470), "日历", "事件管理、冲突检查"),
        ((1260, 350, 1535, 470), "AI", "会话、上下文、工具调用"),
        ((310, 650, 585, 770), "统计", "日总结、质量分析、趋势"),
        ((690, 650, 965, 770), "设置", "模型配置、备份、偏好"),
        ((1070, 650, 1345, 770), "弹窗组件", "任务/事件编辑、提醒、备份"),
    ]
    for box, title, body in pages:
        draw_box(draw, box, title, body, "#f8fafc", "#0ea5e9")
        draw_arrow(draw, (800, 235), ((box[0] + box[2]) // 2, box[1]), "#64748b")
    draw.text((120, 875), "交互原则：常用任务在今日页完成，复杂编辑通过模态框承载；AI 功能显示真实状态，不使用伪造建议替代模型结果。", font=pil_font(24), fill="#334155")
    img.save(path, quality=95)


def set_cell_bg(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, bold=bold)


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def setup_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.2)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    for i in range(1, 4):
        style = styles[f"Heading {i}"]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.bold = True
        style.font.color.rgb = RGBColor(15, 23, 42)
        style.font.size = Pt({1: 16, 2: 14, 3: 12}[i])
        style.paragraph_format.space_before = Pt(10 if i == 1 else 6)
        style.paragraph_format.space_after = Pt(4)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"{PROJECT} 课程提交文档")
    set_run_font(run, size=9, color="64748B")


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text)
    set_run_font(run, size=22, bold=True, color="0F172A")


def add_cover(doc: Document, title: str, doc_type: str) -> None:
    add_title(doc, title)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(PROJECT)
    set_run_font(r, size=16, bold=True, color="2563EB")
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("课程项目提交材料")
    set_run_font(r2, size=13, color="334155")
    table = doc.add_table(rows=7, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    rows = [
        ("课程名称", COURSE),
        ("项目名称", PROJECT),
        ("文档类别", doc_type),
        ("版本日期", f"{VERSION} / {DATE_TEXT}"),
        ("组长", LEADER),
        ("小组成员", "；".join(f"{sid} {name}" for sid, name, _ in MEMBERS[1:])),
        ("适用范围", "课程验收、项目评审、开发与测试追踪"),
    ]
    for row, (k, v) in zip(table.rows, rows):
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(12)
        set_cell_text(row.cells[0], k, bold=True)
        set_cell_text(row.cells[1], v)
        set_cell_bg(row.cells[0], "E2E8F0")
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("说明：本文档根据当前项目代码、构建结果和课程模板整理形成。")
    set_run_font(run, size=10, color="64748B")
    doc.add_page_break()


def add_toc(doc: Document, sections: Sequence[tuple[str, str]]) -> None:
    doc.add_heading("目录", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_text(table.rows[0].cells[0], "章节", bold=True)
    set_cell_text(table.rows[0].cells[1], "标题", bold=True)
    set_cell_bg(table.rows[0].cells[0], "E2E8F0")
    set_cell_bg(table.rows[0].cells[1], "E2E8F0")
    for no, title in sections:
        row = table.add_row()
        set_cell_text(row.cells[0], no)
        set_cell_text(row.cells[1], title)
    doc.add_page_break()


def para(doc: Document, text: str, first_line: bool = True) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(4)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run)


def bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.32)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run("（1）" if False else f"• {item}")
        set_run_font(run)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
        set_cell_bg(hdr[i], "DBEAFE")
        if widths:
            hdr[i].width = Cm(widths[i])
    for row_data in rows:
        row = table.add_row().cells
        for i, text in enumerate(row_data):
            set_cell_text(row[i], text)
            if widths:
                row[i].width = Cm(widths[i])
    doc.add_paragraph()


def add_picture(doc: Document, path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.4))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(caption)
    set_run_font(r, size=9, color="475569")


def build_common_intro(doc: Document, target: str) -> None:
    doc.add_heading("1．引言", level=1)
    doc.add_heading("1.1 编写目的", level=2)
    para(doc, f"本文档用于说明 {PROJECT} 在{target}方面的目标、范围、约束和交付要求。文档面向课程教师、项目小组成员、后续维护人员和验收人员，使需求、设计、实现和测试之间可以互相追踪。")
    doc.add_heading("1.2 项目背景", level=2)
    para(doc, "本项目来源于软件工程课程实践。日常学习和团队协作中，任务、课程安排、临时事项和复习计划常分散在不同工具中，用户需要频繁手工拆分任务、估算时长、处理冲突并调整日程。EvolveFlow 旨在以本地优先的桌面应用承载个人日程数据，并通过真实 AI Agent 理解用户自然语言意图，调用受控工具完成任务创建、日程规划、提醒和总结。")
    para(doc, "项目委托单位可视为软件工程课程实践教学环节，开发单位为 EvolveFlow 项目小组，主管和验收主体为课程任课教师及相关教学组织。系统与操作系统、SQLite 本地数据库和 Anthropic 兼容 AI 服务协同工作；没有网络或未配置模型时，基础任务、日程和备份功能仍可使用。")
    doc.add_heading("1.3 定义", level=2)
    add_table(
        doc,
        ["术语", "含义"],
        [
            ("AI Agent", "由真实大模型驱动的智能体，负责理解用户意图、决定是否调用工具，并将结果解释给用户。"),
            ("Sidecar", "Tauri 桌面端启动的 Node.js 运行时进程，通过 JSON-RPC 与前端通信并承载 AI 工具调用。"),
            ("Capability", "系统内部统一能力接口，例如 task.create、schedule.plan_day、ai.chat、backup.restore 等。"),
            ("Schedule Block", "计划块，表示某一日期内任务或事件占用的时间段，是自动排程的核心输出。"),
            ("Local-first", "数据优先保存在本机 SQLite 数据库中，保证隐私、离线可用性和可备份性。"),
            ("Dream System", "AI 记忆/洞察模块，用于沉淀用户偏好、行为规律和长期优化建议。"),
        ],
        [3.2, 12.2],
    )
    doc.add_heading("1.4 参考资料", level=2)
    bullets(
        doc,
        [
            "课程提供的软件工程文档模板：项目开发计划、需求规格说明书、概要设计说明书、测试分析报告。",
            "项目源码：README、package.json、Tauri 配置、存储层 database.ts、能力层 capabilities.ts、领域服务与运行时源码。",
            "测试执行记录：npm run typecheck、npm test、npm run build 均已在 2026年6月1日通过。",
            "Tauri v2、React、TypeScript、SQLite、Vitest 等技术官方文档与项目依赖说明。",
        ],
    )


def build_plan(path: Path, images: dict[str, Path]) -> None:
    doc = Document()
    setup_doc(doc)
    add_cover(doc, "二、项目开发计划", "项目开发计划")
    add_toc(doc, [("1", "引言"), ("2", "项目概述"), ("3", "实施计划"), ("4", "人员组织及分工"), ("5", "交付期限"), ("6", "专题计划要点")])
    build_common_intro(doc, "项目开发计划")
    doc.add_heading("2．项目概述", level=1)
    doc.add_heading("2.1 工作内容", level=2)
    para(doc, "本项目建设一个可安装、可实际使用的智能日程助手桌面产品。核心工作包括任务管理、事件管理、自动日程规划、提醒通知、历史撤销、备份恢复、AI 对话与工具调用、Dream 洞察、桌面端打包发布和课程文档整理。")
    add_table(
        doc,
        ["工作项", "主要内容", "完成标准"],
        [
            ("任务管理", "创建、编辑、完成、延期、删除、锁定、项目与标签管理。", "用户可以在桌面端和能力接口中稳定管理任务。"),
            ("事件与日历", "创建固定时间事件、检查冲突、锁定不可移动事件。", "日历页与计划服务能统一识别时间占用。"),
            ("智能排程", "按日期生成任务/事件时间块，尊重锁定项、截止日期和时长估计。", "schedule.plan_day、schedule.rebalance 等能力可被 UI 和 AI 调用。"),
            ("AI Agent", "通过 sidecar 接入 Anthropic 兼容模型，支持 ai.chat、上下文读取和工具调用。", "不使用伪造建议替代 AI；模型不可用时明确提示状态。"),
            ("数据与安全", "SQLite 本地存储、迁移、索引、动作日志、撤销、备份与恢复。", "数据可追踪、可恢复、可导出，关键操作有日志。"),
            ("发布交付", "生成 Windows MSI 与 NSIS 安装包，文档和测试报告齐备。", "安装包可分发，正式公开发布前补充签名证书。"),
        ],
        [2.6, 8.0, 4.2],
    )
    add_picture(doc, images["architecture"], "图 2-1 EvolveFlow 总体架构")
    doc.add_heading("2.2 条件与限制", level=2)
    bullets(
        doc,
        [
            "开发条件：Node.js、npm、TypeScript、Rust/Tauri、React/Vite、Vitest 与 Windows 开发环境已经具备。",
            "运行条件：面向 Windows 10/11 x64；安装包内包含运行所需资源，用户无需理解内部工作空间结构。",
            "AI 条件：AI Agent 依赖可访问的 Anthropic 兼容接口及有效模型配置；没有模型配置时不得展示伪造 AI 结论。",
            "数据限制：当前版本为本地优先单机应用，不承诺跨设备云同步；用户需通过备份功能保护重要数据。",
            "发布限制：课程交付可直接安装运行；若公开发布到互联网，建议补充代码签名证书、自动更新签名密钥和更完整的端到端测试矩阵。",
        ],
    )
    doc.add_heading("2.3 产品", level=2)
    doc.add_heading("2.3.1 程序", level=3)
    add_table(
        doc,
        ["交付程序", "位置/形式", "说明"],
        [
            ("桌面安装包 MSI", "apps/desktop-tauri/src-tauri/target/release/bundle/msi/EvolveFlow_0.1.0_x64_en-US.msi", "约 41.10 MB，适合 Windows 安装部署。"),
            ("桌面安装包 NSIS", "apps/desktop-tauri/src-tauri/target/release/bundle/nsis/EvolveFlow_0.1.0_x64-setup.exe", "约 27.44 MB，适合普通用户双击安装。"),
            ("源码工作空间", "packages、runtime、apps/desktop-tauri", "包含存储、领域、能力、运行时、CLI 与桌面端代码。"),
        ],
        [3.0, 8.5, 4.0],
    )
    doc.add_heading("2.3.2 文档", level=3)
    add_table(
        doc,
        ["文档", "用途"],
        [
            ("项目开发计划", "说明计划、组织、资源、进度、风险和验收标准。"),
            ("需求规格说明书", "说明执行者、功能需求、非功能需求和故障处理。"),
            ("概要设计说明书", "说明系统架构、对象模型、动态模型、数据结构和界面设计。"),
            ("测试分析报告", "说明测试执行情况、结果、需求覆盖和项目评价。"),
        ],
        [4.0, 11.0],
    )
    doc.add_heading("2.4 运行环境", level=2)
    add_table(
        doc,
        ["类别", "要求"],
        [
            ("用户端", "Windows 10/11 x64，建议 4GB 以上内存；首次 AI 调用需要网络。"),
            ("桌面框架", "Tauri v2 + WebView，前端使用 React + Vite 构建。"),
            ("运行时", "发布包内置 Node sidecar、runtime/dist 与必要依赖。"),
            ("数据库", "SQLite，本地文件存储，启用迁移、索引与备份机制。"),
            ("AI 服务", "Anthropic 兼容 Messages API，可通过环境变量或设置页配置。"),
        ],
        [3.2, 12.0],
    )
    doc.add_heading("2.5 服务", level=2)
    para(doc, "项目小组提供安装说明、基本使用说明、测试报告、备份恢复说明和后续维护建议。对于课程验收，交付重点为可安装程序、核心功能演示、源码可构建性和文档完整性。")
    doc.add_heading("2.6 验收标准", level=2)
    add_table(
        doc,
        ["验收项", "标准", "当前结果"],
        [
            ("类型检查", "所有 TypeScript 工作空间 tsc --noEmit 通过。", "已通过 npm run typecheck。"),
            ("自动化测试", "存储、领域、能力、运行时等核心测试通过。", "35 个自动化用例通过。"),
            ("生产构建", "所有包和桌面前端可完成生产构建。", "已通过 npm run build。"),
            ("安装包", "生成 Windows 安装包并包含运行资源。", "已生成 MSI 与 NSIS 安装包。"),
            ("AI 真实性", "AI 功能由真实模型接口驱动，不以规则或固定文案伪造。", "sidecar 支持 Anthropic 兼容接口和真实 tool-use。"),
            ("文档完整性", "四份课程提交文档结构完整、无模板占位。", "本次交付已补写并插入图表。"),
        ],
        [3.2, 7.2, 4.2],
    )
    doc.add_heading("3．实施计划", level=1)
    doc.add_heading("3.1 任务分解", level=2)
    add_table(
        doc,
        ["阶段", "任务", "主要交付"],
        [
            ("阶段一：调研与需求", "梳理课程提交要求、同类日程工具功能、用户使用场景。", "需求列表、用例、风险假设。"),
            ("阶段二：架构与数据设计", "确定本地优先架构、能力层、领域服务和 SQLite 表结构。", "架构图、ER 图、接口清单。"),
            ("阶段三：核心实现", "实现任务/事件/排程/提醒/撤销/备份等基础能力。", "可运行的 CLI 与桌面端基础功能。"),
            ("阶段四：AI 集成", "实现 sidecar、AI 会话、真实模型调用、工具调用和 Dream 洞察。", "AI Agent 可通过受控能力修改数据。"),
            ("阶段五：桌面体验与发布", "完善页面、错误提示、资源打包、安装包生成和人工体验验证。", "MSI、NSIS、测试记录。"),
            ("阶段六：课程文档", "整理计划、需求、设计、测试四份文档并插入图表。", "最终 DOCX 交付材料。"),
        ],
        [3.4, 8.0, 4.0],
    )
    doc.add_heading("3.2 进度", level=2)
    add_table(
        doc,
        ["里程碑", "计划内容", "状态"],
        [
            ("M1", "完成项目立项、需求范围和小组分工。", "完成"),
            ("M2", "完成数据模型、能力接口和领域服务骨架。", "完成"),
            ("M3", "完成核心业务功能和自动化测试。", "完成"),
            ("M4", "完成 AI sidecar、真实模型配置和桌面端联调。", "完成"),
            ("M5", "完成 Windows 安装包、测试报告和课程文档。", "完成/提交前复核"),
        ],
        [2.5, 10.0, 3.0],
    )
    doc.add_heading("3.3 预算", level=2)
    para(doc, "本课程项目以学习实践为主，不设置商业采购预算。主要成本为小组成员投入的人时、AI 接口少量测试额度、开发电脑资源和课程文档整理时间。若进入正式产品化阶段，预算应补充代码签名证书、云同步服务、自动更新发布服务和长期模型调用费用。")
    doc.add_heading("3.4 关键问题", level=2)
    add_table(
        doc,
        ["关键问题", "影响", "应对措施"],
        [
            ("AI 模型不可用或网络失败", "AI 对话和智能建议受影响。", "界面明确展示状态；基础本地功能不依赖模型。"),
            ("数据误操作", "任务和日程可能被错误修改。", "动作日志、撤销、备份恢复和 capability 输入校验。"),
            ("安装包缺少运行资源", "用户安装后 sidecar 无法启动。", "发布包内置 Node、runtime/dist 和依赖资源。"),
            ("排程结果不符合预期", "用户信任下降。", "提供锁定、解释、重新规划和手工调整入口。"),
            ("课程文档与代码脱节", "验收材料可信度降低。", "文档以当前源码、测试结果和交付物为依据。"),
        ],
        [4.0, 5.0, 6.0],
    )
    doc.add_heading("4．人员组织及分工", level=1)
    add_table(doc, ["学号", "姓名", "职责"], MEMBERS, [3.0, 2.6, 10.0])
    doc.add_heading("5．交付期限", level=1)
    para(doc, f"课程提交版本计划于 {DATE_TEXT} 完成。交付内容包括四份 DOCX 文档、项目源码、测试结果和 Windows 安装包。后续若课程要求演示，可基于当前安装包进行功能展示。")
    doc.add_heading("6．专题计划要点", level=1)
    bullets(
        doc,
        [
            "质量计划：以类型检查、自动化测试、生产构建和人工体验验证共同构成验收闭环。",
            "配置计划：模型密钥不写入文档；支持通过环境变量或设置页配置 AI 服务。",
            "测试计划：覆盖存储初始化、领域服务、能力注册、AI 客户端兼容性和桌面冒烟流程。",
            "发布计划：课程版本提供本地安装包；公开发布需补充签名、更新通道和跨版本迁移测试。",
        ],
    )
    doc.save(path)


def build_srs(path: Path, images: dict[str, Path]) -> None:
    doc = Document()
    setup_doc(doc)
    add_cover(doc, "三、需求规格说明书", "需求规格说明书")
    add_toc(doc, [("1", "引言"), ("2", "需求概述"), ("3", "功能需求"), ("4", "非功能需求"), ("5", "故障处理"), ("6", "其它需求")])
    build_common_intro(doc, "需求规格说明")
    doc.add_heading("2．需求概述", level=1)
    doc.add_heading("2.1 目标", level=2)
    para(doc, "系统目标是为个人学习、课程任务和日常事务提供一个本地优先、可被 AI 协助的日程管理工具。用户可以用自然语言或图形界面录入任务与事件，系统自动形成当天计划、提醒关键事项，并在需要时解释排程依据。")
    para(doc, "需求层面的关键目标包括：降低手动规划成本、减少任务遗忘、保证数据可控、支持撤销和备份、使 AI 操作透明可追踪，并提供可安装的桌面成品。")
    doc.add_heading("2.2 运行环境", level=2)
    add_table(
        doc,
        ["环境项", "需求"],
        [
            ("操作系统", "Windows 10/11 x64。"),
            ("硬件", "普通学习/办公电脑即可；建议 4GB 以上内存和稳定磁盘空间。"),
            ("本地依赖", "安装包应携带运行时资源；开发环境需要 Node.js、npm、Rust/Tauri。"),
            ("数据库", "SQLite 本地数据库，不要求额外数据库服务器。"),
            ("网络", "基础功能可离线；AI 会话、模型连通性测试和 Dream 分析需要可访问模型服务。"),
        ],
        [3.2, 12.0],
    )
    doc.add_heading("2.3 条件与限制", level=2)
    bullets(
        doc,
        [
            "系统面向单用户本地数据场景，暂不要求多人协作、云端同步和移动端客户端。",
            "AI 输出可能受模型服务质量影响，所有实际数据修改必须通过 capability 层校验并记录。",
            "用户可手工调整计划并锁定重要任务或事件，AI 和自动排程不得覆盖锁定项。",
            "涉及隐私的数据默认保存在本机，课程文档不记录任何真实 API 密钥。",
        ],
    )
    doc.add_heading("3．功能需求", level=1)
    doc.add_heading("3.1 确定执行者", level=2)
    add_table(
        doc,
        ["执行者", "说明", "主要交互"],
        [
            ("普通用户", "使用桌面端管理学习任务、课程事件和日常安排。", "任务/日历/今日/AI/设置页面。"),
            ("AI Agent", "真实模型驱动的助手，理解自然语言并调用受控能力。", "ai.chat、tool_use、schedule.plan_day 等。"),
            ("系统调度器", "自动规划、提醒轮询、日总结和备份相关的系统过程。", "计划块生成、提醒状态更新、备份校验。"),
            ("维护人员", "负责构建、测试、发布和故障排查。", "源码、日志、测试命令、安装包。"),
        ],
        [3.0, 6.0, 6.2],
    )
    doc.add_heading("3.2 确定用例", level=2)
    add_picture(doc, images["use_case"], "图 3-1 需求用例图")
    add_table(
        doc,
        ["编号", "用例", "优先级", "摘要"],
        [
            ("UC-01", "创建与维护任务", "高", "用户创建任务，补充时长、截止日期、项目、标签并可完成或延期。"),
            ("UC-02", "创建与维护事件", "高", "用户录入固定时间事件，系统用于冲突检查和排程避让。"),
            ("UC-03", "自动规划一天", "高", "系统读取任务、事件和偏好，生成当天计划块。"),
            ("UC-04", "与 AI 助手对话", "高", "用户用自然语言提出需求，AI 调用工具并返回解释。"),
            ("UC-05", "提醒与稍后提醒", "中", "系统在触发时间展示提醒，用户可稍后或解除。"),
            ("UC-06", "撤销与历史", "高", "用户查看动作日志，对误操作进行回滚。"),
            ("UC-07", "备份与恢复", "高", "用户创建、验证、恢复或删除本地备份。"),
            ("UC-08", "统计、总结与 Dream 洞察", "中", "系统生成日总结、分析排程质量并保留长期洞察。"),
            ("UC-09", "配置模型与偏好", "中", "用户配置 AI 服务、工作时间和通知等偏好。"),
        ],
        [2.0, 4.0, 2.0, 8.0],
    )
    doc.add_heading("3.3 编写用例文档", level=2)
    add_table(
        doc,
        ["用例", "前置条件", "基本流程", "异常/扩展", "后置条件"],
        [
            ("UC-01 创建任务", "用户已打开桌面端。", "输入标题、时长、日期等信息；系统校验后写入 tasks。", "标题为空时拒绝；可补充项目和标签。", "任务出现在任务列表和今日候选中。"),
            ("UC-03 自动规划", "存在待办任务或事件。", "用户点击自动规划；系统生成 schedule_blocks 并展示解释。", "锁定项不移动；时间不足时给出未排入说明。", "今日页展示新的计划块。"),
            ("UC-04 AI 对话", "已配置可用模型。", "用户输入自然语言；sidecar 调用模型；模型请求工具；capability 执行。", "模型不可用时显示离线/未配置；工具输入非法时返回错误。", "用户获得回复，数据变更被记录。"),
            ("UC-06 撤销操作", "存在可撤销动作日志。", "用户选择历史操作并确认撤销；UndoService 恢复 state_before。", "动作缺少可恢复状态时提示不可撤销。", "数据恢复到指定动作前状态。"),
            ("UC-07 备份恢复", "存在数据库文件。", "用户创建备份；恢复前系统校验备份完整性。", "备份损坏时拒绝恢复并提示原因。", "生成备份或恢复成功。"),
        ],
        [2.2, 3.2, 5.1, 4.0, 3.6],
    )
    add_table(
        doc,
        ["需求编号", "功能需求", "验收方式"],
        [
            ("FR-01", "系统应支持任务创建、编辑、完成、延期、锁定、取消和删除。", "任务服务测试、桌面任务页人工验证。"),
            ("FR-02", "系统应支持事件创建、编辑、删除、冲突检查和锁定。", "事件服务与日历页验证。"),
            ("FR-03", "系统应按日期生成日程计划块，并能解释计划来源。", "ScheduleService 测试和 Today 页面验证。"),
            ("FR-04", "系统应提供真实 AI Agent 对话能力，允许模型通过工具调用受控修改数据。", "runtime AI 测试和 sidecar 联调。"),
            ("FR-05", "系统应记录动作日志并支持撤销可恢复操作。", "UndoService 与 action_logs 验证。"),
            ("FR-06", "系统应支持提醒列表、稍后提醒和解除提醒。", "ReminderService 测试。"),
            ("FR-07", "系统应支持本地备份、校验、恢复和删除。", "backup capability 与人工流程验证。"),
            ("FR-08", "系统应支持日总结、排程质量分析和 Dream 洞察。", "summary、dream capability 与 UI 面板验证。"),
            ("FR-09", "系统应提供设置页显示 AI 配置状态，不泄露密钥。", "设置页和 api_key.status 验证。"),
        ],
        [2.2, 8.2, 5.0],
    )
    doc.add_heading("4．非功能需求", level=1)
    doc.add_heading("4.1 性能需求", level=2)
    add_table(
        doc,
        ["指标", "需求说明"],
        [
            ("启动与响应", "常规页面操作应在用户可感知的短时间内完成；数据库查询使用索引支撑。"),
            ("排程效率", "单日任务和事件规模达到学习/办公常见数量时，计划生成应保持流畅。"),
            ("资源占用", "桌面端应避免长期高 CPU 占用；提醒轮询和 AI 连通性检查需有缓存或节流。"),
            ("构建可靠性", "typecheck、test、build 应作为提交前基本门槛。"),
        ],
        [3.0, 12.0],
    )
    doc.add_heading("4.2 安全需求", level=2)
    add_table(
        doc,
        ["安全点", "需求说明"],
        [
            ("本地数据保护", "用户数据默认保存在本地 SQLite 文件中，支持备份恢复。"),
            ("密钥保护", "文档和日志不得明文公开 API Key；状态接口只返回是否配置、来源和模型等非敏感信息。"),
            ("AI 操作边界", "AI 不能直接写数据库，必须通过 capability 白名单和输入校验。"),
            ("可追溯性", "所有变更能力应记录 actor、origin、input、result 和 state_before。"),
            ("恢复能力", "误操作可撤销，重大变更可通过备份恢复降低损失。"),
        ],
        [3.0, 12.0],
    )
    doc.add_heading("5．故障处理", level=1)
    add_table(
        doc,
        ["故障场景", "系统处理"],
        [
            ("AI 未配置", "AI 页面和状态栏显示未配置，基础本地功能继续可用。"),
            ("AI 网络失败", "返回明确错误并保留用户输入，不生成伪造建议。"),
            ("Sidecar 进程异常", "前端显示服务离线，用户可重启应用；日志用于定位。"),
            ("数据库迁移失败", "停止危险写入并提示备份/修复；迁移版本记录在 app_meta。"),
            ("排程冲突", "保留锁定事件，提示冲突或未排入任务，并允许用户手动调整。"),
            ("备份损坏", "恢复前校验失败则拒绝覆盖当前数据。"),
        ],
        [4.0, 11.0],
    )
    doc.add_heading("6．其它需求", level=1)
    bullets(
        doc,
        [
            "界面需要保持中文友好、页面结构清晰、错误状态可理解。",
            "课程交付应包含可安装程序、源码、测试结论和四份格式化文档。",
            "后续扩展可考虑移动端、云同步、多人共享日历和更细粒度的可观测日志。",
        ],
    )
    doc.save(path)


def build_design(path: Path, images: dict[str, Path]) -> None:
    doc = Document()
    setup_doc(doc)
    add_cover(doc, "四、概要设计说明书", "概要设计说明书")
    add_toc(doc, [("1", "引言"), ("2", "建立分析对象模型"), ("3", "提供交互界面的类"), ("4", "建立动态模型"), ("5", "数据结构设计"), ("6", "用户界面设计")])
    build_common_intro(doc, "概要设计")
    doc.add_heading("2．建立分析对象模型", level=1)
    para(doc, "系统采用分层架构：桌面端负责展示和交互，sidecar 负责 AI 会话与 RPC 调度，能力层统一暴露安全接口，领域层封装业务规则，存储层负责 SQLite 数据、迁移、备份和恢复。")
    add_picture(doc, images["architecture"], "图 2-1 系统分层架构")
    add_table(
        doc,
        ["对象/类", "职责"],
        [
            ("Task", "表示待办任务，包含标题、描述、时长、截止日期、状态、锁定、项目和父任务等信息。"),
            ("Event", "表示固定时间事件，包含开始/结束时间、地点、描述、锁定和可绑定任务。"),
            ("ScheduleBlock", "表示某日的时间安排，可关联任务或事件，并标记来源、锁定状态和说明。"),
            ("Reminder", "表示提醒实体，支持触发时间、状态、稍后提醒和关联任务/事件。"),
            ("ActionLog", "记录 capability 调用、actor、origin、输入输出和变更前状态，为撤销提供基础。"),
            ("Preference", "保存用户偏好、AI 配置状态、工作时间和学习习惯等本地设置。"),
            ("AISession/AiMessage", "保存 AI 会话和消息，为上下文延续、工具调用记录和调试提供依据。"),
            ("DreamInsight", "保存长期偏好洞察、置信度、证据和过期时间。"),
        ],
        [4.0, 11.0],
    )
    add_table(
        doc,
        ["服务类", "核心方法/能力"],
        [
            ("TaskService", "create、update、complete、defer、lock、delete、list。"),
            ("EventService", "create、update、delete、findConflicts、lock、list。"),
            ("ScheduleService", "planDay、planRange、rebalance、getBlocks、explain、analyzeQuality。"),
            ("ReminderService/Poller", "create、snooze、dismiss、pollDueReminders。"),
            ("UndoService", "revertAction、restore state_before、维护 undo_groups。"),
            ("SummaryService", "generateDailySummary、统计完成情况和计划质量。"),
            ("PreferenceService", "set/get preference、学习偏好信号。"),
            ("DreamOrchestrator", "汇总历史数据、调用模型分析、写入 dream_insights。"),
        ],
        [4.0, 11.0],
    )
    doc.add_heading("3．提供交互界面的类", level=1)
    add_table(
        doc,
        ["界面类/组件", "作用"],
        [
            ("TodayPage", "展示今日任务、计划块、快速添加、AI/Dream 今日建议和自动规划入口。"),
            ("TasksPage", "任务列表、筛选、创建/编辑/完成/延期/删除/锁定等操作。"),
            ("CalendarPage", "展示事件和日期视图，支持事件创建、编辑、冲突处理。"),
            ("AIPage", "提供 AI 会话、模型状态、上下文查看和工具调用反馈。"),
            ("AnalyticsPage", "展示日总结、计划质量、完成趋势和洞察信息。"),
            ("SettingsPage", "模型配置、API 状态、备份恢复、偏好设置。"),
            ("TaskEditModal/EventEditModal", "承载复杂编辑表单，避免主页面过载。"),
            ("BackupPanel/DreamInsightsPanel/DailySummaryPanel/BuddyWidget", "分别处理备份、AI 洞察、日总结和轻量提示反馈。"),
        ],
        [4.0, 11.0],
    )
    add_table(
        doc,
        ["Capability", "类别", "说明"],
        [
            ("task.*", "任务", "create/update/complete/defer/lock/delete/cancel/list。"),
            ("event.*", "事件", "create/update/lock/delete/list/find_conflicts。"),
            ("schedule.*", "排程", "plan_day/plan_range/rebalance/get_blocks/explain/analyze_quality。"),
            ("reminder.*", "提醒", "list/snooze。"),
            ("ai.*", "AI", "check_connectivity/chat/stream/get_context/delete_session。"),
            ("dream.*", "记忆", "run/status/get_insights。"),
            ("backup.*", "备份", "list/create/verify/restore/delete。"),
            ("history/undo/memory/preference", "支撑", "历史查询、撤销、清理 AI 记忆、偏好读写。"),
        ],
        [3.4, 2.4, 9.0],
    )
    doc.add_heading("4．建立动态模型", level=1)
    add_picture(doc, images["sequence"], "图 4-1 AI 日程规划时序图")
    para(doc, "动态模型的关键是把 AI 的开放式自然语言能力与系统的受控业务能力分开：模型只产生解释和工具调用意图，实际数据变更由 capability 层进行输入校验、权限控制、领域服务调用和日志记录。")
    add_table(
        doc,
        ["流程", "描述"],
        [
            ("任务创建流程", "用户或 AI 提交 task.create；能力层校验 title 等必填项；TaskService 写入 tasks；ActionLog 记录结果。"),
            ("日程规划流程", "ScheduleService 读取待办、事件和偏好，先保留锁定时间，再把可移动任务分配到空闲窗口。"),
            ("撤销流程", "用户选择历史动作；UndoService 读取 state_before；根据动作类型恢复任务/事件/计划块。"),
            ("提醒流程", "ReminderPoller 周期查询到期提醒；UI 展示提醒；用户选择稍后或解除并更新 reminders。"),
            ("Dream 洞察流程", "系统在满足冷启动阈值后汇总历史数据，调用模型分析偏好并写入 dream_insights。"),
        ],
        [4.0, 11.0],
    )
    doc.add_heading("5．数据结构设计", level=1)
    add_picture(doc, images["er"], "图 5-1 数据库 ER / 表结构关系图")
    add_table(
        doc,
        ["表名", "核心字段", "说明"],
        [
            ("app_meta", "key, value", "记录 schema_version 等元信息。"),
            ("tasks", "id, title, duration_minutes, due_date, status, locked, parent_task_id, project", "任务主表。"),
            ("task_recurrence_rules / task_relations / task_tags / task_reminders", "task_id, rule/tag/reminder", "任务的重复、依赖、标签和提醒扩展。"),
            ("events", "id, title, start_time, end_time, location, bound_task_id, locked", "日历事件主表。"),
            ("event_recurrence_rules / event_reminders", "event_id, rule/reminder", "事件重复和提醒扩展。"),
            ("schedule_blocks", "date, start_time, end_time, task_id, event_id, source, locked", "日程规划结果。"),
            ("reminders", "task_id, event_id, trigger_at, status, snooze_until", "统一提醒列表。"),
            ("action_logs / undo_groups", "capability, actor, input, result, state_before", "动作审计与撤销支撑。"),
            ("preferences / preference_signals", "key, value, scope, source", "偏好配置和学习信号。"),
            ("ai_sessions / ai_messages", "session_id, role, content, tool_calls", "AI 会话与消息历史。"),
            ("daily_summaries / dream_insights", "date/category/confidence/evidence", "总结与长期洞察。"),
        ],
        [4.0, 6.2, 5.0],
    )
    para(doc, "数据库设计采用外键约束和索引提升一致性与查询效率。例如 tasks 按 status、due_date、project 建索引，events 按 start_time/end_time 建索引，schedule_blocks 按 date 和 start/end 建索引，reminders 按 trigger_at/status 建索引，action_logs 按 capability、actor、idempotency_key 和 created_at 建索引。")
    doc.add_heading("6．用户界面设计", level=1)
    add_picture(doc, images["ui_flow"], "图 6-1 桌面端页面与交互流转")
    bullets(
        doc,
        [
            "信息架构：以“今日”为默认工作台，任务、日历、AI、统计、设置为主导航。",
            "交互原则：高频操作在当前页面完成，复杂编辑进入模态框；危险操作必须确认。",
            "状态表达：AI 未配置、sidecar 离线、网络失败等状态必须可见，不用假数据掩盖。",
            "可维护性：页面通过统一 RPC 客户端访问能力层，减少跨层耦合。",
            "可扩展性：后续可增加云同步、移动端或更多 AI 工具，而不破坏现有领域服务。",
        ],
    )
    doc.save(path)


def build_test(path: Path, images: dict[str, Path]) -> None:
    doc = Document()
    setup_doc(doc)
    add_cover(doc, "八、测试分析报告", "测试分析报告")
    add_toc(doc, [("1", "引言"), ("2", "测试计划执行情况"), ("3", "软件需求测试结论"), ("4", "评价")])
    build_common_intro(doc, "测试分析")
    doc.add_heading("2．测试计划执行情况", level=1)
    doc.add_heading("2.1 测试项目", level=2)
    add_table(
        doc,
        ["测试项目", "内容", "目的"],
        [
            ("类型检查", "运行 npm run typecheck，覆盖 capabilities、cli、domain、storage、ui-shared、desktop-tauri、runtime。", "发现类型不一致、接口不匹配和编译期错误。"),
            ("自动化单元/集成测试", "运行 npm test，覆盖 capability registry、CLI、领域服务、数据库、共享 UI 和 runtime AI 客户端。", "验证核心业务逻辑和关键边界。"),
            ("生产构建", "运行 npm run build，依次构建各工作空间与桌面前端。", "确认源码可构建为生产产物。"),
            ("桌面安装包检查", "检查 MSI 与 NSIS 安装包是否生成，资源是否随包发布。", "确认课程交付具备可安装形式。"),
            ("人工体验验证", "从用户角度检查页面导航、任务创建、自动规划、AI 状态和备份入口。", "验证真实使用流程，而不只停留在代码层。"),
        ],
        [3.4, 8.2, 4.0],
    )
    add_picture(doc, images["test_summary"], "图 2-1 桌面端运行界面与核心功能示意")
    doc.add_heading("2.2 测试机构和人员", level=2)
    add_table(doc, ["学号", "姓名", "测试职责"], [(sid, name, role) for sid, name, role in MEMBERS], [3.0, 2.6, 10.0])
    doc.add_heading("2.3 测试结果", level=2)
    add_table(
        doc,
        ["命令/检查项", "结果", "说明"],
        [
            ("npm run typecheck", "通过", "7 个工作空间完成 TypeScript 类型检查。"),
            ("npm test", "通过", "35 个自动化测试用例通过：capabilities 7、cli 1、domain 18、storage 6、ui-shared 1、runtime 2。"),
            ("npm run build", "通过", "各包完成 tsc 构建；桌面前端 Vite 生产构建成功。"),
            ("MSI 安装包", "存在", "EvolveFlow_0.1.0_x64_en-US.msi，约 41.10 MB。"),
            ("NSIS 安装程序", "存在", "EvolveFlow_0.1.0_x64-setup.exe，约 27.44 MB。"),
            ("AI 运行时", "通过设计验证", "sidecar 支持 Anthropic 兼容接口、真实连通性检查和工具调用；文档不暴露密钥。"),
        ],
        [4.0, 2.4, 9.0],
    )
    add_table(
        doc,
        ["用例编号", "测试点", "预期结果", "实际结果"],
        [
            ("TC-01", "数据库初始化", "创建数据库并完成 schema v1/v2 迁移。", "通过。"),
            ("TC-02", "任务创建与状态流转", "任务可创建、完成、延期、锁定并按条件查询。", "通过。"),
            ("TC-03", "事件与排程", "固定事件进入计划，任务可自动排入空闲时间。", "通过。"),
            ("TC-04", "锁定保护", "锁定任务/事件不被自动排程随意移动。", "通过。"),
            ("TC-05", "提醒", "提醒可创建、稍后提醒、解除。", "通过。"),
            ("TC-06", "撤销", "任务创建和完成等操作可按日志恢复。", "通过。"),
            ("TC-07", "能力注册", "生产能力清单完整，未知能力拒绝，幂等调用返回缓存结果。", "通过。"),
            ("TC-08", "AI 客户端兼容性", "DeepSeek/Anthropic 兼容路径处理 cache_control、header 和连通性。", "通过。"),
            ("TC-09", "构建发布", "源码可构建，安装包生成。", "通过。"),
        ],
        [2.0, 4.2, 5.0, 4.0],
    )
    doc.add_heading("3．软件需求测试结论", level=1)
    add_table(
        doc,
        ["需求", "覆盖测试", "结论"],
        [
            ("任务管理 FR-01", "TaskService、capability task.create/task.update/task.complete 等。", "满足。"),
            ("事件管理 FR-02", "EventService 与 schedule 相关测试。", "满足。"),
            ("自动排程 FR-03", "ScheduleService planDay、锁定保护、解释能力。", "满足。"),
            ("AI Agent FR-04", "runtime AI 客户端测试、sidecar 设计验证、状态处理。", "满足课程交付要求。"),
            ("撤销与历史 FR-05", "UndoService、action_logs。", "满足。"),
            ("提醒 FR-06", "ReminderService。", "满足。"),
            ("备份恢复 FR-07", "backup capability 与发布检查。", "基本满足，建议公开发布前扩展破坏性恢复测试。"),
            ("统计与总结 FR-08", "Summary/Dream 面板和能力设计检查。", "满足课程展示要求。"),
            ("配置与安全 FR-09", "api_key.status、设置页状态检查、文档脱敏。", "满足。"),
        ],
        [3.8, 7.0, 4.0],
    )
    doc.add_heading("4．评价", level=1)
    doc.add_heading("4.1 软件能力", level=2)
    bullets(
        doc,
        [
            "系统已经具备任务、事件、排程、提醒、撤销、备份、AI 会话和桌面打包等完整课程项目能力。",
            "架构分层明确，核心业务逻辑位于领域服务和 capability 层，前端与 AI 均通过统一接口访问。",
            "本地优先设计兼顾隐私与离线能力，AI 不可用时基础功能仍能工作。",
            "测试结果显示当前代码可以通过类型检查、自动化测试和生产构建。",
        ],
    )
    doc.add_heading("4.2 缺陷和限制", level=2)
    bullets(
        doc,
        [
            "AI 功能依赖外部模型服务，网络或额度异常会影响智能对话和 Dream 分析。",
            "当前课程版本主要面向 Windows 桌面，未覆盖移动端和多人云同步场景。",
            "公开商业发布前仍应补充代码签名证书、自动更新签名密钥和更完整的端到端回归测试。",
            "桌面端目前以人工体验验证为主，后续可引入 Playwright/Tauri 自动化 UI 测试。",
        ],
    )
    doc.add_heading("4.3 建议", level=2)
    bullets(
        doc,
        [
            "增加跨版本升级与数据库迁移压力测试，确保长期使用稳定。",
            "补充 AI 工具调用的审计界面，使用户更直观看到模型做了什么。",
            "完善备份恢复的异常演练，例如损坏备份、磁盘不足和恢复中断。",
            "在公开发布版本中加入签名、更新、崩溃报告和可选的数据导出向导。",
        ],
    )
    doc.add_heading("4.4 测试结论", level=2)
    para(doc, "截至 2026年6月1日，EvolveFlow 课程提交版本通过类型检查、自动化测试和生产构建，已生成 Windows 安装包。结合人工体验验证和文档复核，项目具备课程验收与演示条件。剩余限制主要属于公开发布前的工程加固项，不影响本次课程交付目标。")
    doc.save(path)


def generate_images(image_dir: Path) -> dict[str, Path]:
    image_dir.mkdir(parents=True, exist_ok=True)
    images = {
        "architecture": image_dir / "01_项目总体架构图.png",
        "use_case": image_dir / "02_需求用例图.png",
        "er": image_dir / "03_数据库ER图.png",
        "sequence": image_dir / "04_AI日程规划时序图.png",
        "test_summary": image_dir / "05_测试结果概览.png",
        "ui_flow": image_dir / "06_界面页面流转图.png",
    }
    imagegen_sources = {
        "architecture": image_dir / "imagegen_01_系统总体架构.png",
        "use_case": image_dir / "imagegen_05_需求用例图.png",
        "er": image_dir / "imagegen_02_本地优先数据模型.png",
        "sequence": image_dir / "imagegen_03_AI日程规划流程.png",
        "test_summary": image_dir / "imagegen_04_桌面运行界面.png",
        "ui_flow": image_dir / "imagegen_04_桌面运行界面.png",
    }
    missing = [str(path) for path in imagegen_sources.values() if not path.exists()]
    if missing:
        raise FileNotFoundError("缺少 imagegen 生成的图表素材，拒绝回退到 Python 绘图：\n" + "\n".join(missing))
    for key, source in imagegen_sources.items():
        shutil.copy2(source, images[key])
    return images


def backup_existing(docx_dir: Path) -> Path:
    backup_dir = docx_dir / f"原始备份_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    backup_dir.mkdir(parents=True, exist_ok=True)
    for path in sorted(docx_dir.glob("*.docx")):
        shutil.copy2(path, backup_dir / path.name)
    return backup_dir


def main() -> int:
    if len(sys.argv) != 2:
        print("Usage: generate_course_submission_docs.py <docx_dir>")
        return 2
    docx_dir = Path(sys.argv[1])
    docx_dir.mkdir(parents=True, exist_ok=True)
    backup_dir = backup_existing(docx_dir)
    image_dir = docx_dir / "图表"
    images = generate_images(image_dir)
    build_plan(docx_dir / "01_项目开发计划.docx", images)
    build_srs(docx_dir / "02_需求规格说明书.docx", images)
    build_design(docx_dir / "03_系统设计说明书.docx", images)
    build_test(docx_dir / "04_测试报告.docx", images)
    print(f"Backed up original DOCX files to: {backup_dir}")
    print(f"Generated diagrams in: {image_dir}")
    for path in sorted(docx_dir.glob("*.docx")):
        print(f"Wrote: {path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
